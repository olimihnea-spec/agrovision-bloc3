"""
BLOC 3 — Deep Learning YOLOv8, Ziua 15
Batch procesare: analiza automata a unui folder intreg cu imagini drone
Autor: Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj

CONCEPT CHEIE:
  Batch procesare = procesezi zeci de parcele dintr-un singur click, nu una cate una.
  Inspector incarca 20+ imagini drone → sistemul le analizeaza pe toate automat →
  genereaza raport consolidat, Excel cu toate rezultatele, si rapoarte individuale per parcela.

  Avantaj operational APIA:
    - Fara batch: 20 parcele x 10 min/parcela = 200 minute (3.3 ore)
    - Cu batch:   20 parcele procesate automat in < 30 secunde
    - Economie: ~95% din timpul de procesare al inspectorului

  Date reale integrate (antrenament AGROVISION):
    - Model: best_v1_mAP083_20260403.pt | mAP50=0.829
    - Parcele LPIS Gorj: GJ_78258-1675, GJ_79157-348, GJ_79237-628,
      GJ_79308-489, GJ_79406-641, GJ_79406-924, GJ_79834-9533,
      GJ_80123-1004, GJ_80123-3737, GJ_80980-2611
    - Reg. UE 2021/2116: prag vegetatie 50% | neconformitate → recuperare plati
"""

import streamlit as st
import numpy as np
from PIL import Image, ImageDraw
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from io import BytesIO
import random
import zipfile
from datetime import date, datetime
from collections import Counter, defaultdict

# ── Word / Excel ──────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

st.set_page_config(page_title="Batch Procesare — Ziua 15", layout="wide")

# ── CSS ────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
.bloc3-header {
    background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
    padding: 1.5rem 2rem; border-radius: 12px;
    margin-bottom: 1.5rem; border-left: 5px solid #e94560;
}
.bloc3-header h1 { color: #e94560; margin: 0; font-size: 1.6rem; }
.bloc3-header p  { color: #a8b2d8; margin: 0.3rem 0 0 0; font-size: 0.9rem; }
.concept-box {
    background: #0f3460; border: 1px solid #e94560;
    border-radius: 8px; padding: 1rem 1.2rem; margin: 1rem 0;
    color: #a8b2d8; font-size: 0.88rem;
}
.concept-box b { color: #e94560; }
.ok-box {
    background: #0d2b0d; border: 1px solid #27ae60;
    border-radius: 8px; padding: 0.8rem 1rem; color: #7dcea0; margin: 0.4rem 0;
}
.warn-box {
    background: #2d1b00; border: 1px solid #e67e22;
    border-radius: 8px; padding: 0.8rem 1rem; color: #f39c12; margin: 0.4rem 0;
}
.err-box {
    background: #2d0000; border: 1px solid #e74c3c;
    border-radius: 8px; padding: 0.8rem 1rem; color: #f1948a; margin: 0.4rem 0;
}
.metric-card {
    background: #16213e; border: 1px solid #0f3460;
    border-radius: 10px; padding: 1.2rem; text-align: center; margin: 0.3rem 0;
}
.metric-card .val { font-size: 2rem; font-weight: bold; color: #e94560; }
.metric-card .lbl { font-size: 0.8rem; color: #a8b2d8; margin-top: 0.3rem; }
.metric-card.verde .val { color: #27ae60; }
.metric-card.portocaliu .val { color: #e67e22; }
.metric-card.albastru .val { color: #3498db; }
.parcela-card {
    background: #16213e; border: 1px solid #0f3460;
    border-radius: 8px; padding: 1rem; margin: 0.5rem 0;
}
.parcela-card h4 { color: #e94560; margin: 0 0 0.5rem 0; font-size: 1rem; }
.stat-row { display: flex; gap: 1rem; flex-wrap: wrap; margin: 0.3rem 0; }
.stat-item { color: #a8b2d8; font-size: 0.85rem; }
.stat-item b { color: white; }
.sectiune-titlu {
    background: #0f3460; color: #e94560;
    padding: 0.5rem 1rem; border-radius: 6px;
    font-weight: bold; font-size: 1rem; margin: 1rem 0 0.5rem 0;
}
</style>
""", unsafe_allow_html=True)

# ── Header ─────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="bloc3-header">
  <h1>Ziua 15 — Batch Procesare: Folder Intreg cu Imagini Drone</h1>
  <p>Bloc 3 YOLOv8 | Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj</p>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div class="concept-box">
<b>CONCEPTE CHEIE AZI:</b><br>
<b>Batch procesare</b> = un singur click analizeaza zeci de imagini drone in paralel<br>
<b>accept_multiple_files=True</b> = st.file_uploader accepta mai multe fisiere simultan<br>
<b>st.progress()</b> = bara de progres vizuala in timp ce se proceseaza fiecare imagine<br>
<b>Raport consolidat</b> = un singur document Word/Excel cu toate parcelele analizate<br>
<b>ZIP in memorie</b> = zipfile.ZipFile(BytesIO, "w") fara niciun fisier pe disk<br>
<b>Prag PAC (Reg. UE 2021/2116)</b> = vegetatie &lt; 50% → NECONFORM → recuperare plati
</div>
""", unsafe_allow_html=True)

# ── Flux vizual ────────────────────────────────────────────────────────────────
st.header("Fluxul batch APIA")
cols_flux = st.columns(6)
pasi = [
    ("1", "Upload multiple", "Toate imaginile dintr-o data"),
    ("2", "Configurare", "Parametri inspector + model"),
    ("3", "Procesare batch", "Detectie automata per imagine"),
    ("4", "Statistici", "Centralizator toate parcelele"),
    ("5", "Raport Word", "Document oficial consolidat"),
    ("6", "Excel + ZIP", "Export complet pentru audit"),
]
for col, (nr, titlu, desc) in zip(cols_flux, pasi):
    with col:
        st.markdown(f"""
        <div style="background:#16213e;border-left:4px solid #e94560;
             border-radius:0 8px 8px 0;padding:0.7rem 0.9rem;margin:0.3rem 0;color:#a8b2d8;font-size:0.82rem;">
        <b style="color:white;">Pasul {nr}: {titlu}</b><br>{desc}
        </div>""", unsafe_allow_html=True)

st.markdown("---")

# ══════════════════════════════════════════════════════════════════════════════
# SECTIUNEA 1 — Configurare generala
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="sectiune-titlu">Sectiunea 1 — Configurare sesiune batch</div>', unsafe_allow_html=True)

col_c1, col_c2, col_c3 = st.columns(3)
with col_c1:
    inspector    = st.text_input("Inspector APIA", "Prof. Asoc. Dr. Oliviu Mihnea Gamulescu")
    unitate      = st.text_input("Unitate APIA", "Centrul Judetean Gorj")
with col_c2:
    conf_thresh  = st.slider("Confidence threshold", 0.10, 0.90, 0.45, 0.05,
                             help="Detectii sub acest prag sunt ignorate")
    prag_veg_pac = st.slider("Prag vegetatie PAC (%)", 30, 70, 50, 5,
                             help="Sub acest prag = neconform (Reg. UE 2021/2116 default=50%)")
with col_c3:
    class_names_in = st.text_input("Clase model YOLOv8",
                                   "vegetatie,sol_gol,apa",
                                   help="Separat prin virgula, in ordinea din data.yaml")
    class_names    = [c.strip() for c in class_names_in.split(",") if c.strip()]
    data_sesiune   = st.date_input("Data sesiunii de control", value=date.today())

# ── Date parcele (pot fi personalizate) ──────────────────────────────────────
st.markdown("**Date parcele** — completati tabelul sau lasati valorile implicite (parcele reale LPIS Gorj):")

PARCELE_DEFAULT = [
    {"cod": "GJ_78258-1675", "fermier": "Ionescu Marin",     "cnp": "1560712182011", "suprafata": 3.42, "cultura": "Grau"},
    {"cod": "GJ_79157-348",  "fermier": "Popescu Ion",       "cnp": "1760312182456", "suprafata": 2.45, "cultura": "Porumb"},
    {"cod": "GJ_79237-628",  "fermier": "Dumitrescu Vasile", "cnp": "1680512182233", "suprafata": 5.10, "cultura": "Floarea-soarelui"},
    {"cod": "GJ_79308-489",  "fermier": "Stanescu Maria",    "cnp": "2710215182301", "suprafata": 1.80, "cultura": "Rapita"},
    {"cod": "GJ_79406-641",  "fermier": "Gheorghiu Aurel",   "cnp": "1650804182512", "suprafata": 4.20, "cultura": "Orz"},
    {"cod": "GJ_79406-924",  "fermier": "Constantin Elena",  "cnp": "2740110182417", "suprafata": 6.75, "cultura": "Lucerna"},
    {"cod": "GJ_79834-9533", "fermier": "Marin Gheorghe",    "cnp": "1780920182388", "suprafata": 2.30, "cultura": "Pasune permanenta"},
    {"cod": "GJ_80123-1004", "fermier": "Popa Nicolae",      "cnp": "1820315182491", "suprafata": 8.60, "cultura": "Grau"},
    {"cod": "GJ_80123-3737", "fermier": "Dima Florin",       "cnp": "1790601182255", "suprafata": 3.15, "cultura": "Porumb"},
    {"cod": "GJ_80980-2611", "fermier": "Olteanu Traian",    "cnp": "1710428182344", "suprafata": 7.40, "cultura": "Floarea-soarelui"},
]

with st.expander("Editeaza datele parcelelor (optional — implicit: 10 parcele reale LPIS Gorj)"):
    st.info("Poti edita direct in tabel. Modificarile sunt active imediat.")
    import pandas as pd
    df_parcele_edit = pd.DataFrame(PARCELE_DEFAULT)
    df_editat = st.data_editor(
        df_parcele_edit,
        column_config={
            "cod":       st.column_config.TextColumn("Cod LPIS", width="medium"),
            "fermier":   st.column_config.TextColumn("Fermier", width="large"),
            "cnp":       st.column_config.TextColumn("CNP/CUI", width="medium"),
            "suprafata": st.column_config.NumberColumn("Suprafata (ha)", min_value=0.1, max_value=500.0, format="%.2f"),
            "cultura":   st.column_config.SelectboxColumn("Cultura", options=[
                "Grau", "Porumb", "Floarea-soarelui", "Rapita", "Orz",
                "Lucerna", "Pasune permanenta", "Legume", "Livada"]),
        },
        num_rows="dynamic",
        use_container_width=True,
        key="editor_parcele"
    )
    PARCELE_DEFAULT = df_editat.to_dict("records")

st.markdown("---")

# ══════════════════════════════════════════════════════════════════════════════
# SECTIUNEA 2 — Upload imagini
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="sectiune-titlu">Sectiunea 2 — Upload imagini drone</div>', unsafe_allow_html=True)

uploaded_files = st.file_uploader(
    "Incarca imaginile drone (.jpg / .png) — poti selecta mai multe simultan",
    type=["jpg", "jpeg", "png"],
    accept_multiple_files=True,
    help="Ctrl+click sau Shift+click pentru selectie multipla"
)

n_imagini   = len(uploaded_files) if uploaded_files else 0
n_parcele_d = len(PARCELE_DEFAULT)

col_info1, col_info2, col_info3 = st.columns(3)
with col_info1:
    st.markdown(f'<div class="metric-card albastru"><div class="val">{n_imagini}</div><div class="lbl">Imagini incarcate</div></div>', unsafe_allow_html=True)
with col_info2:
    st.markdown(f'<div class="metric-card"><div class="val">{n_parcele_d}</div><div class="lbl">Parcele configurate</div></div>', unsafe_allow_html=True)
with col_info3:
    n_proc = min(n_imagini, n_parcele_d) if n_imagini > 0 else n_parcele_d
    st.markdown(f'<div class="metric-card portocaliu"><div class="val">{n_proc}</div><div class="lbl">Parcele de procesat</div></div>', unsafe_allow_html=True)

if n_imagini == 0:
    st.info("Nu ai incarcat imagini — se vor genera imagini sintetice demonstrative pentru toate parcelele configurate.")
elif n_imagini < n_parcele_d:
    st.warning(f"Ai {n_imagini} imagini pentru {n_parcele_d} parcele. Primele {n_imagini} parcele vor folosi imaginile tale, restul vor folosi imagini sintetice.")

st.markdown("---")

# ══════════════════════════════════════════════════════════════════════════════
# BUTON PORNIRE BATCH
# ══════════════════════════════════════════════════════════════════════════════
btn_batch = st.button(
    f"Porneste analiza batch — {n_proc} parcele",
    type="primary",
    use_container_width=True
)

# ══════════════════════════════════════════════════════════════════════════════
# PROCESARE BATCH
# ══════════════════════════════════════════════════════════════════════════════

def genereaza_imagine_sintetica(seed_val: int) -> tuple:
    """Genereaza o imagine drone sintetica reproductibila per parcela."""
    rng = random.Random(seed_val)
    np_rng = np.random.default_rng(seed_val)
    W, H = 640, 480
    arr = np.zeros((H, W, 3), dtype=np.uint8)
    # fond sol
    arr[:, :] = [rng.randint(90, 120), rng.randint(70, 95), rng.randint(40, 65)]
    # zone vegetatie
    n_zone = rng.randint(6, 16)
    for _ in range(n_zone):
        x0 = rng.randint(0, W-150); y0 = rng.randint(0, H-100)
        x1 = min(x0 + rng.randint(60, 180), W)
        y1 = min(y0 + rng.randint(50, 140), H)
        arr[y0:y1, x0:x1] = [rng.randint(15, 40),
                              rng.randint(85, 150),
                              rng.randint(10, 35)]
    # zona apa (uneori)
    if rng.random() < 0.35:
        ax0 = rng.randint(20, W-120); ay0 = rng.randint(10, H-80)
        arr[ay0:ay0+rng.randint(40,80), ax0:ax0+rng.randint(80,160)] = [15, 55, rng.randint(110,155)]
    # zgomot
    noise = np_rng.integers(-12, 12, arr.shape, dtype=np.int16)
    arr = np.clip(arr.astype(np.int16) + noise, 0, 255).astype(np.uint8)
    return Image.fromarray(arr), W, H


def simuleaza_detectie(img_pil: Image.Image, W: int, H: int,
                       n_cls: int, conf_thresh: float,
                       seed_val: int) -> tuple:
    """
    Simuleaza inferenta YOLOv8 cu rezultate reproductibile per parcela.
    In productie: inlocuieste cu model.predict(img_pil, conf=conf_thresh).
    """
    rng = random.Random(seed_val + 1000)
    weights_cls = [0.55, 0.32, 0.13][:n_cls]
    while len(weights_cls) < n_cls:
        weights_cls.append(0.08)
    total_w = sum(weights_cls)
    weights_cls = [w / total_w for w in weights_cls]

    CULORI_HEX = ["#27ae60", "#e67e22", "#3498db", "#e94560", "#9b59b6"]
    CULORI_RGB  = [(39,174,96), (230,126,34), (52,152,219), (233,69,96), (155,89,182)]

    detectii = []
    for _ in range(rng.randint(8, 20)):
        cls_id = rng.choices(range(n_cls), weights=weights_cls)[0]
        conf   = round(rng.uniform(0.28, 0.97), 3)
        if conf < conf_thresh:
            continue
        xc = rng.uniform(0.07, 0.93)
        yc = rng.uniform(0.07, 0.93)
        w  = rng.uniform(0.06, 0.32)
        h  = rng.uniform(0.05, 0.28)
        detectii.append({"cls": cls_id, "conf": conf,
                         "xc": xc, "yc": yc, "w": w, "h": h})

    if not detectii:
        detectii.append({"cls": 0, "conf": 0.71,
                         "xc": 0.5, "yc": 0.5, "w": 0.42, "h": 0.38})

    # Deseneaza BBox
    img_draw = img_pil.copy()
    draw = ImageDraw.Draw(img_draw)
    aria_totala = 0.0
    aria_per_cls = defaultdict(float)

    for det in detectii:
        cls_id = det["cls"]
        x1 = int((det["xc"] - det["w"] / 2) * W)
        y1 = int((det["yc"] - det["h"] / 2) * H)
        x2 = int((det["xc"] + det["w"] / 2) * W)
        y2 = int((det["yc"] + det["h"] / 2) * H)
        color_hex = CULORI_HEX[cls_id % len(CULORI_HEX)]
        draw.rectangle([x1, y1, x2, y2], outline=color_hex, width=3)
        lbl = f"{class_names[cls_id] if cls_id < len(class_names) else cls_id} {det['conf']:.2f}"
        draw.rectangle([x1, y1 - 18, x1 + len(lbl) * 7 + 4, y1], fill=color_hex)
        draw.text((x1 + 2, y1 - 16), lbl, fill="white")
        aria = det["w"] * det["h"]
        aria_per_cls[cls_id] += aria
        aria_totala += aria

    pct_cls = {}
    if aria_totala > 0:
        for i in range(n_cls):
            pct_cls[i] = aria_per_cls[i] / aria_totala * 100
    else:
        for i in range(n_cls):
            pct_cls[i] = 100.0 / n_cls

    return img_draw, detectii, pct_cls


def evalueaza_pac(pct_cls: dict, n_cls: int, prag_veg: float) -> tuple:
    """Evalueaza conformitatea PAC conform Reg. UE 2021/2116."""
    pct_veg     = pct_cls.get(0, 0)
    pct_sol_gol = pct_cls.get(1, 0) if n_cls > 1 else 0
    pct_apa     = pct_cls.get(2, 0) if n_cls > 2 else 0

    probleme = []
    if pct_veg < prag_veg:
        probleme.append(f"Vegetatie {pct_veg:.1f}% < {prag_veg}% minim PAC")
    if pct_sol_gol > 40:
        probleme.append(f"Sol gol {pct_sol_gol:.1f}% > 40% (subutilizare)")

    if not probleme:
        return "CONFORM", "#27ae60", probleme, pct_veg, pct_sol_gol, pct_apa
    else:
        return "NECONFORM", "#e74c3c", probleme, pct_veg, pct_sol_gol, pct_apa


def img_to_bytes(img_pil: Image.Image, dpi: int = 150) -> BytesIO:
    buf = BytesIO()
    img_pil.save(buf, format="PNG", dpi=(dpi, dpi))
    buf.seek(0)
    return buf


def fig_to_bytes(fig, dpi: int = 150) -> BytesIO:
    buf = BytesIO()
    fig.savefig(buf, dpi=dpi, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    buf.seek(0)
    return buf


# ── Functie helper: culoare celula Excel ──────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    fill = PatternFill(start_color=hex_color.lstrip("#"),
                       end_color=hex_color.lstrip("#"),
                       fill_type="solid")
    cell.fill = fill


# ══════════════════════════════════════════════════════════════════════════════
if btn_batch:
    n_cls   = len(class_names)
    parcele = PARCELE_DEFAULT[:n_proc]

    rezultate = []          # lista de dictionare per parcela
    imagini_annotate = {}   # {cod_parcela: BytesIO imagine}

    # ── Bara de progres ───────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown('<div class="sectiune-titlu">Sectiunea 3 — Procesare in curs...</div>', unsafe_allow_html=True)

    progress_bar = st.progress(0, text="Initializare...")
    status_text  = st.empty()

    for idx, parcela in enumerate(parcele):
        cod      = parcela["cod"]
        fermier  = parcela["fermier"]
        sup_ha   = float(parcela.get("suprafata", 2.0))
        cultura  = parcela.get("cultura", "Grau")

        status_text.markdown(f"**Procesez parcela {idx+1}/{n_proc}:** `{cod}` — {fermier}")

        # Imagine: reala daca exista, sintetica altfel
        seed_val = hash(cod) % 99999
        if uploaded_files and idx < len(uploaded_files):
            try:
                img_pil = Image.open(uploaded_files[idx]).convert("RGB")
                W, H = img_pil.size
                if max(W, H) > 1280:
                    f = 1280 / max(W, H)
                    img_pil = img_pil.resize((int(W*f), int(H*f)), Image.LANCZOS)
                    W, H = img_pil.size
                sursa_imagine = "reala"
            except Exception:
                img_pil, W, H = genereaza_imagine_sintetica(seed_val)
                sursa_imagine = "sintetica"
        else:
            img_pil, W, H = genereaza_imagine_sintetica(seed_val)
            sursa_imagine = "sintetica"

        # Detectie
        img_annotata, detectii, pct_cls = simuleaza_detectie(
            img_pil, W, H, n_cls, conf_thresh, seed_val
        )

        # Evaluare PAC
        concluzie, culoare_c, probleme, pct_veg, pct_sol, pct_apa = evalueaza_pac(
            pct_cls, n_cls, prag_veg_pac
        )

        # Salveaza imaginea annotata
        imagini_annotate[cod] = img_to_bytes(img_annotata, dpi=120)

        # Rezultat parcela
        rezultat = {
            "cod":          cod,
            "fermier":      fermier,
            "cnp":          parcela.get("cnp", ""),
            "suprafata_ha": sup_ha,
            "cultura":      cultura,
            "n_detectii":   len(detectii),
            "pct_veg":      round(pct_veg, 1),
            "pct_sol":      round(pct_sol, 1),
            "pct_apa":      round(pct_apa, 1),
            "concluzie":    concluzie,
            "culoare":      culoare_c,
            "probleme":     probleme,
            "pct_cls":      pct_cls,
            "sursa":        sursa_imagine,
        }
        rezultate.append(rezultat)

        progress_bar.progress((idx + 1) / n_proc,
                              text=f"Procesat {idx+1}/{n_proc} ({concluzie})")

    status_text.empty()
    progress_bar.empty()
    st.success(f"Batch complet: {n_proc} parcele procesate in < 5 secunde!")

    # ════════════════════════════════════════════════════════════════════════
    # SECTIUNEA 4 — Statistici centralizate
    # ════════════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.markdown('<div class="sectiune-titlu">Sectiunea 4 — Statistici centralizate</div>', unsafe_allow_html=True)

    n_conforme   = sum(1 for r in rezultate if r["concluzie"] == "CONFORM")
    n_neconforme = n_proc - n_conforme
    pct_conf     = n_conforme / n_proc * 100 if n_proc > 0 else 0
    total_ha     = sum(r["suprafata_ha"] for r in rezultate)
    ha_neconf    = sum(r["suprafata_ha"] for r in rezultate if r["concluzie"] == "NECONFORM")
    medie_veg    = sum(r["pct_veg"] for r in rezultate) / n_proc if n_proc > 0 else 0

    col_s1, col_s2, col_s3, col_s4, col_s5 = st.columns(5)
    with col_s1:
        st.markdown(f'<div class="metric-card albastru"><div class="val">{n_proc}</div><div class="lbl">Total parcele</div></div>', unsafe_allow_html=True)
    with col_s2:
        st.markdown(f'<div class="metric-card verde"><div class="val">{n_conforme}</div><div class="lbl">Conforme PAC</div></div>', unsafe_allow_html=True)
    with col_s3:
        st.markdown(f'<div class="metric-card"><div class="val" style="color:#e74c3c">{n_neconforme}</div><div class="lbl">Neconforme PAC</div></div>', unsafe_allow_html=True)
    with col_s4:
        st.markdown(f'<div class="metric-card portocaliu"><div class="val">{ha_neconf:.2f}</div><div class="lbl">Ha neconforme</div></div>', unsafe_allow_html=True)
    with col_s5:
        st.markdown(f'<div class="metric-card verde"><div class="val">{medie_veg:.1f}%</div><div class="lbl">Medie vegetatie</div></div>', unsafe_allow_html=True)

    # ── Grafice statistici ───────────────────────────────────────────────────
    col_g1, col_g2 = st.columns(2)

    with col_g1:
        fig_pie, ax_pie = plt.subplots(figsize=(5, 4))
        fig_pie.patch.set_facecolor("#16213e")
        ax_pie.set_facecolor("#16213e")
        valori_pie  = [n_conforme, n_neconforme]
        etichete    = [f"CONFORME\n({n_conforme})", f"NECONFORME\n({n_neconforme})"]
        culori_pie  = ["#27ae60", "#e74c3c"]
        explode     = (0.05, 0.05)
        wedges, texts, autotexts = ax_pie.pie(
            valori_pie, labels=etichete, colors=culori_pie,
            autopct="%1.0f%%", explode=explode, startangle=90,
            textprops={"color": "white", "fontsize": 9}
        )
        for at in autotexts:
            at.set_fontsize(10); at.set_color("white")
        ax_pie.set_title("Distributie conformitate PAC", color="white", fontsize=10, pad=10)
        plt.tight_layout()
        st.image(fig_to_bytes(fig_pie, dpi=120), use_container_width=True)
        plt.close(fig_pie)

    with col_g2:
        fig_bar, ax_bar = plt.subplots(figsize=(7, 4))
        fig_bar.patch.set_facecolor("#16213e")
        ax_bar.set_facecolor("#16213e")
        coduri_scurte = [r["cod"].split("_")[-1] for r in rezultate]
        pcts_veg      = [r["pct_veg"] for r in rezultate]
        culori_bar    = ["#27ae60" if r["concluzie"] == "CONFORM" else "#e74c3c"
                         for r in rezultate]
        bars = ax_bar.bar(range(len(rezultate)), pcts_veg, color=culori_bar,
                          edgecolor="#0f3460", linewidth=0.8)
        ax_bar.axhline(prag_veg_pac, color="#e67e22", linestyle="--",
                       linewidth=1.5, label=f"Prag PAC {prag_veg_pac}%")
        ax_bar.set_xticks(range(len(rezultate)))
        ax_bar.set_xticklabels(coduri_scurte, rotation=45, ha="right",
                               color="white", fontsize=7)
        ax_bar.set_ylabel("% Vegetatie", color="white", fontsize=9)
        ax_bar.set_title("Vegetatie per parcela (verde=conform, rosu=neconform)",
                         color="white", fontsize=9, pad=8)
        ax_bar.tick_params(colors="white")
        for spine in ax_bar.spines.values():
            spine.set_edgecolor("#0f3460")
        ax_bar.set_facecolor("#16213e")
        patch_conf  = mpatches.Patch(color="#27ae60", label="CONFORM")
        patch_nconf = mpatches.Patch(color="#e74c3c", label="NECONFORM")
        ax_bar.legend(handles=[patch_conf, patch_nconf,
                                plt.Line2D([0],[0], color="#e67e22",
                                           linestyle="--", label=f"Prag {prag_veg_pac}%")],
                      facecolor="#0f3460", labelcolor="white", fontsize=8)
        plt.tight_layout()
        st.image(fig_to_bytes(fig_bar, dpi=120), use_container_width=True)
        plt.close(fig_bar)

    # ── Tabel centralizator ──────────────────────────────────────────────────
    st.markdown("**Tabel centralizator rezultate:**")
    import pandas as pd
    df_rez = pd.DataFrame([{
        "Cod LPIS":        r["cod"],
        "Fermier":         r["fermier"],
        "Suprafata (ha)":  r["suprafata_ha"],
        "Cultura":         r["cultura"],
        "Detectii":        r["n_detectii"],
        "Vegetatie (%)":   r["pct_veg"],
        "Sol gol (%)":     r["pct_sol"],
        "Apa (%)":         r["pct_apa"],
        "Concluzie PAC":   r["concluzie"],
        "Imagine":         r["sursa"],
    } for r in rezultate])

    def color_concluzie(val):
        if val == "CONFORM":
            return "background-color: #0d2b0d; color: #7dcea0; font-weight: bold"
        else:
            return "background-color: #2d0000; color: #f1948a; font-weight: bold"

    def color_veg(val):
        try:
            v = float(val)
            if v >= prag_veg_pac:
                return "color: #7dcea0"
            else:
                return "color: #f1948a; font-weight: bold"
        except Exception:
            return ""

    st.dataframe(
        df_rez.style
              .applymap(color_concluzie, subset=["Concluzie PAC"])
              .applymap(color_veg, subset=["Vegetatie (%)"]),
        use_container_width=True,
        height=350
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTIUNEA 5 — Detalii per parcela + imagini annotate
    # ════════════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.markdown('<div class="sectiune-titlu">Sectiunea 5 — Detalii per parcela (imagini annotate)</div>', unsafe_allow_html=True)

    cols_per_row = 2
    for i in range(0, len(rezultate), cols_per_row):
        cols = st.columns(cols_per_row)
        for j, col in enumerate(cols):
            if i + j >= len(rezultate):
                break
            r = rezultate[i + j]
            with col:
                culoare_border = "#27ae60" if r["concluzie"] == "CONFORM" else "#e74c3c"
                st.markdown(f"""
                <div style="background:#16213e;border:1px solid {culoare_border};
                     border-radius:8px;padding:0.8rem;margin:0.3rem 0;">
                <h4 style="color:{culoare_border};margin:0 0 0.4rem 0;font-size:0.95rem;">
                  {r['cod']} — {r['concluzie']}</h4>
                <span style="color:#a8b2d8;font-size:0.82rem;">
                  {r['fermier']} | {r['suprafata_ha']} ha | {r['cultura']}<br>
                  Veg: <b style="color:#7dcea0">{r['pct_veg']}%</b> |
                  Sol: <b style="color:#e67e22">{r['pct_sol']}%</b> |
                  Apa: <b style="color:#3498db">{r['pct_apa']}%</b> |
                  Detectii: <b style="color:white">{r['n_detectii']}</b>
                </span>
                </div>""", unsafe_allow_html=True)

                img_bytes = imagini_annotate[r["cod"]]
                img_bytes.seek(0)
                st.image(img_bytes, caption=f"{r['cod']} — imagine {r['sursa']}",
                         use_container_width=True)

                if r["probleme"]:
                    for prob in r["probleme"]:
                        st.markdown(f'<div class="warn-box" style="font-size:0.8rem;">{prob}</div>',
                                    unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════════════════════
    # SECTIUNEA 6 — Exporturi
    # ════════════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.markdown('<div class="sectiune-titlu">Sectiunea 6 — Export documente oficiale</div>', unsafe_allow_html=True)

    col_ex1, col_ex2, col_ex3 = st.columns(3)

    # ── Helper: setare latime coloane Word ──────────────────────────────────
    def set_col_width(table, col_idx, width_cm):
        for row in table.rows:
            cell = row.cells[col_idx]
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW  = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(width_cm * 567)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)

    # ── 1. RAPORT WORD CONSOLIDAT ──────────────────────────────────────────
    with col_ex1:
        if st.button("Genereaza Raport Word Consolidat", type="primary", use_container_width=True):
            doc = Document()
            for section in doc.sections:
                section.top_margin    = Cm(2.5)
                section.bottom_margin = Cm(2.5)
                section.left_margin   = Cm(2.5)
                section.right_margin  = Cm(2.5)

            # Antet document
            p_antet = doc.add_paragraph()
            p_antet.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p_antet.add_run("AGENȚIA DE PLĂȚI ȘI INTERVENȚIE PENTRU AGRICULTURĂ")
            r1.bold = True; r1.font.size = Pt(13)
            r1.font.color.rgb = RGBColor(0x00, 0x4e, 0x92)

            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.add_run(f"{unitate} | Str. I.C. Pompilian nr. 51, Târgu Jiu, jud. Gorj").font.size = Pt(10)

            doc.add_paragraph()
            titlu_doc = doc.add_heading(
                "RAPORT CONSOLIDAT CONTROL UAV/AI — SESIUNE BATCH",
                level=1
            )
            titlu_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in titlu_doc.runs:
                run.font.color.rgb = RGBColor(0x00, 0x4e, 0x92)

            # Informatii sesiune
            p_ses = doc.add_paragraph()
            p_ses.add_run("Inspector: ").bold = True
            p_ses.add_run(f"{inspector}   ")
            p_ses.add_run("Data: ").bold = True
            p_ses.add_run(data_sesiune.strftime("%d.%m.%Y"))
            p_ses.add_run("   Model AI: ").bold = True
            p_ses.add_run(f"YOLOv8n | best_v1_mAP083_20260403.pt | mAP50=0.829")

            doc.add_paragraph()

            # Rezumat executiv
            doc.add_heading("1. Rezumat executiv", level=2)
            tbl_rez = doc.add_table(rows=6, cols=2)
            tbl_rez.style = "Table Grid"
            date_rezumat = [
                ("Total parcele analizate", str(n_proc)),
                ("Parcele conforme PAC",    f"{n_conforme} ({pct_conf:.0f}%)"),
                ("Parcele neconforme PAC",  f"{n_neconforme} ({100-pct_conf:.0f}%)"),
                ("Total suprafata analizata", f"{total_ha:.2f} ha"),
                ("Suprafata neconforma",    f"{ha_neconf:.2f} ha"),
                ("Medie vegetatie",         f"{medie_veg:.1f}%"),
            ]
            for i, (k, v) in enumerate(date_rezumat):
                tbl_rez.rows[i].cells[0].text = k
                tbl_rez.rows[i].cells[0].paragraphs[0].runs[0].bold = True
                tbl_rez.rows[i].cells[1].text = v

            doc.add_paragraph()

            # Tabel rezultate per parcela
            doc.add_heading("2. Rezultate per parcela", level=2)
            tbl_main = doc.add_table(rows=1 + n_proc, cols=7)
            tbl_main.style = "Table Grid"
            antete = ["Cod LPIS", "Fermier", "Suprafata",
                      "Cultura", "Veg.(%)", "Sol(%)", "Concluzie"]
            for j, h in enumerate(antete):
                cell = tbl_main.rows[0].cells[j]
                cell.text = h
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(9)
                tcPr = cell._tc.get_or_add_tcPr()
                shd  = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "004e92")
                tcPr.append(shd)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            for i, r in enumerate(rezultate):
                row = tbl_main.rows[i + 1]
                valori = [r["cod"], r["fermier"], f"{r['suprafata_ha']:.2f} ha",
                          r["cultura"], f"{r['pct_veg']:.1f}%",
                          f"{r['pct_sol']:.1f}%", r["concluzie"]]
                for j, val in enumerate(valori):
                    cell = row.cells[j]
                    cell.text = val
                    cell.paragraphs[0].runs[0].font.size = Pt(8)
                    if j == 6:  # coloana concluzie
                        run_c = cell.paragraphs[0].runs[0]
                        run_c.bold = True
                        if r["concluzie"] == "CONFORM":
                            run_c.font.color.rgb = RGBColor(0x1a, 0x80, 0x1a)
                        else:
                            run_c.font.color.rgb = RGBColor(0xcc, 0x00, 0x00)

            doc.add_paragraph()

            # Detalii neconforme
            neconforme_lista = [r for r in rezultate if r["concluzie"] == "NECONFORM"]
            if neconforme_lista:
                doc.add_heading("3. Parcele neconforme — detalii", level=2)
                for r in neconforme_lista:
                    doc.add_heading(f"Parcela {r['cod']} — {r['fermier']}", level=3)
                    for prob in r["probleme"]:
                        p_prob = doc.add_paragraph()
                        p_prob.add_run("Problema identificata: ").bold = True
                        run_p = p_prob.add_run(prob)
                        run_p.font.color.rgb = RGBColor(0xcc, 0x00, 0x00)
                    doc.add_paragraph(
                        f"Se recomanda verificare fizica a parcelei {r['cod']} "
                        f"({r['suprafata_ha']:.2f} ha, {r['cultura']}) "
                        f"conform procedurii APIA si Reg. UE 2021/2116."
                    )
                    doc.add_paragraph()

            # Nota metodologica
            doc.add_heading("4. Nota metodologica", level=2)
            nota = doc.add_paragraph(
                "Prezentul raport a fost generat automat de sistemul AGROVISION "
                f"(YOLOv8n, transfer learning, mAP50=0.829) pe baza analizei "
                f"imaginilor UAV pentru {n_proc} parcele LPIS Gorj, "
                f"sesiunea din {data_sesiune.strftime('%d.%m.%Y')}. "
                "Concluziile au caracter orientativ si se completeaza cu "
                "verificarea documentelor justificative ale fermierului. "
                "Baza legala: Reg. UE 2021/2116 (IACS), Reg. UE 2022/1173."
            )
            nota.runs[0].font.size = Pt(9)
            nota.runs[0].font.italic = True

            # Semnatura
            doc.add_paragraph()
            doc.add_heading("5. Semnatura inspector", level=2)
            tbl_s = doc.add_table(rows=3, cols=2)
            tbl_s.style = "Table Grid"
            tbl_s.rows[0].cells[0].text = "Inspector APIA"
            tbl_s.rows[0].cells[1].text = inspector
            tbl_s.rows[1].cells[0].text = "Data"
            tbl_s.rows[1].cells[1].text = date.today().strftime("%d.%m.%Y")
            tbl_s.rows[2].cells[0].text = "Semnatura"
            tbl_s.rows[2].cells[1].text = "______________________________"

            buf_word = BytesIO()
            doc.save(buf_word)
            buf_word.seek(0)
            st.download_button(
                "Descarca Raport Word",
                data=buf_word,
                file_name=f"Raport_Batch_APIA_{data_sesiune.strftime('%Y%m%d')}_{n_proc}parcele.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.markdown('<div class="ok-box">Raport Word generat!</div>', unsafe_allow_html=True)

    # ── 2. EXCEL CENTRALIZATOR ──────────────────────────────────────────────
    with col_ex2:
        if st.button("Genereaza Excel Centralizator", type="primary", use_container_width=True):
            wb = openpyxl.Workbook()

            # Sheet 1 — Rezultate
            ws1 = wb.active
            ws1.title = "Rezultate_Batch"

            # Header
            antet_xl = ["Cod LPIS", "Fermier", "CNP/CUI", "Suprafata (ha)",
                        "Cultura", "Nr. Detectii", "Vegetatie (%)",
                        "Sol gol (%)", "Apa (%)", "Concluzie PAC",
                        "Imagine sursa"]
            for j, h in enumerate(antet_xl, 1):
                cell = ws1.cell(row=1, column=j, value=h)
                cell.font      = Font(bold=True, color="FFFFFF", size=10)
                cell.fill      = PatternFill("solid", fgColor="004e92")
                cell.alignment = Alignment(horizontal="center", wrap_text=True)

            # Date
            for i, r in enumerate(rezultate, 2):
                row_data = [
                    r["cod"], r["fermier"], r["cnp"],
                    r["suprafata_ha"], r["cultura"], r["n_detectii"],
                    r["pct_veg"], r["pct_sol"], r["pct_apa"],
                    r["concluzie"], r["sursa"]
                ]
                for j, val in enumerate(row_data, 1):
                    cell = ws1.cell(row=i, column=j, value=val)
                    cell.alignment = Alignment(horizontal="center")
                    if j == 10:  # Concluzie
                        if val == "CONFORM":
                            cell.font = Font(bold=True, color="1a801a")
                            cell.fill = PatternFill("solid", fgColor="d5f5d5")
                        else:
                            cell.font = Font(bold=True, color="cc0000")
                            cell.fill = PatternFill("solid", fgColor="ffe5e5")
                    if j == 7:  # Vegetatie
                        if float(val) < prag_veg_pac:
                            cell.font = Font(color="cc0000", bold=True)
                        else:
                            cell.font = Font(color="1a801a")

            # Latimi coloane
            latimi = [18, 22, 16, 14, 18, 12, 13, 12, 10, 15, 12]
            for j, lat in enumerate(latimi, 1):
                ws1.column_dimensions[get_column_letter(j)].width = lat

            # Sheet 2 — Rezumat
            ws2 = wb.create_sheet("Rezumat")
            ws2["A1"] = "REZUMAT SESIUNE BATCH APIA"
            ws2["A1"].font = Font(bold=True, size=14, color="004e92")

            rezumat_date = [
                ("Inspector", inspector),
                ("Data sesiunii", data_sesiune.strftime("%d.%m.%Y")),
                ("Model AI", "YOLOv8n | best_v1_mAP083_20260403.pt | mAP50=0.829"),
                ("Total parcele", n_proc),
                ("Conforme PAC", f"{n_conforme} ({pct_conf:.0f}%)"),
                ("Neconforme PAC", f"{n_neconforme} ({100-pct_conf:.0f}%)"),
                ("Total suprafata", f"{total_ha:.2f} ha"),
                ("Suprafata neconforma", f"{ha_neconf:.2f} ha"),
                ("Medie vegetatie", f"{medie_veg:.1f}%"),
                ("Prag PAC aplicat", f"{prag_veg_pac}%"),
            ]
            for row_idx, (k, v) in enumerate(rezumat_date, 3):
                ws2.cell(row=row_idx, column=1, value=k).font = Font(bold=True)
                ws2.cell(row=row_idx, column=2, value=str(v))
            ws2.column_dimensions["A"].width = 28
            ws2.column_dimensions["B"].width = 50

            # Sheet 3 — Neconforme
            ws3 = wb.create_sheet("Neconforme")
            ws3["A1"] = "PARCELE NECONFORME PAC"
            ws3["A1"].font = Font(bold=True, size=12, color="cc0000")
            antet_nc = ["Cod LPIS", "Fermier", "Suprafata", "Cultura",
                        "Vegetatie (%)", "Probleme identificate"]
            for j, h in enumerate(antet_nc, 1):
                cell = ws3.cell(row=2, column=j, value=h)
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="cc0000")
                cell.alignment = Alignment(horizontal="center")

            row_nc = 3
            for r in rezultate:
                if r["concluzie"] == "NECONFORM":
                    ws3.cell(row=row_nc, column=1, value=r["cod"])
                    ws3.cell(row=row_nc, column=2, value=r["fermier"])
                    ws3.cell(row=row_nc, column=3, value=r["suprafata_ha"])
                    ws3.cell(row=row_nc, column=4, value=r["cultura"])
                    ws3.cell(row=row_nc, column=5, value=r["pct_veg"])
                    ws3.cell(row=row_nc, column=6, value="; ".join(r["probleme"]))
                    ws3.cell(row=row_nc, column=6).font = Font(color="cc0000")
                    row_nc += 1

            for j, lat in enumerate([18, 22, 14, 18, 13, 55], 1):
                ws3.column_dimensions[get_column_letter(j)].width = lat

            buf_xl = BytesIO()
            wb.save(buf_xl)
            buf_xl.seek(0)
            st.download_button(
                "Descarca Excel Centralizator",
                data=buf_xl,
                file_name=f"Centralizator_Batch_APIA_{data_sesiune.strftime('%Y%m%d')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            st.markdown('<div class="ok-box">Excel generat (3 sheet-uri: Rezultate, Rezumat, Neconforme)!</div>', unsafe_allow_html=True)

    # ── 3. ZIP CU TOT ──────────────────────────────────────────────────────
    with col_ex3:
        if st.button("Genereaza ZIP Complet (toate rapoartele)", type="primary", use_container_width=True):
            buf_zip = BytesIO()
            with zipfile.ZipFile(buf_zip, "w", zipfile.ZIP_DEFLATED) as zf:

                # Imagini annotate
                for cod, img_bytes in imagini_annotate.items():
                    img_bytes.seek(0)
                    zf.writestr(f"imagini/{cod}.png", img_bytes.read())

                # Rapoarte Word individuale per parcela
                for r in rezultate:
                    doc_ind = Document()
                    for section in doc_ind.sections:
                        section.top_margin  = Cm(2.5); section.bottom_margin = Cm(2.5)
                        section.left_margin = Cm(2.5); section.right_margin  = Cm(2.5)

                    p_a = doc_ind.add_paragraph()
                    p_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_a = p_a.add_run("AGENȚIA DE PLĂȚI ȘI INTERVENȚIE PENTRU AGRICULTURĂ")
                    run_a.bold = True; run_a.font.size = Pt(12)
                    run_a.font.color.rgb = RGBColor(0x00, 0x4e, 0x92)

                    doc_ind.add_paragraph(
                        f"{unitate} | Str. I.C. Pompilian nr. 51, Târgu Jiu"
                    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

                    titlu_ind = doc_ind.add_heading(
                        f"RAPORT CONTROL UAV/AI — PARCELA {r['cod']}", level=1
                    )
                    titlu_ind.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    doc_ind.add_heading("Date fermier si parcela", level=2)
                    tbl_i = doc_ind.add_table(rows=5, cols=2)
                    tbl_i.style = "Table Grid"
                    for row_i, (k, v) in enumerate([
                        ("Fermier",          r["fermier"]),
                        ("CNP/CUI",          r["cnp"]),
                        ("Cod LPIS",         r["cod"]),
                        ("Suprafata",        f"{r['suprafata_ha']:.2f} ha"),
                        ("Cultura declarata",r["cultura"]),
                    ]):
                        tbl_i.rows[row_i].cells[0].text = k
                        tbl_i.rows[row_i].cells[0].paragraphs[0].runs[0].bold = True
                        tbl_i.rows[row_i].cells[1].text = v

                    doc_ind.add_paragraph()
                    doc_ind.add_heading("Rezultate analiza AI", level=2)
                    p_ai = doc_ind.add_paragraph()
                    p_ai.add_run("Nr. detectii: ").bold = True
                    p_ai.add_run(f"{r['n_detectii']}\n")
                    p_ai.add_run("Vegetatie: ").bold = True
                    p_ai.add_run(f"{r['pct_veg']:.1f}%   ")
                    p_ai.add_run("Sol gol: ").bold = True
                    p_ai.add_run(f"{r['pct_sol']:.1f}%   ")
                    p_ai.add_run("Apa: ").bold = True
                    p_ai.add_run(f"{r['pct_apa']:.1f}%")

                    doc_ind.add_paragraph()
                    doc_ind.add_heading("Concluzie PAC", level=2)
                    p_conc = doc_ind.add_paragraph()
                    run_c = p_conc.add_run(f"CONCLUZIE: {r['concluzie']}")
                    run_c.bold = True; run_c.font.size = Pt(12)
                    if r["concluzie"] == "CONFORM":
                        run_c.font.color.rgb = RGBColor(0x1a, 0x80, 0x1a)
                        doc_ind.add_paragraph(
                            f"Parcela {r['cod']} este conforma cu declaratia "
                            f"({r['cultura']}, {r['suprafata_ha']:.2f} ha). "
                            "Vegetatie suficienta detectata. Eligibila PAC 2023-2027."
                        )
                    else:
                        run_c.font.color.rgb = RGBColor(0xcc, 0x00, 0x00)
                        for prob in r["probleme"]:
                            doc_ind.add_paragraph(f"- {prob}")
                        doc_ind.add_paragraph(
                            "Se recomanda control fizic suplimentar conform procedurii APIA."
                        )

                    buf_ind = BytesIO()
                    doc_ind.save(buf_ind)
                    buf_ind.seek(0)
                    zf.writestr(
                        f"rapoarte_individuale/Raport_{r['cod']}_{data_sesiune.strftime('%Y%m%d')}.docx",
                        buf_ind.read()
                    )

                # Metadata JSON
                import json
                metadata = {
                    "sesiune": data_sesiune.strftime("%d.%m.%Y"),
                    "inspector": inspector,
                    "model": "YOLOv8n | best_v1_mAP083_20260403.pt | mAP50=0.829",
                    "n_parcele": n_proc,
                    "conforme": n_conforme,
                    "neconforme": n_neconforme,
                    "prag_vegetatie_pac": prag_veg_pac,
                    "rezultate": [{
                        "cod": r["cod"], "fermier": r["fermier"],
                        "suprafata_ha": r["suprafata_ha"],
                        "pct_veg": r["pct_veg"], "concluzie": r["concluzie"]
                    } for r in rezultate]
                }
                zf.writestr("metadata_sesiune.json", json.dumps(metadata, indent=2, ensure_ascii=False))

            buf_zip.seek(0)
            st.download_button(
                "Descarca ZIP Complet",
                data=buf_zip,
                file_name=f"Batch_APIA_{data_sesiune.strftime('%Y%m%d')}_{n_proc}parcele.zip",
                mime="application/zip"
            )
            st.markdown(f'<div class="ok-box">ZIP generat: {n_proc} rapoarte individuale + imagini annotate + metadata.json</div>', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# REZUMAT LECTIE
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("---")
with st.expander("Rezumat Ziua 15 — Ce am invatat"):
    st.markdown("""
**Batch procesare YOLOv8 — concepte noi fata de Ziua 14:**

| Concept | Ziua 14 (single) | Ziua 15 (batch) |
|---|---|---|
| Upload | 1 imagine | N imagini simultan |
| Procesare | Manuala per parcela | Automata in bucla for |
| Progres | - | st.progress((i+1)/total) |
| Rezultate | 1 raport Word | Excel 3 sheet-uri + ZIP |
| Timp | ~10 min/parcela | < 5 sec pentru 10+ parcele |

**Functii noi folosite:**
```
# Upload multiple
uploaded_files = st.file_uploader(..., accept_multiple_files=True)

# Progres vizual
progress_bar = st.progress(0, text="...")
progress_bar.progress((i+1)/n_proc, text=f"Parcela {i+1}/{n_proc}")

# ZIP in memorie (zero fisiere pe disk)
buf_zip = BytesIO()
with zipfile.ZipFile(buf_zip, "w", zipfile.ZIP_DEFLATED) as zf:
    zf.writestr("fisier.docx", continut_bytes)

# JSON metadata
import json
zf.writestr("metadata.json", json.dumps(dict, indent=2, ensure_ascii=False))
```

**Valoarea operationala APIA:**
- Inspector incarca 10 imagini drone → 10 rapoarte in < 10 secunde
- Export ZIP cu toate documentele pentru audit UE
- Trasabilitate completa: metadata.json cu toate rezultatele sesiunii
- Baza legala: Reg. UE 2021/2116 (IACS) | Reg. UE 2022/1173

**Urmatoarea zi — Ziua 16:** Comparatie temporala batch — T1 vs T2 pe aceleasi parcele,
detectie automata schimbari de vegetatie intre doua sesiuni de zbor.
    """)
