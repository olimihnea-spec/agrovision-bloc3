"""
BLOC 3 — Deep Learning YOLOv8, Ziua 10
Validare dataset final — verificare completitudine, distributie clase, statistici BBox, vizualizare
Autor: Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj

CONCEPT CHEIE:
  Un dataset YOLO valid are:
    1. Perechi complete imagine + adnotare (.jpg/.png + .txt)
    2. Adnotari cu format corect: class_id x_c y_c w h (5 valori, toate in [0,1])
    3. Distributie echilibrata de clase (nu 90% clasa 0 si 10% clasa 1)
    4. BBox-uri cu dimensiuni rezonabile (nu 1px, nu toata imaginea)
  Inainte de antrenament: valideaza dataset! Greselile aici = model slab antrenat.
"""

import streamlit as st
import numpy as np
from PIL import Image, ImageDraw
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from io import BytesIO
import zipfile
import os
from collections import defaultdict, Counter
from datetime import date

st.set_page_config(page_title="Validare Dataset — Ziua 10", layout="wide")

st.markdown("""
<style>
.bloc3-header {
    background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
    padding: 1.5rem 2rem;
    border-radius: 12px;
    margin-bottom: 1.5rem;
    border-left: 5px solid #e94560;
}
.bloc3-header h1 { color: #e94560; margin: 0; font-size: 1.6rem; }
.bloc3-header p { color: #a8b2d8; margin: 0.3rem 0 0 0; font-size: 0.9rem; }
.concept-box {
    background: #0f3460;
    border: 1px solid #e94560;
    border-radius: 8px;
    padding: 1rem 1.2rem;
    margin: 1rem 0;
    color: #a8b2d8;
    font-size: 0.88rem;
}
.concept-box b { color: #e94560; }
.ok-box {
    background: #0d2b0d;
    border: 1px solid #27ae60;
    border-radius: 8px;
    padding: 0.8rem 1rem;
    color: #7dcea0;
    font-size: 0.9rem;
    margin: 0.4rem 0;
}
.warn-box {
    background: #2d1b00;
    border: 1px solid #e67e22;
    border-radius: 8px;
    padding: 0.8rem 1rem;
    color: #f39c12;
    font-size: 0.9rem;
    margin: 0.4rem 0;
}
.err-box {
    background: #2d0000;
    border: 1px solid #e74c3c;
    border-radius: 8px;
    padding: 0.8rem 1rem;
    color: #f1948a;
    font-size: 0.9rem;
    margin: 0.4rem 0;
}
.stat-card {
    background: #16213e;
    border: 1px solid #0f3460;
    border-radius: 8px;
    padding: 1rem;
    text-align: center;
    color: #a8b2d8;
}
.stat-card .val { font-size: 2rem; font-weight: bold; color: #e94560; }
.stat-card .lbl { font-size: 0.8rem; margin-top: 0.2rem; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="bloc3-header">
  <h1>Ziua 10 — Validare Dataset Final</h1>
  <p>Bloc 3 YOLOv8 | Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj</p>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div class="concept-box">
<b>CONCEPTE CHEIE AZI:</b><br>
<b>Validare dataset</b> = verificare inainte de antrenament ca totul e corect<br>
<b>Perechi complete</b> = fiecare imagine .jpg/.png are un fisier .txt cu acelasi nume<br>
<b>Format YOLO</b> = class_id x_c y_c w h (toate valorile in [0,1])<br>
<b>Class imbalance</b> = daca o clasa are &lt;10% din adnotari, modelul o va invata slab<br>
<b>BBox degenerat</b> = w sau h &lt; 0.01 (prea mic) sau &gt; 0.95 (prea mare) = eroare de adnotare
</div>
""", unsafe_allow_html=True)

# ── Sectiunea 1: Upload dataset ZIP ──────────────────────────────────────────

st.header("1. Incarca dataset YOLO (ZIP)")

st.info("""
**Format asteptat in ZIP:**
```
images/  (sau orice folder cu imagini .jpg/.png)
labels/  (fisiere .txt cu adnotari YOLO)
data.yaml  (optional)
```
Poti folosi ZIP-ul generat la Zilele 7 (Split) sau 9 (Augmentare).
""")

uploaded_zip = st.file_uploader(
    "Incarca ZIP cu dataset YOLO",
    type=["zip"],
    help="ZIP cu foldere images/ si labels/"
)

# ── Generare dataset demo ─────────────────────────────────────────────────────

st.markdown("---")
st.subheader("Sau foloseste dataset demonstrativ")

col_demo1, col_demo2, col_demo3 = st.columns(3)
with col_demo1:
    n_imagini = st.slider("Numar imagini", 20, 200, 80, 10)
with col_demo2:
    n_clase = st.slider("Numar clase", 2, 6, 3)
with col_demo3:
    pct_lipsa = st.slider("% adnotari lipsa (erori simulate)", 0, 30, 5)

class_names_input = st.text_input(
    "Nume clase (separate prin virgula)",
    value="vegetatie,sol_gol,apa" if n_clase == 3 else ",".join([f"clasa_{i}" for i in range(n_clase)])
)
class_names = [c.strip() for c in class_names_input.split(",")][:n_clase]
while len(class_names) < n_clase:
    class_names.append(f"clasa_{len(class_names)}")

btn_demo = st.button("Genereaza dataset demonstrativ", type="primary")

# ── Stare sesiune ─────────────────────────────────────────────────────────────

if "dataset_incarcat" not in st.session_state:
    st.session_state.dataset_incarcat = False
    st.session_state.ds_images = {}      # nume -> bytes
    st.session_state.ds_labels = {}      # nume -> list of lines
    st.session_state.ds_class_names = []
    st.session_state.ds_perechi_ok = []
    st.session_state.ds_perechi_fara_label = []
    st.session_state.ds_label_fara_img = []

def parse_label_lines(lines):
    """Parseaza liniile unui fisier .txt YOLO. Returneaza lista de (cls, xc, yc, w, h)."""
    result = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        parts = line.split()
        if len(parts) != 5:
            continue
        try:
            cls_id = int(parts[0])
            xc, yc, w, h = float(parts[1]), float(parts[2]), float(parts[3]), float(parts[4])
            result.append((cls_id, xc, yc, w, h))
        except ValueError:
            continue
    return result

# ── Procesare ZIP incarcat ────────────────────────────────────────────────────

if uploaded_zip is not None:
    images_dict = {}
    labels_dict = {}
    class_names_yaml = []

    zf = zipfile.ZipFile(uploaded_zip)
    for name in zf.namelist():
        lower = name.lower()
        basename = os.path.splitext(os.path.basename(name))[0]
        if lower.endswith((".jpg", ".jpeg", ".png")) and basename:
            images_dict[basename] = zf.read(name)
        elif lower.endswith(".txt") and os.path.basename(name).lower() not in ("classes.txt", "notes.txt"):
            lines = zf.read(name).decode("utf-8", errors="ignore").splitlines()
            if basename:
                labels_dict[basename] = lines
        elif lower.endswith(".yaml") or lower.endswith(".yml"):
            content = zf.read(name).decode("utf-8", errors="ignore")
            for line in content.splitlines():
                if line.strip().startswith("names"):
                    # extrage lista de clase
                    rest = line.split(":", 1)[1].strip()
                    if rest.startswith("["):
                        class_names_yaml = [x.strip().strip("'\"") for x in rest.strip("[]").split(",")]

    if class_names_yaml:
        class_names = class_names_yaml

    perechi_ok = sorted(set(images_dict.keys()) & set(labels_dict.keys()))
    fara_label = sorted(set(images_dict.keys()) - set(labels_dict.keys()))
    label_fara_img = sorted(set(labels_dict.keys()) - set(images_dict.keys()))

    st.session_state.ds_images = images_dict
    st.session_state.ds_labels = labels_dict
    st.session_state.ds_class_names = class_names
    st.session_state.ds_perechi_ok = perechi_ok
    st.session_state.ds_perechi_fara_label = fara_label
    st.session_state.ds_label_fara_img = label_fara_img
    st.session_state.dataset_incarcat = True
    st.success(f"ZIP incarcat: {len(images_dict)} imagini, {len(labels_dict)} adnotari.")

# ── Generare dataset demo ─────────────────────────────────────────────────────

if btn_demo:
    import random
    random.seed(42)
    np.random.seed(42)

    images_dict = {}
    labels_dict = {}
    n_lipsa = int(n_imagini * pct_lipsa / 100)
    fara_label_set = set(random.sample(range(n_imagini), min(n_lipsa, n_imagini)))

    # distributie clase: prima clasa dominanta
    weights = [0.5] + [0.5 / (n_clase - 1)] * (n_clase - 1) if n_clase > 1 else [1.0]

    for i in range(n_imagini):
        name = f"img_{i:04d}"
        # imagine sintetica 320x320
        arr = np.random.randint(30, 200, (320, 320, 3), dtype=np.uint8)
        img = Image.fromarray(arr)
        buf = BytesIO()
        img.save(buf, format="JPEG")
        images_dict[name] = buf.getvalue()

        if i not in fara_label_set:
            n_bbox = random.randint(1, 5)
            lines = []
            for _ in range(n_bbox):
                cls_id = random.choices(range(n_clase), weights=weights)[0]
                xc = round(random.uniform(0.1, 0.9), 4)
                yc = round(random.uniform(0.1, 0.9), 4)
                w  = round(random.uniform(0.05, 0.4), 4)
                h  = round(random.uniform(0.05, 0.4), 4)
                # cateva BBox-uri intentionat gresite
                if random.random() < 0.03:
                    w = round(random.uniform(0.96, 1.0), 4)  # prea mare
                if random.random() < 0.02:
                    w = round(random.uniform(0.001, 0.008), 4)  # prea mic
                lines.append(f"{cls_id} {xc} {yc} {w} {h}")
            labels_dict[name] = lines

    perechi_ok = sorted(set(images_dict.keys()) & set(labels_dict.keys()))
    fara_label = sorted(set(images_dict.keys()) - set(labels_dict.keys()))
    label_fara_img = []

    st.session_state.ds_images = images_dict
    st.session_state.ds_labels = labels_dict
    st.session_state.ds_class_names = class_names
    st.session_state.ds_perechi_ok = perechi_ok
    st.session_state.ds_perechi_fara_label = fara_label
    st.session_state.ds_label_fara_img = label_fara_img
    st.session_state.dataset_incarcat = True
    st.success(f"Dataset demo generat: {n_imagini} imagini, {n_clase} clase, {n_lipsa} adnotari lipsa simulate.")

# ══════════════════════════════════════════════════════════════════════════════
# ANALIZA DATASET
# ══════════════════════════════════════════════════════════════════════════════

if st.session_state.dataset_incarcat:

    images_dict   = st.session_state.ds_images
    labels_dict   = st.session_state.ds_labels
    cls_names     = st.session_state.ds_class_names
    perechi_ok    = st.session_state.ds_perechi_ok
    fara_label    = st.session_state.ds_perechi_fara_label
    label_fara_img = st.session_state.ds_label_fara_img

    st.markdown("---")
    st.header("2. Statistici generale")

    total_img = len(images_dict)
    total_lbl = len(labels_dict)
    total_perechi = len(perechi_ok)

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown(f'<div class="stat-card"><div class="val">{total_img}</div><div class="lbl">Imagini totale</div></div>', unsafe_allow_html=True)
    with c2:
        st.markdown(f'<div class="stat-card"><div class="val">{total_lbl}</div><div class="lbl">Fisiere adnotari</div></div>', unsafe_allow_html=True)
    with c3:
        st.markdown(f'<div class="stat-card"><div class="val">{total_perechi}</div><div class="lbl">Perechi complete</div></div>', unsafe_allow_html=True)
    with c4:
        pct_complete = round(total_perechi / max(total_img, 1) * 100, 1)
        st.markdown(f'<div class="stat-card"><div class="val">{pct_complete}%</div><div class="lbl">Completitudine</div></div>', unsafe_allow_html=True)

    # ── Sectiunea 3: Verificare perechi ──────────────────────────────────────

    st.markdown("---")
    st.header("3. Verificare perechi imagine — adnotare")

    if not fara_label and not label_fara_img:
        st.markdown('<div class="ok-box">Toate imaginile au adnotari si toate adnotarile au imagini corespondente.</div>', unsafe_allow_html=True)
    else:
        if fara_label:
            st.markdown(f'<div class="warn-box">Imagini FARA adnotare ({len(fara_label)}): {", ".join(fara_label[:10])}{"..." if len(fara_label) > 10 else ""}</div>', unsafe_allow_html=True)
        if label_fara_img:
            st.markdown(f'<div class="warn-box">Adnotari FARA imagine ({len(label_fara_img)}): {", ".join(label_fara_img[:10])}{"..." if len(label_fara_img) > 10 else ""}</div>', unsafe_allow_html=True)

    # ── Sectiunea 4: Validare format adnotari ────────────────────────────────

    st.markdown("---")
    st.header("4. Validare format adnotari YOLO")

    erori_format = []       # (fisier, linia, motiv)
    bbox_prea_mare = []
    bbox_prea_mic = []
    all_bbox = []           # (cls_id, xc, yc, w, h)
    cls_counter = Counter()

    for name in perechi_ok:
        lines = labels_dict[name]
        parsed = parse_label_lines(lines)
        if not parsed and lines:
            erori_format.append((name, "—", "Nicio adnotare valida parsata"))
        for cls_id, xc, yc, w, h in parsed:
            all_bbox.append((cls_id, xc, yc, w, h))
            cls_counter[cls_id] += 1
            # verificare valori in [0,1]
            for val, vname in [(xc,"xc"),(yc,"yc"),(w,"w"),(h,"h")]:
                if not (0.0 <= val <= 1.0):
                    erori_format.append((name, f"cls={cls_id}", f"{vname}={val:.4f} in afara [0,1]"))
            # BBox prea mare
            if w > 0.95 or h > 0.95:
                bbox_prea_mare.append((name, cls_id, w, h))
            # BBox prea mic
            if w < 0.01 or h < 0.01:
                bbox_prea_mic.append((name, cls_id, w, h))

    col_v1, col_v2, col_v3 = st.columns(3)
    with col_v1:
        if erori_format:
            st.markdown(f'<div class="err-box">Erori format: <b>{len(erori_format)}</b> probleme gasite</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="ok-box">Format corect: toate adnotarile respecta structura YOLO</div>', unsafe_allow_html=True)
    with col_v2:
        if bbox_prea_mare:
            st.markdown(f'<div class="warn-box">BBox prea mari (w/h > 0.95): <b>{len(bbox_prea_mare)}</b></div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="ok-box">Niciun BBox exagerat de mare</div>', unsafe_allow_html=True)
    with col_v3:
        if bbox_prea_mic:
            st.markdown(f'<div class="warn-box">BBox prea mici (w/h < 0.01): <b>{len(bbox_prea_mic)}</b></div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="ok-box">Niciun BBox exagerat de mic</div>', unsafe_allow_html=True)

    if erori_format:
        with st.expander(f"Vezi detalii erori format ({len(erori_format)})"):
            for fisier, linia, motiv in erori_format[:30]:
                st.write(f"- `{fisier}` | {linia} | {motiv}")

    # ── Sectiunea 5: Distributie clase ───────────────────────────────────────

    st.markdown("---")
    st.header("5. Distributie clase")

    if cls_counter:
        total_bbox_count = sum(cls_counter.values())

        fig_cls, ax_cls = plt.subplots(figsize=(8, 3))
        fig_cls.patch.set_facecolor("#16213e")
        ax_cls.set_facecolor("#0f3460")

        cls_ids_sorted = sorted(cls_counter.keys())
        cls_labels = []
        cls_counts = []
        for cid in cls_ids_sorted:
            lbl = cls_names[cid] if cid < len(cls_names) else f"cls_{cid}"
            cls_labels.append(lbl)
            cls_counts.append(cls_counter[cid])

        colors = ["#e94560", "#27ae60", "#3498db", "#f39c12", "#9b59b6", "#1abc9c"]
        bars = ax_cls.bar(cls_labels, cls_counts,
                          color=[colors[i % len(colors)] for i in range(len(cls_labels))],
                          edgecolor="white", linewidth=0.5)

        for bar, cnt in zip(bars, cls_counts):
            pct = cnt / total_bbox_count * 100
            ax_cls.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                        f"{cnt}\n({pct:.1f}%)", ha="center", va="bottom",
                        color="white", fontsize=8)

        ax_cls.set_xlabel("Clasa", color="white")
        ax_cls.set_ylabel("Nr. adnotari", color="white")
        ax_cls.set_title("Distributia claselor in dataset", color="white", fontsize=11)
        ax_cls.tick_params(colors="white")
        for spine in ax_cls.spines.values():
            spine.set_edgecolor("#0f3460")
        plt.tight_layout()

        buf_cls = BytesIO()
        fig_cls.savefig(buf_cls, dpi=150, bbox_inches="tight", facecolor="#16213e")
        buf_cls.seek(0)
        st.image(buf_cls, use_container_width=True)
        plt.close(fig_cls)

        # Avertizare class imbalance
        max_cls = max(cls_counts)
        for lbl, cnt in zip(cls_labels, cls_counts):
            pct = cnt / total_bbox_count * 100
            if pct < 10:
                st.markdown(f'<div class="warn-box">Clasa <b>{lbl}</b> are doar {pct:.1f}% din adnotari — posibil class imbalance. Considera adaugarea mai multor exemple pentru aceasta clasa.</div>', unsafe_allow_html=True)

    # ── Sectiunea 6: Statistici BBox ─────────────────────────────────────────

    st.markdown("---")
    st.header("6. Statistici dimensiuni BBox")

    if all_bbox:
        widths  = [b[3] for b in all_bbox]
        heights = [b[4] for b in all_bbox]
        areas   = [b[3]*b[4] for b in all_bbox]

        fig_bbox, axes = plt.subplots(1, 3, figsize=(12, 3))
        fig_bbox.patch.set_facecolor("#16213e")

        for ax, data, title, color in [
            (axes[0], widths,  "Latime BBox (w)", "#e94560"),
            (axes[1], heights, "Inaltime BBox (h)", "#27ae60"),
            (axes[2], areas,   "Aria BBox (w*h)", "#3498db"),
        ]:
            ax.set_facecolor("#0f3460")
            ax.hist(data, bins=20, color=color, edgecolor="white", linewidth=0.4, alpha=0.85)
            ax.axvline(np.mean(data), color="yellow", linewidth=1.5, linestyle="--", label=f"Media: {np.mean(data):.3f}")
            ax.set_title(title, color="white", fontsize=9)
            ax.tick_params(colors="white")
            ax.legend(fontsize=7, labelcolor="white", facecolor="#0f3460", edgecolor="white")
            for spine in ax.spines.values():
                spine.set_edgecolor("#0f3460")

        plt.tight_layout()
        buf_bbox = BytesIO()
        fig_bbox.savefig(buf_bbox, dpi=150, bbox_inches="tight", facecolor="#16213e")
        buf_bbox.seek(0)
        st.image(buf_bbox, use_container_width=True)
        plt.close(fig_bbox)

        col_s1, col_s2, col_s3 = st.columns(3)
        with col_s1:
            st.metric("Latime medie", f"{np.mean(widths):.3f}", help="Medie w normalizata [0,1]")
        with col_s2:
            st.metric("Inaltime medie", f"{np.mean(heights):.3f}", help="Medie h normalizata [0,1]")
        with col_s3:
            st.metric("Total adnotari valide", len(all_bbox))

    # ── Sectiunea 7: Vizualizare adnotari pe imagine ──────────────────────────

    st.markdown("---")
    st.header("7. Vizualizare adnotari pe imagine")

    if perechi_ok:
        img_selectata = st.selectbox(
            "Alege imaginea",
            options=perechi_ok[:50],
            help="Primele 50 de imagini cu perechi complete"
        )

        try:
            img_bytes = images_dict[img_selectata]
            img_pil = Image.open(BytesIO(img_bytes)).convert("RGB")
            W, H = img_pil.size

            draw = ImageDraw.Draw(img_pil)
            colors_draw = ["#e94560", "#27ae60", "#3498db", "#f39c12", "#9b59b6", "#1abc9c"]

            bbox_list = parse_label_lines(labels_dict[img_selectata])
            for cls_id, xc, yc, w, h in bbox_list:
                x1 = int((xc - w/2) * W)
                y1 = int((yc - h/2) * H)
                x2 = int((xc + w/2) * W)
                y2 = int((yc + h/2) * H)
                color = colors_draw[cls_id % len(colors_draw)]
                draw.rectangle([x1, y1, x2, y2], outline=color, width=3)
                lbl = cls_names[cls_id] if cls_id < len(cls_names) else f"cls_{cls_id}"
                draw.rectangle([x1, y1-16, x1+len(lbl)*7+4, y1], fill=color)
                draw.text((x1+2, y1-14), lbl, fill="white")

            col_img1, col_img2 = st.columns([2, 1])
            with col_img1:
                st.image(img_pil, caption=f"{img_selectata} — {len(bbox_list)} adnotari", use_container_width=True)
            with col_img2:
                st.markdown("**Adnotari brute:**")
                for line in labels_dict[img_selectata]:
                    st.code(line)

        except Exception as e:
            st.error(f"Nu pot afisa imaginea: {e}")

    # ── Sectiunea 8: Raport final ─────────────────────────────────────────────

    st.markdown("---")
    st.header("8. Raport final — Dataset GATA / PROBLEME")

    probleme = []
    recomandari = []

    if fara_label:
        probleme.append(f"{len(fara_label)} imagini fara adnotare")
        recomandari.append("Adnoteaza imaginile lipsa sau sterge-le din dataset")
    if label_fara_img:
        probleme.append(f"{len(label_fara_img)} adnotari fara imagine")
        recomandari.append("Sterge fisierele .txt orfane")
    if erori_format:
        probleme.append(f"{len(erori_format)} erori de format in adnotari")
        recomandari.append("Corecteaza valorile in afara [0,1]")
    if bbox_prea_mare:
        probleme.append(f"{len(bbox_prea_mare)} BBox-uri cu w/h > 0.95")
        recomandari.append("Re-adnoteaza obiectele cu bounding box exagerat de mare")
    if bbox_prea_mic:
        probleme.append(f"{len(bbox_prea_mic)} BBox-uri cu w/h < 0.01")
        recomandari.append("Sterge adnotarile pentru obiecte prea mici (< 10px la 640px)")
    if cls_counter:
        for cid, cnt in cls_counter.items():
            pct = cnt / sum(cls_counter.values()) * 100
            if pct < 10:
                lbl = cls_names[cid] if cid < len(cls_names) else f"cls_{cid}"
                probleme.append(f"Clasa '{lbl}' sub-reprezentata ({pct:.1f}%)")
                recomandari.append(f"Adauga mai multe exemple pentru clasa '{lbl}' (minim 10%)")

    if not probleme:
        st.markdown(f"""
        <div class="ok-box" style="font-size:1rem; padding:1.2rem;">
        <b>Dataset VALID — Gata de antrenament!</b><br>
        {total_perechi} perechi complete | {len(all_bbox)} adnotari valide | {len(cls_counter)} clase<br>
        Nicio problema critica detectata.
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown(f"""
        <div class="warn-box" style="font-size:1rem; padding:1.2rem;">
        <b>Dataset cu {len(probleme)} probleme — Recomandat: rezolva inainte de antrenament</b>
        </div>
        """, unsafe_allow_html=True)
        for i, (prob, rec) in enumerate(zip(probleme, recomandari), 1):
            st.markdown(f"**{i}. Problema:** {prob}  \n**Recomandare:** {rec}")

    # ── Export raport Word ────────────────────────────────────────────────────

    st.markdown("---")
    st.subheader("Exporta raport validare")

    if st.button("Genereaza raport Word"):
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # margini
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # titlu
        titlu = doc.add_heading("RAPORT VALIDARE DATASET YOLO", level=1)
        titlu.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titlu.runs:
            run.font.color.rgb = RGBColor(0x23, 0x6a, 0x8e)

        doc.add_paragraph(f"Data: {date.today().strftime('%d.%m.%Y')} | Autor: Prof. Asoc. Dr. Oliviu Mihnea Gamulescu")
        doc.add_paragraph(f"Universitatea Constantin Brancusi Targu Jiu | APIA CJ Gorj")
        doc.add_paragraph()

        doc.add_heading("1. Statistici generale", level=2)
        tbl = doc.add_table(rows=5, cols=2)
        tbl.style = "Table Grid"
        for row, (label, val) in enumerate([
            ("Total imagini", str(total_img)),
            ("Total adnotari", str(total_lbl)),
            ("Perechi complete", str(total_perechi)),
            ("Completitudine (%)", f"{pct_complete}%"),
            ("Total BBox valide", str(len(all_bbox))),
        ]):
            tbl.rows[row].cells[0].text = label
            tbl.rows[row].cells[1].text = val

        doc.add_paragraph()
        doc.add_heading("2. Distributie clase", level=2)
        if cls_counter:
            tbl2 = doc.add_table(rows=1 + len(cls_counter), cols=3)
            tbl2.style = "Table Grid"
            for i, hdr in enumerate(["Clasa", "Nr. adnotari", "Procent"]):
                tbl2.rows[0].cells[i].text = hdr
            for row_i, cid in enumerate(sorted(cls_counter.keys()), 1):
                lbl = cls_names[cid] if cid < len(cls_names) else f"cls_{cid}"
                cnt = cls_counter[cid]
                pct_c = cnt / sum(cls_counter.values()) * 100
                tbl2.rows[row_i].cells[0].text = lbl
                tbl2.rows[row_i].cells[1].text = str(cnt)
                tbl2.rows[row_i].cells[2].text = f"{pct_c:.1f}%"

        doc.add_paragraph()
        doc.add_heading("3. Probleme identificate", level=2)
        if not probleme:
            doc.add_paragraph("Nicio problema critica — dataset valid pentru antrenament.")
        else:
            for i, (prob, rec) in enumerate(zip(probleme, recomandari), 1):
                p = doc.add_paragraph(style="List Number")
                run_prob = p.add_run(f"Problema: {prob}")
                run_prob.bold = True
                p.add_run(f"\n   Recomandare: {rec}")

        doc.add_paragraph()
        doc.add_heading("4. Concluzie", level=2)
        concl = "DATASET VALID — gata de antrenament YOLOv8." if not probleme else f"Dataset cu {len(probleme)} probleme — recomandat a fi corectate inainte de antrenament."
        doc.add_paragraph(concl)

        buf_word = BytesIO()
        doc.save(buf_word)
        buf_word.seek(0)

        st.download_button(
            label="Descarca raport Word",
            data=buf_word,
            file_name=f"Raport_Validare_Dataset_{date.today().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # ── Rezumat lectie ────────────────────────────────────────────────────────

    st.markdown("---")
    with st.expander("Rezumat Ziua 10 — Ce am invatat"):
        st.markdown("""
**Validarea dataset-ului YOLO** este pasul obligatoriu inainte de antrenament.

| Verificare | Ce detecteaza |
|---|---|
| Perechi imagine-adnotare | Imagini fara .txt sau .txt fara imagine |
| Format YOLO | Valori in afara [0,1], linii malformate |
| BBox degenerate | Prea mari (>0.95) sau prea mici (<0.01) |
| Distributie clase | Class imbalance (<10% dintr-o clasa) |
| Vizualizare | Adnotarile sunt pozitionate corect pe imagine |

**Regula de baza:** Un model YOLOv8 antrenat pe date invalide va produce detectii gresite,
chiar daca arhitectura si parametrii sunt perfecti. **Garbage in = Garbage out.**

**Urmatoarea zi — Ziua 11:** Antrenament YOLOv8 pe dataset propriu (transfer learning, yolov8n.pt)
        """)

else:
    st.info("Incarca un ZIP sau genereaza un dataset demonstrativ pentru a incepe validarea.")
