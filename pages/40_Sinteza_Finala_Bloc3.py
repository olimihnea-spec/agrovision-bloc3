"""
AGROVISION -- Sinteza Finala Bloc 3
Ziua 40 | Autor: Prof. Asoc. Dr. Oliviu Mihnea Gamulescu

BLOC 3 FINALIZAT COMPLET -- 40/40 zile (9 aprilie 2026)

Continut:
  - Timeline interactiv 40 zile
  - Radar chart competente dobandite
  - Statistici aplicatie (module, linii cod, concepte)
  - Certificat Word oficial UCB
  - Roadmap viitor: IEEE FINE 2026 -> MDPI -> PCE UEFISCDI -> Horizon Europe
"""

import streamlit as st
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import io
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── CONFIGURARE ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Sinteza Finala Bloc 3 | AGROVISION",
    page_icon="40",
    layout="wide"
)

# ── DATE BLOC 3 ───────────────────────────────────────────────────────────────
ZILE_BLOC3 = [
    # (nr, titlu_scurt, categorie)
    (1,  "Arhitectura YOLOv8",           "Fundamente"),
    (2,  "Prima detectie",               "Fundamente"),
    (3,  "Detectie drone",               "Fundamente"),
    (4,  "Batch analiza imagini",         "Pipeline"),
    (5,  "Format Dataset YOLO",          "Dataset"),
    (6,  "Adnotare LabelImg",            "Dataset"),
    (7,  "Split automat dataset",        "Dataset"),
    (8,  "Tiling imagini mari",          "Dataset"),
    (9,  "Augmentare dataset",           "Dataset"),
    (10, "Validare dataset",             "Dataset"),
    (11, "Antrenament YOLOv8",           "Model"),
    (12, "Inferenta model antrenat",     "Model"),
    (13, "Evaluare model",               "Model"),
    (14, "Pipeline APIA complet",        "Integrare"),
    (15, "Batch procesare drone",        "Pipeline"),
    (16, "Comparatie temporala T1/T2",   "Analiza"),
    (17, "Export GIS (GeoJSON/SHP/GPX)", "Export"),
    (18, "Dashboard AGROVISION",         "Dashboard"),
    (19, "Autentificare + Roluri",       "Securitate"),
    (20, "Rapoarte PDF (fpdf2)",         "Export"),
    (21, "Deployment Streamlit Cloud",   "Deployment"),
    (22, "Baza de date SQLite",          "Date"),
    (23, "Notificari email SMTP",        "Comunicare"),
    (24, "API REST FastAPI",             "API"),
    (25, "Hugging Face Hub",             "AI/ML"),
    (26, "Analiza NDVI Spectral",        "Analiza"),
    (27, "Clustering Spatial K-Means",   "AI/ML"),
    (28, "Random Forest Clasificare",    "AI/ML"),
    (29, "Random Forest Regresie",       "AI/ML"),
    (30, "Sinteza Bloc 3 (Z30)",         "Sinteza"),
    (31, "QGIS WMS Live",               "GIS"),
    (32, "Analiza Spatiala",            "GIS"),
    (33, "Raport Control Teren GIS",    "GIS"),
    (34, "Dashboard PAC Ministerial",   "Dashboard"),
    (35, "Export Excel Multi-Sheet",    "Export"),
    (36, "Raport PDF Ministerial",      "Export"),
    (37, "Multi-Agenti Simulare",       "Agenti"),
    (38, "Multi-Agenti Avansat",        "Agenti"),
    (39, "Generator Articol ISI",       "Academic"),
    (40, "Sinteza Finala",              "Sinteza"),
]

CATEGORII_CULORI = {
    "Fundamente":  "#3498DB",
    "Dataset":     "#27AE60",
    "Model":       "#E74C3C",
    "Pipeline":    "#9B59B6",
    "Integrare":   "#E67E22",
    "Analiza":     "#1ABC9C",
    "Export":      "#F39C12",
    "Dashboard":   "#2980B9",
    "Securitate":  "#C0392B",
    "Deployment":  "#8E44AD",
    "Date":        "#16A085",
    "Comunicare":  "#D35400",
    "API":         "#7F8C8D",
    "AI/ML":       "#27AE60",
    "GIS":         "#2471A3",
    "Agenti":      "#6C3483",
    "Academic":    "#1A5276",
    "Sinteza":     "#717D7E",
}

STATISTICI = {
    "Module create":         40,
    "Linii de cod (est.)":   "~18.000",
    "Concepte Python noi":   85,
    "Librarii folosite":     22,
    "Figuri academice":      120,
    "Documente Word/PDF":    "40+",
    "Zile de lucru":         40,
    "Model mAP50":           "0.829",
    "Articol depus":         "IEEE FINE 2026 Osaka",
    "App live":              "Streamlit Cloud",
}

COMPETENTE = {
    "Python":           95,
    "Streamlit":        92,
    "YOLOv8 / ML":      78,
    "Computer Vision":  75,
    "GIS / QGIS":       70,
    "SQL / SQLite":     72,
    "API REST":         68,
    "PDF / Word":       85,
    "Articol ISI":      65,
    "Multi-Agenti":     70,
}

# ─────────────────────────────────────────────────────────────────────────────
# FIGURI
# ─────────────────────────────────────────────────────────────────────────────

def fig_timeline() -> bytes:
    fig, ax = plt.subplots(figsize=(14, 6))

    for nr, titlu, cat in ZILE_BLOC3:
        culoare = CATEGORII_CULORI.get(cat, "#BDC3C7")
        ax.barh(nr, 1, left=nr - 1, color=culoare, edgecolor="white", linewidth=0.5)
        if nr % 5 == 0 or nr == 1 or nr == 40:
            ax.text(nr - 0.5, nr, f"Z{nr}", ha="center", va="center",
                    fontsize=7, color="white", fontweight="bold")

    ax.set_xlim(0, 40)
    ax.set_ylim(0.5, 40.5)
    ax.set_xlabel("Ziua", fontsize=10)
    ax.set_ylabel("Nr. Ziua", fontsize=10)
    ax.set_title("Timeline AGROVISION Bloc 3 -- 40 zile", fontsize=13, fontweight="bold")
    ax.invert_yaxis()
    ax.grid(axis="x", alpha=0.2)

    # Legenda
    patches = [mpatches.Patch(color=c, label=cat)
               for cat, c in CATEGORII_CULORI.items()
               if any(z[2] == cat for z in ZILE_BLOC3)]
    ax.legend(handles=patches, loc="lower right", fontsize=7,
              ncol=3, framealpha=0.9)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def fig_radar() -> bytes:
    competente = list(COMPETENTE.keys())
    valori     = list(COMPETENTE.values())
    N = len(competente)

    angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()
    valori_c  = valori + [valori[0]]
    angles_c  = angles + [angles[0]]
    labels_c  = competente + [competente[0]]

    fig, ax = plt.subplots(figsize=(7, 7), subplot_kw=dict(polar=True))
    ax.plot(angles_c, valori_c, color="#27AE60", linewidth=2)
    ax.fill(angles_c, valori_c, color="#27AE60", alpha=0.25)
    ax.set_xticks(angles)
    ax.set_xticklabels(competente, size=9)
    ax.set_ylim(0, 100)
    ax.set_yticks([25, 50, 75, 100])
    ax.set_yticklabels(["25", "50", "75", "100"], size=7)
    ax.set_title("Competente dobandite Bloc 3\n(autoevaluare %)",
                 size=13, fontweight="bold", pad=20)
    ax.grid(color="gray", linestyle="--", linewidth=0.5, alpha=0.5)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def fig_categorii() -> bytes:
    from collections import Counter
    cats = Counter(z[2] for z in ZILE_BLOC3)
    labels = list(cats.keys())
    sizes  = list(cats.values())
    culori = [CATEGORII_CULORI.get(l, "#BDC3C7") for l in labels]

    fig, ax = plt.subplots(figsize=(8, 5))
    bars = ax.barh(labels, sizes, color=culori, edgecolor="white")
    ax.set_xlabel("Nr. zile")
    ax.set_title("Distributie tematici Bloc 3", fontsize=12, fontweight="bold")
    for bar, val in zip(bars, sizes):
        ax.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height() / 2,
                str(val), va="center", fontsize=9)
    ax.set_xlim(0, max(sizes) + 2)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ── CERTIFICAT WORD ───────────────────────────────────────────────────────────
def genereaza_certificat() -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)

    # Antet UCB
    h = doc.add_heading("UNIVERSITATEA \"CONSTANTIN BRANCUSI\" TARGU-JIU", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        run.font.size = Pt(14)

    h2 = doc.add_heading("Facultatea de Inginerie | Departamentul EEA", level=2)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Titlu certificat
    p_cert = doc.add_paragraph()
    p_cert.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cert = p_cert.add_run("CERTIFICAT DE COMPETENTA DIGITALA")
    r_cert.bold      = True
    r_cert.font.size = Pt(18)
    r_cert.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

    doc.add_paragraph("")

    # Corp
    p_corp = doc.add_paragraph()
    p_corp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_corp.add_run(
        "Se acorda prezentul certificat\n"
        "domnului / doamnei\n\n"
    )
    r.font.size = Pt(12)

    p_nume = doc.add_paragraph()
    p_nume.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_n = p_nume.add_run("Prof. Asoc. Dr. OLIVIU MIHNEA GAMULESCU")
    r_n.bold      = True
    r_n.font.size = Pt(16)
    r_n.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    doc.add_paragraph("")

    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_d = p_desc.add_run(
        "pentru parcurgerea cu succes a programului de formare profesionala "
        "AGROVISION -- 90 Zile Deep Learning si Streamlit, Bloc 3: YOLOv8 si "
        "Inteligenta Artificiala in Agricultura, "
        "compus din 40 de module de instruire practica, "
        "finalizate in perioada 2 -- 9 aprilie 2026."
    )
    r_d.font.size = Pt(12)

    doc.add_paragraph("")

    # Competente
    doc.add_heading("Competente certificate:", level=2)
    for comp, val in COMPETENTE.items():
        p_c = doc.add_paragraph(style="List Bullet")
        r_c = p_c.add_run(f"{comp}: {val}%")
        r_c.font.size = Pt(11)

    doc.add_paragraph("")

    # Realizari
    doc.add_heading("Realizari notabile:", level=2)
    realizari = [
        f"Model YOLOv8n antrenat real: mAP50 = {STATISTICI['Model mAP50']} (dataset LPIS Gorj)",
        f"Aplicatie web live: agrovision-bloc3.streamlit.app (28 module functionale)",
        f"Articol depus: IEEE FINE 2026 Osaka (paper_28, in peer review)",
        f"Model publicat: HuggingFace Hub (oliviu-gamulescu/agrovision-yolov8)",
        f"Draft articol ISI generat automat: structura IMRaD completa, date reale",
        f"Sistem multi-agent cu paralelism, memorie SQLite si orchestrator dinamic",
        f"Export GIS: GeoJSON, Shapefile Stereo70, GPX -- compatibil QGIS",
        f"Rapoarte PDF oficiale multi-pagina destinat APIA Central / Prefectura Gorj",
    ]
    for r_text in realizari:
        p_r = doc.add_paragraph(style="List Bullet")
        run = p_r.add_run(r_text)
        run.font.size = Pt(11)

    doc.add_paragraph("")

    # Statistici
    doc.add_heading("Statistici program:", level=2)
    for cheie, val in STATISTICI.items():
        p_s = doc.add_paragraph(style="List Bullet")
        run = p_s.add_run(f"{cheie}: {val}")
        run.font.size = Pt(11)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Semnatura
    p_data = doc.add_paragraph(f"Targu-Jiu, {date.today().strftime('%d %B %Y')}")
    p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p_sem = doc.add_paragraph("Prof. Asoc. Dr. Oliviu Mihnea Gamulescu")
    p_sem.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sem.runs[0].bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# INTERFATA STREAMLIT
# ─────────────────────────────────────────────────────────────────────────────

# ── HEADER SPECIAL ────────────────────────────────────────────────────────────
st.markdown(
    """
    <div style='background:linear-gradient(135deg,#1A5276,#27AE60);
                padding:30px;border-radius:12px;text-align:center;color:white;margin-bottom:20px'>
        <h1 style='margin:0;font-size:2.2rem'>AGROVISION</h1>
        <h2 style='margin:8px 0;font-weight:300'>Bloc 3 Finalizat -- 40/40 Zile</h2>
        <p style='margin:0;opacity:0.85'>9 aprilie 2026 | Prof. Asoc. Dr. Oliviu Mihnea Gamulescu</p>
    </div>
    """,
    unsafe_allow_html=True
)

# KPI-uri principale
k1, k2, k3, k4, k5 = st.columns(5)
k1.metric("Module create",   "40 / 40", delta="COMPLET")
k2.metric("Linii cod (est.)", "~18.000", delta="+617 azi")
k3.metric("Model mAP50",     "0.829",   delta="Real antrenat")
k4.metric("App live",        "Online",  delta="Streamlit Cloud")
k5.metric("Articol depus",   "IEEE FINE 2026", delta="In review")

st.divider()

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "Timeline 40 Zile",
    "Competente & Statistici",
    "Realizari",
    "Certificat",
    "Roadmap Viitor",
])

# ── TAB 1: TIMELINE ────────────────────────────────────────────────────────────
with tab1:
    st.subheader("Toate cele 40 de zile Bloc 3")
    tl_bytes = fig_timeline()
    st.image(tl_bytes, use_container_width=True)

    # Tabel complet
    df_zile = pd.DataFrame(ZILE_BLOC3, columns=["Ziua", "Titlu", "Categorie"])
    st.dataframe(
        df_zile,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Ziua":      st.column_config.NumberColumn("Nr.", width="small"),
            "Titlu":     st.column_config.TextColumn("Titlu modul", width="large"),
            "Categorie": st.column_config.TextColumn("Categorie", width="medium"),
        }
    )

# ── TAB 2: COMPETENTE ─────────────────────────────────────────────────────────
with tab2:
    col_r, col_c = st.columns(2)

    with col_r:
        st.subheader("Radar competente")
        radar_bytes = fig_radar()
        st.image(radar_bytes, use_container_width=True)

    with col_c:
        st.subheader("Distributie tematici")
        cat_bytes = fig_categorii()
        st.image(cat_bytes, use_container_width=True)

    st.subheader("Statistici program")
    df_stat = pd.DataFrame(
        [(k, str(v)) for k, v in STATISTICI.items()],
        columns=["Indicator", "Valoare"]
    )
    st.dataframe(df_stat, use_container_width=True, hide_index=True)

    st.subheader("Librarii Python folosite")
    librarii = [
        "streamlit", "ultralytics (YOLOv8)", "opencv-python", "Pillow",
        "numpy", "pandas", "matplotlib", "scikit-learn", "shap",
        "fpdf2", "python-docx", "openpyxl", "folium", "pyshp",
        "pyproj", "sqlite3", "fastapi", "uvicorn", "huggingface-hub",
        "smtplib (stdlib)", "threading (stdlib)", "zipfile (stdlib)",
    ]
    cols = st.columns(4)
    for i, lib in enumerate(librarii):
        cols[i % 4].markdown(f"- `{lib}`")

# ── TAB 3: REALIZARI ──────────────────────────────────────────────────────────
with tab3:
    st.subheader("Ce ai construit in 40 de zile")

    st.markdown("""
### Aplicatia AGROVISION (live pe Streamlit Cloud)
- **40 module** functionale intr-o singura aplicatie web
- Detectie culturi agricole cu YOLOv8 (model real, mAP50=0.829)
- Integrare LPIS -- verificare conformitate PAC automata
- Dashboard ministerial cu KPI-uri, harti Folium, grafice
- Autentificare cu roluri (admin / inspector / viewer)
- Export in 6 formate: Word, PDF, Excel, GeoJSON, Shapefile, GPX
- Baza de date SQLite cu log actiuni si cache LPIS
- API REST FastAPI cu 7 endpoint-uri si documentatie Swagger
- Model publicat pe HuggingFace Hub
- Sistem multi-agent cu paralelism si orchestrator dinamic

### Model YOLOv8 antrenat real
- Dataset: imagini drone reale din Gorj (parcele LPIS)
- Antrenament: 50 epoch-uri, Intel i7-7500U (fara GPU)
- Performanta: mAP50=0.829 | Precision=0.641 | Recall=0.667
- Model salvat: best_v1_mAP083_20260403.pt

### Publicatii
- **IEEE FINE 2026 Osaka** -- paper_28 depus 4 apr 2026 (in peer review)
  - Titlu: "UAV-Assisted IoT Network for Precision Agriculture: Real-Time Crop
    Detection Using YOLOv8 for LPIS Compliance Monitoring in Romania"
- **Draft MDPI** -- generat automat cu date reale (Ziua 39, gata de completat)

### Competente dobandite
Ai inceput Bloc 3 stiind Python si Streamlit de baza.
Acum stii: arhitectura YOLOv8, antrenament transfer learning, inferenta,
augmentare dataset, evaluare ML, GIS, API REST, autentificare, multi-agenti,
rapoarte oficiale PDF/Word, deployment cloud, articol stiintific ISI.
""")

# ── TAB 4: CERTIFICAT ─────────────────────────────────────────────────────────
with tab4:
    st.subheader("Certificat de competenta digitala")

    col_info, col_btn = st.columns([3, 1])
    with col_info:
        st.info(
            "Certificatul Word contine: antet UCB, lista competente cu procente, "
            "realizari notabile, statistici program, semnatura si data."
        )

    with col_btn:
        cert_bytes = genereaza_certificat()
        st.download_button(
            "Descarca Certificat Word",
            data=cert_bytes,
            file_name=f"Certificat_AGROVISION_Bloc3_{date.today()}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True,
        )

    # Previzualizare competente
    st.subheader("Competente certificate")
    df_comp = pd.DataFrame(
        [(comp, val, "█" * (val // 10)) for comp, val in COMPETENTE.items()],
        columns=["Competenta", "Nivel (%)", "Progres"]
    )
    st.dataframe(df_comp, use_container_width=True, hide_index=True)

# ── TAB 5: ROADMAP ────────────────────────────────────────────────────────────
with tab5:
    st.subheader("Ce urmeaza dupa Bloc 3")

    st.markdown("""
### Roadmap 2026 -- Oliviu Mihnea Gamulescu

```
ACUM (apr 2026)
   |
   v
IEEE FINE 2026 Osaka (aug 2026)
   Asteptam decizia peer review (4-6 saptamani)
   Daca acceptat --> prezentare la conferinta
   |
   v
MDPI Remote Sensing / Agriculture / Drones
   Extindere dataset (>500 imagini, mai multe judete)
   Colaborare APIA CJ Gorj, Dolj, Olt
   Submitere: trim III 2026
   |
   v
PCE UEFISCDI 2026
   Conditie eligibilitate: cel putin 1 articol publicat ISI
   Titlu proiect propus: "Sistem integrat UAV-AI pentru
   monitorizarea conformitatii PAC in Romania"
   Parteneriate: UCB + APIA Central + Directii Agricole
   |
   v
Horizon Europe (2027)
   Call: Agriculture, Forestry and Rural Areas
   Consortiu international (UCB + universitate EU)
   Instrument principal: AGROVISION extins la nivel national
```

---

### Urmatorul proiect tehnic: Bloc 4 (optional)

Teme propuse pentru continuare:
- **NLP + APIA**: procesare automata documente administrative cu Claude API
- **Multimodal**: combina imagini drone + date meteo + sol pentru predictii
- **Mobile**: aplicatie Android pentru inspectori teren (offline + sync)
- **Real-time**: procesare live stream drone cu YOLOv8 + RTSP

---

### Mesaj final

In 40 de zile ai parcurs un drum de la "instalez Python" la:
- model de deep learning antrenat real
- aplicatie web live pe internet
- articol depus la conferinta IEEE din Japonia
- sistem multi-agent cu inteligenta artificiala
- generator automat de articole ISI

Aceasta nu e o realizare medie. Este exact ce inseamna
transformarea digitala a unui profesionist cu experienta
in domeniu -- combinatia dintre expertiza ta APIA/UCB
si instrumentele AI pe care le stapanesti acum.

**Bloc 3: COMPLET. Felicitari.**
""")

    st.success(
        "AGROVISION Bloc 3 -- 40/40 zile finalizate. "
        "Model live: agrovision-bloc3.streamlit.app | "
        "Articol: IEEE FINE 2026 Osaka (in review)"
    )
