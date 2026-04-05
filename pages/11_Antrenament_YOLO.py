"""
BLOC 3 — Deep Learning YOLOv8, Ziua 11
Antrenament YOLOv8 — transfer learning, parametri, monitorizare, rezultate
Autor: Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj

CONCEPT CHEIE:
  Transfer learning = pornesti de la un model pre-antrenat (yolov8n.pt)
  in loc sa antrenezi de la zero. Modelul "stie" deja sa detecteze forme,
  margini, texturi — tu il inveti doar clasele tale noi.

  Parametri cheie:
    epochs     = cate treceri prin tot dataset-ul (10-100 pentru inceput)
    imgsz      = dimensiunea imaginii la antrenament (640 standard)
    batch      = cate imagini simultan (8-16 pe CPU, 32-64 pe GPU)
    lr0        = learning rate initial (0.01 default)
    patience   = oprire automata daca nu se mai imbunatateste (early stopping)

  Metrici rezultate:
    mAP50      = mean Average Precision la IoU=0.50 (principal)
    mAP50-95   = mAP la IoU=0.50:0.95 (mai strict)
    Precision  = din ce detectezi, cate sunt corecte
    Recall     = din ce exista real, cate le-ai gasit
"""

import streamlit as st
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from io import BytesIO
import time
import random
from datetime import date

st.set_page_config(page_title="Antrenament YOLO — Ziua 11", layout="wide")

st.markdown("""
<style>
.bloc3-header {
    background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
    padding: 1.5rem 2rem; border-radius: 12px;
    margin-bottom: 1.5rem; border-left: 5px solid #e94560;
}
.bloc3-header h1 { color: #e94560; margin: 0; font-size: 1.6rem; }
.bloc3-header p { color: #a8b2d8; margin: 0.3rem 0 0 0; font-size: 0.9rem; }
.concept-box {
    background: #0f3460; border: 1px solid #e94560;
    border-radius: 8px; padding: 1rem 1.2rem; margin: 1rem 0;
    color: #a8b2d8; font-size: 0.88rem;
}
.concept-box b { color: #e94560; }
.ok-box {
    background: #0d2b0d; border: 1px solid #27ae60;
    border-radius: 8px; padding: 0.8rem 1rem; color: #7dcea0;
    font-size: 0.9rem; margin: 0.4rem 0;
}
.warn-box {
    background: #2d1b00; border: 1px solid #e67e22;
    border-radius: 8px; padding: 0.8rem 1rem; color: #f39c12;
    font-size: 0.9rem; margin: 0.4rem 0;
}
.metric-card {
    background: #16213e; border: 1px solid #0f3460;
    border-radius: 8px; padding: 1rem; text-align: center;
}
.metric-card .val { font-size: 2rem; font-weight: bold; color: #e94560; }
.metric-card .lbl { font-size: 0.8rem; color: #a8b2d8; margin-top: 0.2rem; }
.epoch-row {
    background: #16213e; border-radius: 6px; padding: 0.4rem 0.8rem;
    margin: 0.2rem 0; font-family: monospace; font-size: 0.85rem; color: #a8b2d8;
}
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="bloc3-header">
  <h1>Ziua 11 — Antrenament YOLOv8 (Transfer Learning)</h1>
  <p>Bloc 3 YOLOv8 | Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj</p>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div class="concept-box">
<b>CONCEPTE CHEIE AZI:</b><br>
<b>Transfer learning</b> = pornesti de la yolov8n.pt (antrenat pe COCO cu 80 clase) — nu de la zero<br>
<b>epochs</b> = cate treceri complete prin dataset (mai multe = mai bun, dar risc de overfitting)<br>
<b>batch size</b> = cate imagini procesate simultan (mare = mai rapid, necesita mai multa memorie RAM/VRAM)<br>
<b>mAP50</b> = metrica principala de evaluare — cat de bine gaseste modelul obiectele<br>
<b>Overfitting</b> = modelul "memoreaza" datele de antrenament dar esueaza pe date noi<br>
<b>Early stopping</b> = antrenamentul se opreste automat daca mAP nu se mai imbunatateste
</div>
""", unsafe_allow_html=True)

# ── Sectiunea 1: Arhitectura transfer learning ───────────────────────────────

st.header("1. Ce este Transfer Learning?")

col_tl1, col_tl2 = st.columns(2)
with col_tl1:
    st.markdown("""
    **Antrenament de la zero (scratch):**
    - Necesita zeci de mii de imagini
    - Dureaza ore / zile pe GPU
    - Costisitor, impractical pentru cercetare
    """)
with col_tl2:
    st.markdown("""
    **Transfer learning (recomandat):**
    - Pornesti de la `yolov8n.pt` (pre-antrenat COCO)
    - Cateva sute de imagini sunt suficiente
    - Rezultate bune in 10-50 epoch-uri
    """)

st.markdown("""
```
yolov8n.pt  (COCO: 80 clase, 3 mil. imagini, antrenat de Ultralytics)
     |
     v  [Freeze backbone — straturi de baza raman neschimbate]
  Backbone (detectie margini, forme, texturi)
     |
     v  [Antrenezi doar detection head — ultimele straturi]
Detection Head  <-- CLASELE TALE: vegetatie, sol_gol, apa
     |
     v
model_custom.pt  (modelul tau final)
```
""")

# ── Sectiunea 2: Configurare antrenament ────────────────────────────────────

st.header("2. Configurare parametri antrenament")

st.info("""
**Nota:** Antrenamentul real necesita `ultralytics` instalat si un dataset pregatit.
In aceasta pagina simulam antrenamentul pentru a intelege parametrii si metrici.
Codul real pentru linia de comanda este afisat in Sectiunea 3.
""")

col_p1, col_p2, col_p3 = st.columns(3)
with col_p1:
    model_variant = st.selectbox(
        "Varianta model",
        ["yolov8n.pt (Nano — cel mai rapid)", "yolov8s.pt (Small)", "yolov8m.pt (Medium)"],
        help="n=cel mai mic/rapid, m=mai precis dar mai lent"
    )
    epochs = st.slider("Epochs", 5, 100, 30, 5,
        help="Cate treceri complete prin dataset. Incepe cu 30.")

with col_p2:
    imgsz = st.selectbox("Dimensiune imagine (imgsz)", [416, 640, 1280], index=1,
        help="640 este standardul YOLOv8. Mai mare = mai precis dar mai lent.")
    batch = st.selectbox("Batch size", [4, 8, 16, 32], index=1,
        help="8-16 pentru CPU. Reduce daca primesti eroare de memorie.")

with col_p3:
    lr0 = st.select_slider("Learning rate (lr0)", [0.001, 0.005, 0.01, 0.05], value=0.01,
        help="0.01 este default-ul recomandat pentru transfer learning.")
    patience = st.slider("Early stopping (patience)", 5, 50, 20, 5,
        help="Oprire automata daca mAP50 nu creste timp de N epoch-uri.")

n_clase = st.slider("Numar clase in dataset", 2, 6, 3,
    help="Cate clase are dataset-ul tau")
class_names_in = st.text_input("Nume clase", "vegetatie,sol_gol,apa")
class_names = [c.strip() for c in class_names_in.split(",")][:n_clase]

n_train = st.slider("Imagini antrenament", 50, 500, 200, 50)
n_val   = st.slider("Imagini validare", 20, 100, 50, 10)

# ── Sectiunea 3: Cod real ────────────────────────────────────────────────────

st.header("3. Codul real pentru antrenament")

model_name = model_variant.split(" ")[0]

st.code(f"""
# Instalare (o singura data):
# pip install ultralytics

from ultralytics import YOLO

# Incarca modelul pre-antrenat
model = YOLO("{model_name}")

# Antreneaza pe dataset-ul tau
results = model.train(
    data="dataset/data.yaml",   # calea catre fisierul data.yaml
    epochs={epochs},
    imgsz={imgsz},
    batch={batch},
    lr0={lr0},
    patience={patience},
    project="runs/train",       # folder unde se salveaza rezultatele
    name="agrovision_v1",       # numele experimentului
    device="cpu",               # sau "0" pentru GPU NVIDIA
    verbose=True
)

# Dupa antrenament — evalueaza pe setul de validare
metrics = model.val()
print(f"mAP50: {{metrics.box.map50:.3f}}")
print(f"mAP50-95: {{metrics.box.map:.3f}}")

# Salveaza modelul antrenat
# Se salveaza automat in: runs/train/agrovision_v1/weights/best.pt
""", language="python")

st.markdown("""
**Structura `data.yaml` necesara:**
```yaml
path: dataset/          # folderul radacina
train: images/train     # imagini antrenament
val:   images/val       # imagini validare
nc: 3                   # numarul de clase
names: [vegetatie, sol_gol, apa]
```
""")

# ── Sectiunea 4: Simulare antrenament ───────────────────────────────────────

st.header("4. Simulare antrenament (demo vizual)")

st.markdown("""
<div class="concept-box">
<b>De ce simulare?</b> Antrenamentul real pe CPU dureaza 30-120 minute pentru 30 epoch-uri.
Simulam progresul pentru a intelege cum evolueaza metrici si cum arata log-ul real.
</div>
""", unsafe_allow_html=True)

btn_simul = st.button("Ruleaza simulare antrenament", type="primary")

if btn_simul:
    random.seed(42)
    np.random.seed(42)

    # Generam curbe realiste de invatare
    def curba_invatare(epochs, start, plateau, noise=0.02):
        """Curba sigmoida crescatoare cu zgomot."""
        x = np.linspace(-4, 4, epochs)
        sigmoid = 1 / (1 + np.exp(-x))
        curba = start + (plateau - start) * sigmoid
        curba += np.random.normal(0, noise, epochs)
        return np.clip(curba, 0, 1)

    map50_train  = curba_invatare(epochs, 0.05, 0.82, noise=0.015)
    map50_val    = curba_invatare(epochs, 0.04, 0.76, noise=0.025)
    precision    = curba_invatare(epochs, 0.3,  0.88, noise=0.02)
    recall       = curba_invatare(epochs, 0.2,  0.79, noise=0.02)
    loss_box     = np.exp(-np.linspace(0, 3, epochs)) * 0.8 + np.random.normal(0, 0.01, epochs)
    loss_cls     = np.exp(-np.linspace(0, 3, epochs)) * 1.2 + np.random.normal(0, 0.015, epochs)

    # Afisare epoch cu epoch
    st.markdown("**Log antrenament:**")
    log_container = st.container()
    progress_bar = st.progress(0)

    epoch_logs = []
    best_map = 0
    best_epoch = 0

    with log_container:
        for ep in range(1, min(epochs + 1, epochs + 1)):
            pct = ep / epochs
            progress_bar.progress(pct)

            map_val = map50_val[ep-1]
            prec = precision[ep-1]
            rec = recall[ep-1]
            lbox = loss_box[ep-1]
            lcls = loss_cls[ep-1]

            is_best = map_val > best_map
            if is_best:
                best_map = map_val
                best_epoch = ep

            star = " BEST" if is_best else ""
            log_line = f"Epoch {ep:3d}/{epochs} | box_loss={lbox:.4f} | cls_loss={lcls:.4f} | mAP50={map_val:.3f} | P={prec:.3f} | R={rec:.3f}{star}"
            epoch_logs.append(log_line)

            if ep <= 5 or ep % 5 == 0 or ep == epochs:
                color = "#27ae60" if is_best else "#a8b2d8"
                st.markdown(
                    f'<div class="epoch-row" style="color:{color};">{log_line}</div>',
                    unsafe_allow_html=True
                )
            time.sleep(0.03)  # animatie

    st.markdown(f'<div class="ok-box">Antrenament finalizat! Best model: Epoch {best_epoch} cu mAP50={best_map:.3f}</div>', unsafe_allow_html=True)

    # ── Grafice evolutie ─────────────────────────────────────────────────────

    st.markdown("---")
    st.subheader("Curbe de invatare")

    fig, axes = plt.subplots(1, 3, figsize=(14, 4))
    fig.patch.set_facecolor("#16213e")
    epoch_ax = range(1, epochs + 1)

    # mAP50
    ax = axes[0]
    ax.set_facecolor("#0f3460")
    ax.plot(epoch_ax, map50_train, color="#e94560", linewidth=2, label="Train mAP50")
    ax.plot(epoch_ax, map50_val,   color="#27ae60", linewidth=2, label="Val mAP50", linestyle="--")
    ax.axvline(best_epoch, color="yellow", linewidth=1, linestyle=":", alpha=0.7)
    ax.text(best_epoch, 0.02, f"  Best\n  ep.{best_epoch}", color="yellow", fontsize=7)
    ax.set_title("mAP50 — principal", color="white", fontsize=10)
    ax.set_xlabel("Epoch", color="white"); ax.set_ylabel("mAP50", color="white")
    ax.legend(fontsize=8, labelcolor="white", facecolor="#0f3460", edgecolor="white")
    ax.tick_params(colors="white")
    for sp in ax.spines.values(): sp.set_edgecolor("#0f3460")

    # Precision + Recall
    ax = axes[1]
    ax.set_facecolor("#0f3460")
    ax.plot(epoch_ax, precision, color="#3498db", linewidth=2, label="Precision")
    ax.plot(epoch_ax, recall,    color="#f39c12", linewidth=2, label="Recall")
    ax.set_title("Precision & Recall", color="white", fontsize=10)
    ax.set_xlabel("Epoch", color="white"); ax.set_ylabel("Valoare", color="white")
    ax.legend(fontsize=8, labelcolor="white", facecolor="#0f3460", edgecolor="white")
    ax.tick_params(colors="white")
    for sp in ax.spines.values(): sp.set_edgecolor("#0f3460")

    # Loss
    ax = axes[2]
    ax.set_facecolor("#0f3460")
    ax.plot(epoch_ax, loss_box, color="#e94560", linewidth=2, label="Box loss")
    ax.plot(epoch_ax, loss_cls, color="#9b59b6", linewidth=2, label="Cls loss")
    ax.set_title("Loss (scade = bine)", color="white", fontsize=10)
    ax.set_xlabel("Epoch", color="white"); ax.set_ylabel("Loss", color="white")
    ax.legend(fontsize=8, labelcolor="white", facecolor="#0f3460", edgecolor="white")
    ax.tick_params(colors="white")
    for sp in ax.spines.values(): sp.set_edgecolor("#0f3460")

    plt.tight_layout()
    buf_fig = BytesIO()
    fig.savefig(buf_fig, dpi=150, bbox_inches="tight", facecolor="#16213e")
    buf_fig.seek(0)
    st.image(buf_fig, use_container_width=True)
    plt.close(fig)

    # ── Metrici finale ───────────────────────────────────────────────────────

    st.markdown("---")
    st.subheader("Metrici finale (Best model)")

    final_map50    = round(best_map, 3)
    final_map5095  = round(best_map * 0.65, 3)
    final_prec     = round(precision[best_epoch-1], 3)
    final_recall   = round(recall[best_epoch-1], 3)
    final_f1       = round(2 * final_prec * final_recall / max(final_prec + final_recall, 0.001), 3)

    c1, c2, c3, c4, c5 = st.columns(5)
    for col, val, lbl in [
        (c1, f"{final_map50:.3f}",   "mAP50"),
        (c2, f"{final_map5095:.3f}", "mAP50-95"),
        (c3, f"{final_prec:.3f}",    "Precision"),
        (c4, f"{final_recall:.3f}",  "Recall"),
        (c5, f"{final_f1:.3f}",      "F1-Score"),
    ]:
        with col:
            st.markdown(f'<div class="metric-card"><div class="val">{val}</div><div class="lbl">{lbl}</div></div>', unsafe_allow_html=True)

    # interpretare mAP50
    st.markdown("---")
    st.subheader("Interpretare rezultate")

    if final_map50 >= 0.75:
        st.markdown(f'<div class="ok-box"><b>mAP50 = {final_map50}</b> — Model excelent! Gata pentru publicare/utilizare productie.</div>', unsafe_allow_html=True)
    elif final_map50 >= 0.5:
        st.markdown(f'<div class="warn-box"><b>mAP50 = {final_map50}</b> — Model bun. Poti creste epoch-urile sau adauga mai multe date pentru imbunatatire.</div>', unsafe_allow_html=True)
    else:
        st.markdown(f'<div class="warn-box"><b>mAP50 = {final_map50}</b> — Model slab. Verifica dataset-ul (Ziua 10) si creste numarul de imagini.</div>', unsafe_allow_html=True)

    st.markdown("""
    | mAP50 | Interpretare | Actiune recomandata |
    |---|---|---|
    | > 0.85 | Excelent | Publica rezultatele, trece la inferenta |
    | 0.70 – 0.85 | Bun | Acceptabil pentru articol ISI |
    | 0.50 – 0.70 | Mediu | Adauga date, creste epochs, tuneaza parametrii |
    | < 0.50 | Slab | Verifica dataset, re-adnoteaza, schimba modelul |
    """)

    # ── Export raport antrenament ────────────────────────────────────────────

    st.markdown("---")
    st.subheader("Export raport antrenament")

    if st.button("Genereaza raport Word antrenament"):
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

        titlu = doc.add_heading("RAPORT ANTRENAMENT YOLOv8", level=1)
        titlu.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titlu.runs:
            run.font.color.rgb = RGBColor(0x23, 0x6a, 0x8e)

        doc.add_paragraph(f"Data: {date.today().strftime('%d.%m.%Y')} | Model: {model_name}")
        doc.add_paragraph("Autor: Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj")
        doc.add_paragraph()

        doc.add_heading("Parametri antrenament", level=2)
        tbl = doc.add_table(rows=7, cols=2)
        tbl.style = "Table Grid"
        params = [
            ("Model (pretrained)", model_name),
            ("Epochs", str(epochs)),
            ("Image size (imgsz)", str(imgsz)),
            ("Batch size", str(batch)),
            ("Learning rate (lr0)", str(lr0)),
            ("Early stopping (patience)", str(patience)),
            ("Clase", ", ".join(class_names)),
        ]
        for i, (k, v) in enumerate(params):
            tbl.rows[i].cells[0].text = k
            tbl.rows[i].cells[1].text = v

        doc.add_paragraph()
        doc.add_heading("Metrici finale (Best model)", level=2)
        tbl2 = doc.add_table(rows=6, cols=2)
        tbl2.style = "Table Grid"
        metrici = [
            ("Best Epoch", str(best_epoch)),
            ("mAP50", f"{final_map50:.3f}"),
            ("mAP50-95", f"{final_map5095:.3f}"),
            ("Precision", f"{final_prec:.3f}"),
            ("Recall", f"{final_recall:.3f}"),
            ("F1-Score", f"{final_f1:.3f}"),
        ]
        for i, (k, v) in enumerate(metrici):
            tbl2.rows[i].cells[0].text = k
            tbl2.rows[i].cells[1].text = v

        doc.add_paragraph()
        doc.add_heading("Interpretare rezultate", level=2)
        if final_map50 >= 0.75:
            doc.add_paragraph(f"mAP50 = {final_map50} — Model excelent. Gata pentru publicare si utilizare in productie.")
        elif final_map50 >= 0.5:
            doc.add_paragraph(f"mAP50 = {final_map50} — Model bun. Se recomanda cresterea numarului de epoch-uri sau adaugarea de date suplimentare.")
        else:
            doc.add_paragraph(f"mAP50 = {final_map50} — Model slab. Se recomanda verificarea dataset-ului si re-adnotarea imaginilor.")

        doc.add_paragraph()
        doc.add_heading("Fisiere generate de antrenamentul real", level=2)
        doc.add_paragraph("runs/train/agrovision_v1/weights/best.pt — modelul cu cele mai bune performante")
        doc.add_paragraph("runs/train/agrovision_v1/weights/last.pt — modelul de la ultimul epoch")
        doc.add_paragraph("runs/train/agrovision_v1/results.csv — metrici per epoch")
        doc.add_paragraph("runs/train/agrovision_v1/confusion_matrix.png — matricea de confuzie")

        buf_word = BytesIO()
        doc.save(buf_word)
        buf_word.seek(0)
        st.download_button(
            "Descarca raport Word",
            data=buf_word,
            file_name=f"Raport_Antrenament_YOLOv8_{date.today().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # ── Log complet ──────────────────────────────────────────────────────────

    with st.expander("Vezi log complet antrenament"):
        st.code("\n".join(epoch_logs))

# ── Sectiunea 5: Fisiere generate ────────────────────────────────────────────

st.markdown("---")
st.header("5. Fisierele generate de antrenamentul real")

st.markdown("""
Dupa `model.train()`, YOLOv8 creeaza automat urmatoarea structura:

```
runs/
  train/
    agrovision_v1/
      weights/
        best.pt       <-- ACESTA il folosesti pentru inferenta (detectie)
        last.pt       <-- modelul de la ultimul epoch
      results.csv     <-- mAP50, Precision, Recall per epoch
      results.png     <-- grafice automate (curbe invatare)
      confusion_matrix.png  <-- ce confunda modelul
      val_batch0_pred.jpg   <-- exemple detectii pe validare
      args.yaml       <-- toti parametrii salvati automat
```

**Cel mai important fisier: `best.pt`** — acesta este modelul tau antrenat.
Il vei folosi in Ziua 12 pentru inferenta pe imagini noi.
""")

st.markdown("""
```python
# Ziua 12 — cum incarci si folosesti modelul antrenat:
from ultralytics import YOLO

model = YOLO("runs/train/agrovision_v1/weights/best.pt")
results = model.predict("imagine_noua.jpg", conf=0.5)
```
""")

# ── Rezumat lectie ────────────────────────────────────────────────────────────

st.markdown("---")
with st.expander("Rezumat Ziua 11 — Ce am invatat"):
    st.markdown("""
**Transfer Learning YOLOv8** — principii cheie:

| Concept | Explicatie |
|---|---|
| `yolov8n.pt` | Model pre-antrenat pe COCO (80 clase, 3M imagini) |
| Transfer learning | Pornesti de la ce stie deja, inveti clasele tale |
| `epochs` | Cate treceri prin dataset — incepe cu 30 |
| `batch` | 8-16 pe CPU, 32-64 pe GPU |
| `mAP50` | Metrica principala — vrei > 0.70 pentru articol ISI |
| Early stopping | Oprire automata = eviti overfitting |
| `best.pt` | Fisierul cu modelul tau antrenat — pastreaza-l! |

**Curba de invatare normala:**
- Loss scade — bine
- mAP50 creste — bine
- mAP val < mAP train — normal (daca diferenta e mica)
- mAP val mult mai mic decat train — overfitting

**Urmatoarea zi — Ziua 12:** Inferenta cu modelul antrenat — detectie pe imagini noi
    """)
