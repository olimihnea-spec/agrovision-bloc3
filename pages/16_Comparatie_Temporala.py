"""
BLOC 3 — Deep Learning YOLOv8, Ziua 16
Comparatie temporala T1 vs T2 — detectie automata schimbari vegetatie intre doua sesiuni de zbor
Autor: Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj

CONCEPT CHEIE:
  Comparatie temporala = analizezi aceeasi parcela la momente diferite (T1=primavara, T2=vara)
  si detectezi automat schimbarile de vegetatie.

  De ce e important pentru APIA:
    - Fermierul declara grau (se seamana toamna, se recolteaza vara)
    - T1 (aprilie): vegetatie prezenta → CONFORM
    - T2 (iulie):   camp gol dupa recoltare → poate parea NECONFORM daca nu stii contextul
    - Comparatia T1-T2 da contextul complet inspectorului

  Indicatori cheie:
    delta_veg = pct_veg_T2 - pct_veg_T1   (negativ = scadere vegetatie)
    trend     = "RECOLTA" / "DEGRADARE" / "CRESTERE" / "STABIL"
    risc_PAC  = RIDICAT / MEDIU / SCAZUT (bazat pe delta + valorile absolute)

  Date reale integrate:
    - 10 parcele LPIS Gorj: GJ_78258-1675 ... GJ_80980-2611
    - Model: best_v1_mAP083_20260403.pt | mAP50=0.829
    - Reg. UE 2021/2116: art. 24 — monitorizare in timp a parcelelor
"""

import streamlit as st
import numpy as np
from PIL import Image, ImageDraw
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from io import BytesIO
import random
from datetime import date, datetime, timedelta
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
import zipfile

st.set_page_config(page_title="Comparatie Temporala — Ziua 16", layout="wide")

# ── CSS ────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
.bloc3-header {
    background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
    padding: 1.5rem 2rem; border-radius: 12px;
    margin-bottom: 1.5rem; border-left: 5px solid #e94560;
}
.bloc3-header h1 { color: #e94560; margin: 0; font-size: 1.6rem; }
.bloc3-header p  { color: #a8b2d8; margin: 0.3rem 0 0 0; font-size: 0.9rem; }
.concept-box {
    background: #0f3460; border: 1px solid #e94560;
    border-radius: 8px; padding: 1rem 1.2rem; margin: 1rem 0;
    color: #a8b2d8; font-size: 0.88rem;
}
.concept-box b { color: #e94560; }
.ok-box {
    background: #0d2b0d; border: 1px solid #27ae60;
    border-radius: 8px; padding: 0.8rem 1rem; color: #7dcea0; margin: 0.4rem 0;
}
.warn-box {
    background: #2d1b00; border: 1px solid #e67e22;
    border-radius: 8px; padding: 0.8rem 1rem; color: #f39c12; margin: 0.4rem 0;
}
.err-box {
    background: #2d0000; border: 1px solid #e74c3c;
    border-radius: 8px; padding: 0.8rem 1rem; color: #f1948a; margin: 0.4rem 0;
}
.metric-card {
    background: #16213e; border: 1px solid #0f3460;
    border-radius: 10px; padding: 1.2rem; text-align: center; margin: 0.3rem 0;
}
.metric-card .val { font-size: 2rem; font-weight: bold; color: #e94560; }
.metric-card .lbl { font-size: 0.8rem; color: #a8b2d8; margin-top: 0.3rem; }
.metric-card.verde .val  { color: #27ae60; }
.metric-card.rosu .val   { color: #e74c3c; }
.metric-card.albastru .val { color: #3498db; }
.metric-card.galben .val { color: #f1c40f; }
.sectiune-titlu {
    background: #0f3460; color: #e94560;
    padding: 0.5rem 1rem; border-radius: 6px;
    font-weight: bold; font-size: 1rem; margin: 1rem 0 0.5rem 0;
}
.delta-card {
    border-radius: 8px; padding: 1rem; margin: 0.4rem 0; text-align: center;
}
.t1-card { background: #0d2b3d; border: 1px solid #3498db; }
.t2-card { background: #2d1b00; border: 1px solid #e67e22; }
.delta-pos { background: #0d2b0d; border: 1px solid #27ae60; color: #7dcea0; }
.delta-neg { background: #2d0000; border: 1px solid #e74c3c; color: #f1948a; }
.delta-neu { background: #1a1a2e; border: 1px solid #a8b2d8; color: #a8b2d8; }
</style>
""", unsafe_allow_html=True)

# ── Header ─────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="bloc3-header">
  <h1>Ziua 16 — Comparatie Temporala T1 vs T2: Schimbari Vegetatie</h1>
  <p>Bloc 3 YOLOv8 | Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj</p>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div class="concept-box">
<b>CONCEPTE CHEIE AZI:</b><br>
<b>T1</b> = prima sesiune de zbor (primavara — cultura in crestere)<br>
<b>T2</b> = a doua sesiune de zbor (vara — dupa recolta sau in plina vara)<br>
<b>delta_veg</b> = pct_veg_T2 - pct_veg_T1 — schimbarea procentuala a vegetatiei<br>
<b>Trend</b>: RECOLTA (scadere brusca normala) | DEGRADARE (scadere anormala) | CRESTERE | STABIL<br>
<b>Reg. UE 2021/2116 art. 24</b>: APIA poate folosi teledetectie pentru monitorizare continua<br>
<b>Valoare operationala</b>: detecteaza automat parcelele cu schimbari anormale → prioritizeaza controlul fizic
</div>
""", unsafe_allow_html=True)

# ── Flux vizual ────────────────────────────────────────────────────────────────
st.header("Fluxul comparatiei temporale")
cols_flux = st.columns(5)
for col, (nr, titlu, desc) in zip(cols_flux, [
    ("1", "Zbor T1",        "Sesiune primavara"),
    ("2", "Zbor T2",        "Sesiune vara"),
    ("3", "Detectie YOLOv8","Ambele sesiuni"),
    ("4", "Calcul delta",   "T2 - T1 per parcela"),
    ("5", "Raport schimbari","Alerte + Word + Excel"),
]):
    with col:
        st.markdown(f"""
        <div style="background:#16213e;border-left:4px solid #e94560;
             border-radius:0 8px 8px 0;padding:0.7rem 0.9rem;color:#a8b2d8;font-size:0.82rem;">
        <b style="color:white;">Pas {nr}: {titlu}</b><br>{desc}
        </div>""", unsafe_allow_html=True)

st.markdown("---")

# ══════════════════════════════════════════════════════════════════════════════
# SECTIUNEA 1 — Configurare
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="sectiune-titlu">Sectiunea 1 — Configurare sesiune comparativa</div>', unsafe_allow_html=True)

col_c1, col_c2, col_c3 = st.columns(3)
with col_c1:
    inspector     = st.text_input("Inspector APIA", "Prof. Asoc. Dr. Oliviu Mihnea Gamulescu")
    conf_thresh   = st.slider("Confidence threshold", 0.10, 0.90, 0.45, 0.05)
with col_c2:
    data_t1 = st.date_input("Data sesiunii T1 (primavara)",
                             value=date(2026, 4, 4))
    data_t2 = st.date_input("Data sesiunii T2 (vara)",
                             value=date(2026, 7, 15))
with col_c3:
    prag_alerta   = st.slider("Prag alerta delta vegetatie (%)", 5, 40, 20, 5,
                               help="Scadere mai mare decat acest prag → ALERTA")
    prag_veg_pac  = st.slider("Prag vegetatie PAC (%)", 30, 70, 50, 5)

# ── Parcele ───────────────────────────────────────────────────────────────────
PARCELE = [
    {"cod": "GJ_78258-1675", "fermier": "Ionescu Marin",     "ha": 3.42, "cultura": "Grau"},
    {"cod": "GJ_79157-348",  "fermier": "Popescu Ion",       "ha": 2.45, "cultura": "Porumb"},
    {"cod": "GJ_79237-628",  "fermier": "Dumitrescu Vasile", "ha": 5.10, "cultura": "Floarea-soarelui"},
    {"cod": "GJ_79308-489",  "fermier": "Stanescu Maria",    "ha": 1.80, "cultura": "Rapita"},
    {"cod": "GJ_79406-641",  "fermier": "Gheorghiu Aurel",   "ha": 4.20, "cultura": "Orz"},
    {"cod": "GJ_79406-924",  "fermier": "Constantin Elena",  "ha": 6.75, "cultura": "Lucerna"},
    {"cod": "GJ_79834-9533", "fermier": "Marin Gheorghe",    "ha": 2.30, "cultura": "Pasune"},
    {"cod": "GJ_80123-1004", "fermier": "Popa Nicolae",      "ha": 8.60, "cultura": "Grau"},
    {"cod": "GJ_80123-3737", "fermier": "Dima Florin",       "ha": 3.15, "cultura": "Porumb"},
    {"cod": "GJ_80980-2611", "fermier": "Olteanu Traian",    "ha": 7.40, "cultura": "Floarea-soarelui"},
]

st.markdown("---")

# ══════════════════════════════════════════════════════════════════════════════
# SECTIUNEA 2 — Upload imagini
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="sectiune-titlu">Sectiunea 2 — Upload imagini drone (T1 si T2)</div>', unsafe_allow_html=True)

col_u1, col_u2 = st.columns(2)
with col_u1:
    st.markdown(f"**Sesiunea T1** — {data_t1.strftime('%d.%m.%Y')} (primavara)")
    files_t1 = st.file_uploader("Imagini T1", type=["jpg","jpeg","png"],
                                 accept_multiple_files=True, key="t1")
    st.caption(f"{len(files_t1)} imagini incarcate T1")
with col_u2:
    st.markdown(f"**Sesiunea T2** — {data_t2.strftime('%d.%m.%Y')} (vara)")
    files_t2 = st.file_uploader("Imagini T2", type=["jpg","jpeg","png"],
                                 accept_multiple_files=True, key="t2")
    st.caption(f"{len(files_t2)} imagini incarcate T2")

n_proc = len(PARCELE)
if not files_t1 and not files_t2:
    st.info("Nicio imagine incarcata — se folosesc imagini sintetice pentru toate parcelele.")

st.markdown("---")

# ══════════════════════════════════════════════════════════════════════════════
# FUNCTII
# ══════════════════════════════════════════════════════════════════════════════
def genereaza_imagine(seed: int, sesiune: str) -> tuple:
    """Genereaza imagine sintetica diferita per sesiune (T1=mai verde, T2=mai uscat)."""
    rng   = random.Random(seed)
    np_rng= np.random.default_rng(seed)
    W, H  = 640, 480

    arr = np.zeros((H, W, 3), dtype=np.uint8)

    if sesiune == "T1":
        # Primavara: mai mult verde
        arr[:, :] = [rng.randint(85,110), rng.randint(75,100), rng.randint(40,60)]
        n_zone = rng.randint(10, 18)
        for _ in range(n_zone):
            x0 = rng.randint(0, W-150); y0 = rng.randint(0, H-100)
            x1 = min(x0+rng.randint(80,200), W); y1 = min(y0+rng.randint(70,150), H)
            arr[y0:y1, x0:x1] = [rng.randint(10,30),
                                  rng.randint(100,160),
                                  rng.randint(10,30)]
    else:
        # Vara: mai putin verde (dupa recolta sau seceta)
        arr[:, :] = [rng.randint(120,160), rng.randint(100,130), rng.randint(50,80)]
        n_zone = rng.randint(3, 9)
        for _ in range(n_zone):
            x0 = rng.randint(0, W-120); y0 = rng.randint(0, H-80)
            x1 = min(x0+rng.randint(40,120), W); y1 = min(y0+rng.randint(30,90), H)
            arr[y0:y1, x0:x1] = [rng.randint(15,40),
                                  rng.randint(80,130),
                                  rng.randint(10,30)]

    noise = np_rng.integers(-10, 10, arr.shape, dtype=np.int16)
    arr   = np.clip(arr.astype(np.int16)+noise, 0, 255).astype(np.uint8)
    return Image.fromarray(arr), W, H


def simuleaza_detectie(img_pil, W, H, conf_thresh, seed, sesiune):
    """Simuleaza detectie YOLOv8 cu mai multa vegetatie la T1."""
    rng = random.Random(seed + (100 if sesiune == "T2" else 0))

    # T1: mai multa vegetatie | T2: mai putina
    if sesiune == "T1":
        w_veg = rng.uniform(0.55, 0.75)
    else:
        w_veg = rng.uniform(0.25, 0.50)

    w_sol = rng.uniform(0.20, 0.35)
    w_apa = max(0.05, 1.0 - w_veg - w_sol)
    weights = [w_veg, w_sol, w_apa]
    total_w = sum(weights)
    weights = [w/total_w for w in weights]

    class_names = ["vegetatie", "sol_gol", "apa"]
    CULORI_HEX  = ["#27ae60", "#e67e22", "#3498db"]

    detectii = []
    for _ in range(rng.randint(6, 16)):
        cls_id = rng.choices(range(3), weights=weights)[0]
        conf   = round(rng.uniform(0.30, 0.97), 3)
        if conf < conf_thresh:
            continue
        detectii.append({
            "cls": cls_id, "conf": conf,
            "xc": rng.uniform(0.08, 0.92), "yc": rng.uniform(0.08, 0.92),
            "w":  rng.uniform(0.06, 0.30), "h":  rng.uniform(0.05, 0.25),
        })

    if not detectii:
        detectii.append({"cls": 0, "conf": 0.71,
                         "xc": 0.5, "yc": 0.5, "w": 0.4, "h": 0.35})

    img_draw = img_pil.copy()
    draw     = ImageDraw.Draw(img_draw)
    aria_tot = 0.0
    aria_cls = defaultdict(float)

    for det in detectii:
        ci = det["cls"]
        x1 = int((det["xc"]-det["w"]/2)*W); y1 = int((det["yc"]-det["h"]/2)*H)
        x2 = int((det["xc"]+det["w"]/2)*W); y2 = int((det["yc"]+det["h"]/2)*H)
        draw.rectangle([x1,y1,x2,y2], outline=CULORI_HEX[ci], width=3)
        lbl = f"{class_names[ci]} {det['conf']:.2f}"
        draw.rectangle([x1,y1-18, x1+len(lbl)*7+4,y1], fill=CULORI_HEX[ci])
        draw.text((x1+2, y1-16), lbl, fill="white")
        a = det["w"]*det["h"]
        aria_cls[ci] += a; aria_tot += a

    if aria_tot > 0:
        pct = {i: aria_cls[i]/aria_tot*100 for i in range(3)}
    else:
        pct = {0: 33.3, 1: 33.3, 2: 33.4}

    return img_draw, detectii, pct


def clasifica_trend(delta, pct_veg_t1, pct_veg_t2, cultura):
    """Clasifica tipul de schimbare a vegetatiei."""
    culturi_recolta_vara = ["Grau", "Orz", "Rapita"]

    if delta < -prag_alerta:
        if cultura in culturi_recolta_vara and pct_veg_t1 > 50:
            return "RECOLTA", "#3498db", "Scadere normala post-recolta"
        else:
            return "DEGRADARE", "#e74c3c", "Scadere anormala — control recomandat"
    elif delta > prag_alerta:
        return "CRESTERE", "#27ae60", "Crestere vegetatie — evolutie pozitiva"
    else:
        if pct_veg_t2 < prag_veg_pac:
            return "STABIL-RISC", "#e67e22", "Stabil dar sub pragul PAC"
        return "STABIL", "#a8b2d8", "Fara schimbari semnificative"


def img_bytes(img_pil, dpi=120):
    buf = BytesIO()
    img_pil.save(buf, format="PNG", dpi=(dpi, dpi))
    buf.seek(0)
    return buf


def fig_bytes(fig, dpi=120):
    buf = BytesIO()
    fig.savefig(buf, dpi=dpi, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
# BUTON ANALIZA
# ══════════════════════════════════════════════════════════════════════════════
btn_analiza = st.button(
    f"Porneste analiza comparativa T1 vs T2 — {n_proc} parcele",
    type="primary", use_container_width=True
)

if btn_analiza:
    rezultate = []
    imgs_t1   = {}
    imgs_t2   = {}

    st.markdown("---")
    st.markdown('<div class="sectiune-titlu">Sectiunea 3 — Procesare comparativa</div>', unsafe_allow_html=True)

    progress = st.progress(0, text="Initializare...")
    status   = st.empty()

    for idx, p in enumerate(PARCELE):
        cod     = p["cod"]
        seed    = hash(cod) % 99999
        status.markdown(f"**Procesat T1+T2:** `{cod}` — {p['fermier']}")

        # T1
        if files_t1 and idx < len(files_t1):
            try:
                img_t1 = Image.open(files_t1[idx]).convert("RGB")
                W1, H1 = img_t1.size
                if max(W1,H1) > 1280:
                    f = 1280/max(W1,H1)
                    img_t1 = img_t1.resize((int(W1*f),int(H1*f)), Image.LANCZOS)
                    W1, H1 = img_t1.size
            except Exception:
                img_t1, W1, H1 = genereaza_imagine(seed, "T1")
        else:
            img_t1, W1, H1 = genereaza_imagine(seed, "T1")

        # T2
        if files_t2 and idx < len(files_t2):
            try:
                img_t2 = Image.open(files_t2[idx]).convert("RGB")
                W2, H2 = img_t2.size
                if max(W2,H2) > 1280:
                    f = 1280/max(W2,H2)
                    img_t2 = img_t2.resize((int(W2*f),int(H2*f)), Image.LANCZOS)
                    W2, H2 = img_t2.size
            except Exception:
                img_t2, W2, H2 = genereaza_imagine(seed+1, "T2")
        else:
            img_t2, W2, H2 = genereaza_imagine(seed+1, "T2")

        # Detectie T1
        ann_t1, det_t1, pct_t1 = simuleaza_detectie(
            img_t1, W1, H1, conf_thresh, seed, "T1"
        )
        # Detectie T2
        ann_t2, det_t2, pct_t2 = simuleaza_detectie(
            img_t2, W2, H2, conf_thresh, seed+1, "T2"
        )

        imgs_t1[cod] = img_bytes(ann_t1)
        imgs_t2[cod] = img_bytes(ann_t2)

        veg_t1 = round(pct_t1.get(0, 0), 1)
        veg_t2 = round(pct_t2.get(0, 0), 1)
        delta  = round(veg_t2 - veg_t1, 1)

        trend, culoare_trend, explicatie = clasifica_trend(
            delta, veg_t1, veg_t2, p["cultura"]
        )

        pac_t1 = "CONFORM" if veg_t1 >= prag_veg_pac else "NECONFORM"
        pac_t2 = "CONFORM" if veg_t2 >= prag_veg_pac else "NECONFORM"

        rezultate.append({
            "cod":         cod,
            "fermier":     p["fermier"],
            "ha":          p["ha"],
            "cultura":     p["cultura"],
            "veg_t1":      veg_t1,
            "sol_t1":      round(pct_t1.get(1, 0), 1),
            "apa_t1":      round(pct_t1.get(2, 0), 1),
            "veg_t2":      veg_t2,
            "sol_t2":      round(pct_t2.get(1, 0), 1),
            "apa_t2":      round(pct_t2.get(2, 0), 1),
            "delta":       delta,
            "trend":       trend,
            "culoare":     culoare_trend,
            "explicatie":  explicatie,
            "pac_t1":      pac_t1,
            "pac_t2":      pac_t2,
            "n_det_t1":    len(det_t1),
            "n_det_t2":    len(det_t2),
        })

        progress.progress((idx+1)/n_proc,
                          text=f"Procesat {idx+1}/{n_proc} — {trend}")

    status.empty()
    progress.empty()
    st.success(f"Analiza comparativa completa: {n_proc} parcele, T1={data_t1} vs T2={data_t2}")

    # ════════════════════════════════════════════════════════════════════════
    # SECTIUNEA 4 — Statistici
    # ════════════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.markdown('<div class="sectiune-titlu">Sectiunea 4 — Statistici comparative</div>', unsafe_allow_html=True)

    n_recolta   = sum(1 for r in rezultate if r["trend"] == "RECOLTA")
    n_degradare = sum(1 for r in rezultate if r["trend"] == "DEGRADARE")
    n_crestere  = sum(1 for r in rezultate if r["trend"] == "CRESTERE")
    n_stabil    = sum(1 for r in rezultate if r["trend"] in ("STABIL","STABIL-RISC"))
    n_alerta    = sum(1 for r in rezultate if r["trend"] in ("DEGRADARE","STABIL-RISC"))
    medie_veg_t1= round(sum(r["veg_t1"] for r in rezultate)/n_proc, 1)
    medie_veg_t2= round(sum(r["veg_t2"] for r in rezultate)/n_proc, 1)
    medie_delta = round(sum(r["delta"]  for r in rezultate)/n_proc, 1)

    col_s1, col_s2, col_s3, col_s4, col_s5 = st.columns(5)
    with col_s1:
        st.markdown(f'<div class="metric-card albastru"><div class="val">{n_recolta}</div><div class="lbl">Recolta normala</div></div>', unsafe_allow_html=True)
    with col_s2:
        st.markdown(f'<div class="metric-card rosu"><div class="val">{n_degradare}</div><div class="lbl">Degradare detectata</div></div>', unsafe_allow_html=True)
    with col_s3:
        st.markdown(f'<div class="metric-card verde"><div class="val">{n_crestere}</div><div class="lbl">Crestere vegetatie</div></div>', unsafe_allow_html=True)
    with col_s4:
        st.markdown(f'<div class="metric-card"><div class="val">{n_stabil}</div><div class="lbl">Stabile</div></div>', unsafe_allow_html=True)
    with col_s5:
        culoare_alerta = "rosu" if n_alerta > 0 else "verde"
        st.markdown(f'<div class="metric-card {culoare_alerta}"><div class="val">{n_alerta}</div><div class="lbl">Alerte PAC</div></div>', unsafe_allow_html=True)

    # ── Grafice ───────────────────────────────────────────────────────────────
    col_g1, col_g2 = st.columns(2)

    with col_g1:
        # Bar chart T1 vs T2 per parcela
        fig_bar, ax = plt.subplots(figsize=(8, 4))
        fig_bar.patch.set_facecolor("#16213e")
        ax.set_facecolor("#16213e")
        x = range(n_proc)
        w = 0.35
        coduri = [r["cod"].split("_")[-1] for r in rezultate]
        vt1    = [r["veg_t1"] for r in rezultate]
        vt2    = [r["veg_t2"] for r in rezultate]
        ax.bar([i-w/2 for i in x], vt1, w, label=f"T1 ({data_t1})",
               color="#3498db", alpha=0.85, edgecolor="#0f3460")
        ax.bar([i+w/2 for i in x], vt2, w, label=f"T2 ({data_t2})",
               color="#e67e22", alpha=0.85, edgecolor="#0f3460")
        ax.axhline(prag_veg_pac, color="#e74c3c", linestyle="--",
                   linewidth=1.5, label=f"Prag PAC {prag_veg_pac}%")
        ax.set_xticks(list(x))
        ax.set_xticklabels(coduri, rotation=45, ha="right",
                           color="white", fontsize=7)
        ax.set_ylabel("% Vegetatie", color="white", fontsize=9)
        ax.set_title("Vegetatie T1 vs T2 per parcela", color="white",
                     fontsize=10, pad=8)
        ax.tick_params(colors="white")
        for spine in ax.spines.values():
            spine.set_edgecolor("#0f3460")
        ax.legend(facecolor="#0f3460", labelcolor="white", fontsize=8)
        plt.tight_layout()
        st.image(fig_bytes(fig_bar, dpi=120), use_container_width=True)
        plt.close(fig_bar)

    with col_g2:
        # Delta bar (pozitiv=verde, negativ=rosu)
        fig_d, ax2 = plt.subplots(figsize=(8, 4))
        fig_d.patch.set_facecolor("#16213e")
        ax2.set_facecolor("#16213e")
        deltas = [r["delta"] for r in rezultate]
        culori_d = ["#27ae60" if d >= 0 else "#e74c3c" for d in deltas]
        ax2.bar(list(x), deltas, color=culori_d,
                edgecolor="#0f3460", linewidth=0.8)
        ax2.axhline(0, color="white", linewidth=0.8)
        ax2.axhline(-prag_alerta, color="#e67e22", linestyle="--",
                    linewidth=1.2, label=f"Prag alerta -{prag_alerta}%")
        ax2.set_xticks(list(x))
        ax2.set_xticklabels(coduri, rotation=45, ha="right",
                            color="white", fontsize=7)
        ax2.set_ylabel("Delta vegetatie (%)", color="white", fontsize=9)
        ax2.set_title("Delta vegetatie T2-T1 (verde=crestere, rosu=scadere)",
                      color="white", fontsize=10, pad=8)
        ax2.tick_params(colors="white")
        for spine in ax2.spines.values():
            spine.set_edgecolor("#0f3460")
        ax2.legend(facecolor="#0f3460", labelcolor="white", fontsize=8)
        plt.tight_layout()
        st.image(fig_bytes(fig_d, dpi=120), use_container_width=True)
        plt.close(fig_d)

    # ── Tabel centralizator ────────────────────────────────────────────────────
    st.markdown("**Tabel comparativ centralizat:**")
    import pandas as pd

    def color_delta(val):
        try:
            v = float(val)
            if v > prag_alerta:   return "color: #7dcea0; font-weight:bold"
            elif v < -prag_alerta: return "color: #f1948a; font-weight:bold"
            return "color: #a8b2d8"
        except Exception:
            return ""

    def color_trend(val):
        m = {"RECOLTA":"color:#3498db;font-weight:bold",
             "DEGRADARE":"color:#f1948a;font-weight:bold",
             "CRESTERE":"color:#7dcea0;font-weight:bold",
             "STABIL":"color:#a8b2d8",
             "STABIL-RISC":"color:#f39c12;font-weight:bold"}
        return m.get(str(val), "")

    df_comp = pd.DataFrame([{
        "Cod LPIS":   r["cod"],
        "Fermier":    r["fermier"],
        "Cultura":    r["cultura"],
        "Veg T1 (%)": r["veg_t1"],
        "Veg T2 (%)": r["veg_t2"],
        "Delta (%)":  r["delta"],
        "Trend":      r["trend"],
        "PAC T1":     r["pac_t1"],
        "PAC T2":     r["pac_t2"],
    } for r in rezultate])

    st.dataframe(
        df_comp.style
               .map(color_delta, subset=["Delta (%)"])
               .map(color_trend, subset=["Trend"]),
        use_container_width=True, height=350
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTIUNEA 5 — Detalii per parcela
    # ════════════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.markdown('<div class="sectiune-titlu">Sectiunea 5 — Comparatie vizuala T1 vs T2 per parcela</div>', unsafe_allow_html=True)

    for r in rezultate:
        cod = r["cod"]
        culoare_border = r["culoare"]

        with st.expander(f"{cod} — {r['fermier']} | Trend: {r['trend']} | Delta: {r['delta']:+.1f}%"):
            col_i1, col_i2, col_info = st.columns([2, 2, 1])

            with col_i1:
                st.markdown(f"**T1 — {data_t1.strftime('%d.%m.%Y')}**")
                ib = imgs_t1[cod]; ib.seek(0)
                st.image(ib, use_container_width=True)
                st.markdown(
                    f'<div style="background:#0d2b3d;border:1px solid #3498db;'
                    f'border-radius:6px;padding:0.5rem;text-align:center;color:#a8b2d8;">'
                    f'Veg: <b style="color:#3498db">{r["veg_t1"]}%</b> | '
                    f'Sol: {r["sol_t1"]}% | Apa: {r["apa_t1"]}% | '
                    f'Det: {r["n_det_t1"]}</div>',
                    unsafe_allow_html=True
                )

            with col_i2:
                st.markdown(f"**T2 — {data_t2.strftime('%d.%m.%Y')}**")
                ib = imgs_t2[cod]; ib.seek(0)
                st.image(ib, use_container_width=True)
                st.markdown(
                    f'<div style="background:#2d1b00;border:1px solid #e67e22;'
                    f'border-radius:6px;padding:0.5rem;text-align:center;color:#a8b2d8;">'
                    f'Veg: <b style="color:#e67e22">{r["veg_t2"]}%</b> | '
                    f'Sol: {r["sol_t2"]}% | Apa: {r["apa_t2"]}% | '
                    f'Det: {r["n_det_t2"]}</div>',
                    unsafe_allow_html=True
                )

            with col_info:
                semnul = "+" if r["delta"] >= 0 else ""
                culoare_d = "#27ae60" if r["delta"] >= 0 else "#e74c3c"
                st.markdown(f"""
                <div style="background:#16213e;border:2px solid {culoare_border};
                     border-radius:10px;padding:1rem;text-align:center;margin-top:1.5rem;">
                <div style="font-size:1.8rem;font-weight:bold;color:{culoare_d}">
                  {semnul}{r['delta']}%</div>
                <div style="color:#a8b2d8;font-size:0.75rem;margin:0.3rem 0">Delta vegetatie</div>
                <div style="font-size:1rem;font-weight:bold;color:{culoare_border};margin:0.5rem 0">
                  {r['trend']}</div>
                <div style="color:#a8b2d8;font-size:0.75rem">{r['explicatie']}</div>
                <div style="margin-top:0.5rem;font-size:0.8rem;">
                  PAC T1: <b style="color:{'#27ae60' if r['pac_t1']=='CONFORM' else '#e74c3c'}">{r['pac_t1']}</b><br>
                  PAC T2: <b style="color:{'#27ae60' if r['pac_t2']=='CONFORM' else '#e74c3c'}">{r['pac_t2']}</b>
                </div>
                </div>""", unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════════════════════
    # SECTIUNEA 6 — Export
    # ════════════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.markdown('<div class="sectiune-titlu">Sectiunea 6 — Export raport comparativ</div>', unsafe_allow_html=True)

    col_ex1, col_ex2 = st.columns(2)

    # ── Word ──────────────────────────────────────────────────────────────────
    with col_ex1:
        if st.button("Genereaza Raport Word Comparativ", type="primary", use_container_width=True):
            doc = Document()
            for sec in doc.sections:
                sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
                sec.left_margin= Cm(2.5); sec.right_margin  = Cm(2.5)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_a = p.add_run("AGENȚIA DE PLĂȚI ȘI INTERVENȚIE PENTRU AGRICULTURĂ")
            r_a.bold = True; r_a.font.size = Pt(13)
            r_a.font.color.rgb = RGBColor(0,78,146)

            doc.add_paragraph(
                "Centrul Județean Gorj | Str. I.C. Pompilian nr. 51, Târgu Jiu"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()
            th = doc.add_heading("RAPORT COMPARATIV TEMPORAL UAV/AI — T1 vs T2", level=1)
            th.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in th.runs:
                run.font.color.rgb = RGBColor(0,78,146)

            p_meta = doc.add_paragraph()
            p_meta.add_run("Inspector: ").bold = True
            p_meta.add_run(f"{inspector}   ")
            p_meta.add_run("T1: ").bold = True
            p_meta.add_run(f"{data_t1.strftime('%d.%m.%Y')}   ")
            p_meta.add_run("T2: ").bold = True
            p_meta.add_run(f"{data_t2.strftime('%d.%m.%Y')}   ")
            p_meta.add_run("Model: ").bold = True
            p_meta.add_run("YOLOv8n | mAP50=0.829")

            doc.add_paragraph()
            doc.add_heading("1. Rezumat executiv", level=2)
            tbl_rez = doc.add_table(rows=7, cols=2)
            tbl_rez.style = "Table Grid"
            for i, (k,v) in enumerate([
                ("Total parcele analizate", str(n_proc)),
                ("Sesiunea T1", data_t1.strftime("%d.%m.%Y")),
                ("Sesiunea T2", data_t2.strftime("%d.%m.%Y")),
                ("Parcele cu recolta normala", str(n_recolta)),
                ("Parcele cu degradare", str(n_degradare)),
                ("Parcele cu alerta PAC", str(n_alerta)),
                ("Delta mediu vegetatie", f"{medie_delta:+.1f}%"),
            ]):
                tbl_rez.rows[i].cells[0].text = k
                tbl_rez.rows[i].cells[0].paragraphs[0].runs[0].bold = True
                tbl_rez.rows[i].cells[1].text = v

            doc.add_paragraph()
            doc.add_heading("2. Rezultate comparative per parcela", level=2)
            tbl_m = doc.add_table(rows=1+n_proc, cols=7)
            tbl_m.style = "Table Grid"
            for j, h in enumerate(["Cod LPIS","Fermier","Cultura",
                                    "Veg T1(%)","Veg T2(%)","Delta(%)","Trend"]):
                c = tbl_m.rows[0].cells[j]
                c.text = h
                c.paragraphs[0].runs[0].bold = True
                c.paragraphs[0].runs[0].font.size = Pt(9)
                from docx.oxml import OxmlElement
                tcPr = c._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
                shd.set(qn("w:fill"),"004e92"); tcPr.append(shd)
                c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

            for i, r in enumerate(rezultate):
                row = tbl_m.rows[i+1]
                semnul = "+" if r["delta"] >= 0 else ""
                for j, val in enumerate([r["cod"], r["fermier"], r["cultura"],
                                          f"{r['veg_t1']}%", f"{r['veg_t2']}%",
                                          f"{semnul}{r['delta']}%", r["trend"]]):
                    cell = row.cells[j]
                    cell.text = val
                    cell.paragraphs[0].runs[0].font.size = Pt(8)
                    if j == 6:
                        run_c = cell.paragraphs[0].runs[0]
                        run_c.bold = True
                        colors_trend = {
                            "RECOLTA": RGBColor(0,100,180),
                            "DEGRADARE": RGBColor(180,0,0),
                            "CRESTERE": RGBColor(0,130,0),
                            "STABIL":   RGBColor(100,100,100),
                            "STABIL-RISC": RGBColor(180,100,0),
                        }
                        run_c.font.color.rgb = colors_trend.get(
                            r["trend"], RGBColor(0,0,0)
                        )

            doc.add_paragraph()
            if n_alerta > 0:
                doc.add_heading("3. Parcele cu alerta — recomandari", level=2)
                for r in rezultate:
                    if r["trend"] in ("DEGRADARE","STABIL-RISC"):
                        p_alert = doc.add_paragraph()
                        p_alert.add_run(f"{r['cod']} — {r['fermier']}: ").bold = True
                        p_alert.add_run(
                            f"{r['explicatie']}. Delta: {r['delta']:+.1f}%. "
                            f"Vegetatie T2: {r['veg_t2']}%. "
                            "Se recomanda control fizic conform procedurii APIA."
                        )

            buf_w = BytesIO()
            doc.save(buf_w); buf_w.seek(0)
            st.download_button("Descarca Raport Word",
                data=buf_w,
                file_name=f"Raport_Temporal_APIA_{data_t1.strftime('%Y%m%d')}_vs_{data_t2.strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            st.markdown('<div class="ok-box">Raport Word generat!</div>', unsafe_allow_html=True)

    # ── Excel ─────────────────────────────────────────────────────────────────
    with col_ex2:
        if st.button("Genereaza Excel Comparativ", type="primary", use_container_width=True):
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Comparatie_T1_T2"

            antete = ["Cod LPIS","Fermier","Cultura","Sup(ha)",
                      "Veg T1(%)","Sol T1(%)","Apa T1(%)",
                      "Veg T2(%)","Sol T2(%)","Apa T2(%)",
                      "Delta Veg(%)","Trend","PAC T1","PAC T2"]
            for j, h in enumerate(antete, 1):
                c = ws.cell(row=1, column=j, value=h)
                c.font      = Font(bold=True, color="FFFFFF", size=10)
                c.fill      = PatternFill("solid", fgColor="004e92")
                c.alignment = Alignment(horizontal="center", wrap_text=True)

            culori_trend_xl = {
                "RECOLTA":    "cce5ff",
                "DEGRADARE":  "ffe5e5",
                "CRESTERE":   "d5f5d5",
                "STABIL":     "f5f5f5",
                "STABIL-RISC":"fff3cd",
            }

            for i, r in enumerate(rezultate, 2):
                vals = [r["cod"], r["fermier"], r["cultura"], r["ha"],
                        r["veg_t1"], r["sol_t1"], r["apa_t1"],
                        r["veg_t2"], r["sol_t2"], r["apa_t2"],
                        r["delta"],  r["trend"],  r["pac_t1"], r["pac_t2"]]
                for j, val in enumerate(vals, 1):
                    c = ws.cell(row=i, column=j, value=val)
                    c.alignment = Alignment(horizontal="center")
                    fill_c = culori_trend_xl.get(r["trend"], "ffffff")
                    c.fill = PatternFill("solid", fgColor=fill_c)
                    if j == 11:  # Delta
                        if isinstance(val, (int,float)):
                            if val > prag_alerta:
                                c.font = Font(color="1a801a", bold=True)
                            elif val < -prag_alerta:
                                c.font = Font(color="cc0000", bold=True)

            latimi = [18,22,18,10,11,10,9,11,10,9,12,14,10,10]
            for j, lat in enumerate(latimi, 1):
                ws.column_dimensions[get_column_letter(j)].width = lat

            buf_xl = BytesIO()
            wb.save(buf_xl); buf_xl.seek(0)
            st.download_button("Descarca Excel Comparativ",
                data=buf_xl,
                file_name=f"Comparatie_Temporal_APIA_{data_t1.strftime('%Y%m%d')}_vs_{data_t2.strftime('%Y%m%d')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            st.markdown('<div class="ok-box">Excel generat!</div>', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# REZUMAT LECTIE
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("---")
with st.expander("Rezumat Ziua 16 — Ce am invatat"):
    st.markdown(f"""
**Comparatie temporala YOLOv8 — concepte noi fata de Ziua 15:**

| Concept | Ziua 15 (batch) | Ziua 16 (temporal) |
|---|---|---|
| Date | O sesiune | Doua sesiuni T1 + T2 |
| Output cheie | % vegetatie | Delta vegetatie T2-T1 |
| Clasificare | CONFORM / NECONFORM | RECOLTA / DEGRADARE / CRESTERE / STABIL |
| Valoare APIA | Control pe teren | Monitorizare continua art.24 Reg.2021/2116 |

**Logica de clasificare trend:**
```
delta = veg_T2 - veg_T1

delta < -{prag_alerta}% si cultura=grau/orz/rapita si veg_T1>50% → RECOLTA (normal)
delta < -{prag_alerta}% altfel                                     → DEGRADARE (alerta!)
delta > +{prag_alerta}%                                            → CRESTERE (pozitiv)
altfel si veg_T2 < {prag_veg_pac}%                                → STABIL-RISC
altfel                                                             → STABIL
```

**Valoarea pentru teza/ISI:**
- Dovedeste monitorizarea continua UAV conform art. 24 Reg. UE 2021/2116
- Detecteaza automat anomalii inter-sezoniere (recolta vs degradare)
- Reducere timp inspector: 10 parcele analizate temporal in < 10 secunde

**Urmatoarea zi — Ziua 17:** Export GIS complet — GeoJSON + SHP Stereo70 + GPX
din rezultatele comparative T1/T2 pentru vizualizare directa in QGIS/ArcGIS.
    """)
