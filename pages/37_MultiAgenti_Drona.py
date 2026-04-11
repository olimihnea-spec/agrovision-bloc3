"""
AGROVISION — Multi-Agenti Simulare Drona
Ziua 37 | Autor: Prof. Asoc. Dr. Oliviu Mihnea Gamulescu

Scop:
    Arhitectura multi-agent pentru procesarea imaginilor drone APIA.
    Fiecare agent are un rol specializat; orchestratorul le coordoneaza.

    Agenti:
      AgentDrona       — analizeaza imaginea, calculeaza indici ExG/VARI/GLI
      AgentLPIS        — verifica parcela in baza LPIS (suprafata, cultura)
      AgentConformitate — compara detectia cu declaratia, emite verdict PAC
      AgentRaportare   — genereaza raport Word oficial cu toate rezultatele

    Orchestrator:
      Apeleaza agentii in ordine, paseaza rezultatele de la un agent la altul.
      Afiseaza progresul in timp real cu st.status().

CONCEPT CHEIE — arhitectura multi-agent:
    Fiecare agent = o clasa Python cu metoda proceseaza(input) -> dict.
    Orchestratorul nu stie detaliile fiecarui agent — doar ii apeleaza.
    Iesirea unui agent devine intrarea urmatorului (pipeline).

    Agent 1 → rezultat_1 → Agent 2 → rezultat_2 → Agent 3 → ... → Raport final

CONCEPT CHEIE — st.status():
    Afiseaza un bloc expandabil cu progresul in timp real.
    st.status("Procesez...") as s:
        s.update(label="Pasul 1...", state="running")
        # ... lucru ...
        s.update(label="Gata!", state="complete")
"""

import streamlit as st
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from PIL import Image
import io
import time
import random
from datetime import date, datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─── CONFIGURARE PAGINA ───────────────────────────────────────────────────────
st.set_page_config(
    page_title="Multi-Agenti Drona | AGROVISION",
    page_icon="UAV",
    layout="wide"
)

# ─── BAZA DE DATE LPIS SIMULATA ───────────────────────────────────────────────
LPIS_DB = {
    "GJ_78258-1675": {"fermier": "Popescu Ion",     "cultura_declarata": "grau",             "suprafata_declarata": 4.32, "uat": "Targu Jiu"},
    "GJ_78301-0892": {"fermier": "Ionescu Maria",   "cultura_declarata": "porumb",           "suprafata_declarata": 6.78, "uat": "Rovinari"},
    "GJ_78445-2341": {"fermier": "Dumitrescu Gh.",  "cultura_declarata": "rapita",           "suprafata_declarata": 3.15, "uat": "Motru"},
    "GJ_78512-0077": {"fermier": "Stanescu Petre",  "cultura_declarata": "grau",             "suprafata_declarata": 8.90, "uat": "Bumbesti-Jiu"},
    "GJ_78634-1129": {"fermier": "Olteanu Vasile",  "cultura_declarata": "porumb",           "suprafata_declarata": 2.44, "uat": "Novaci"},
    "GJ_78720-3388": {"fermier": "Marinescu Ana",   "cultura_declarata": "floarea_soarelui", "suprafata_declarata": 5.67, "uat": "Targu Jiu"},
    "GJ_78834-0055": {"fermier": "Constantin Gh.",  "cultura_declarata": "lucerna",          "suprafata_declarata": 7.20, "uat": "Aninoasa"},
    "GJ_78901-2277": {"fermier": "Draghici Marin",  "cultura_declarata": "pasune",           "suprafata_declarata": 12.5, "uat": "Targu Jiu"},
    "GJ_79023-0814": {"fermier": "Barbu Nicolae",   "cultura_declarata": "orz",              "suprafata_declarata": 3.88, "uat": "Carbunesti"},
    "GJ_79157-1563": {"fermier": "Munteanu Elena",  "cultura_declarata": "grau",             "suprafata_declarata": 5.11, "uat": "Tismana"},
}

# ─── AGENT 1: DRONA ───────────────────────────────────────────────────────────

class AgentDrona:
    """
    Analizeaza imaginea drone si calculeaza indicii de vegetatie.
    Input:  imagine PIL + cod_parcela
    Output: dict cu indici ExG/VARI/GLI, procent vegetatie, heatmap PNG
    """
    NUME = "Agent Drona"
    ICONA = "UAV"

    def proceseaza(self, imagine: Image.Image, cod_parcela: str) -> dict:
        arr = np.array(imagine).astype(float)
        R = arr[:, :, 0] / 255.0
        G = arr[:, :, 1] / 255.0
        B = arr[:, :, 2] / 255.0

        eps = 1e-6
        ExG  = 2*G - R - B
        VARI = (G - R) / (G + R - B + eps)
        GLI  = (2*G - R - B) / (2*G + R + B + eps)

        # Masca vegetatie (ExG > 0.1)
        masca_veg = ExG > 0.1
        pct_vegetatie = masca_veg.mean() * 100

        # Heatmap ExG
        fig, axes = plt.subplots(1, 3, figsize=(12, 4))
        for ax, index, titlu, cmap in zip(
            axes,
            [ExG, VARI, GLI],
            ["ExG", "VARI", "GLI"],
            ["RdYlGn", "PiYG", "Greens"]
        ):
            im = ax.imshow(index, cmap=cmap, vmin=-0.5, vmax=0.8)
            ax.set_title(titlu, fontsize=11, fontweight="bold")
            ax.axis("off")
            plt.colorbar(im, ax=ax, fraction=0.046)
        fig.suptitle(f"Indici vegetatie — {cod_parcela}", fontsize=12, fontweight="bold")
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=120, bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)
        heatmap_png = buf.read()

        return {
            "cod_parcela":     cod_parcela,
            "exg_mediu":       float(np.mean(ExG)),
            "vari_mediu":      float(np.mean(VARI)),
            "gli_mediu":       float(np.mean(GLI)),
            "pct_vegetatie":   round(pct_vegetatie, 2),
            "rezolutie":       f"{imagine.width}x{imagine.height}px",
            "heatmap_png":     heatmap_png,
        }


# ─── AGENT 2: LPIS ────────────────────────────────────────────────────────────

class AgentLPIS:
    """
    Verifica parcela in baza de date LPIS.
    Input:  cod_parcela (din rezultatul AgentDrona)
    Output: dict cu datele declarate (fermier, cultura, suprafata)
    """
    NUME = "Agent LPIS"
    ICONA = "DB"

    def proceseaza(self, cod_parcela: str) -> dict:
        if cod_parcela in LPIS_DB:
            date_lpis = LPIS_DB[cod_parcela].copy()
            date_lpis["gasit_in_lpis"] = True
            date_lpis["cod_parcela"]   = cod_parcela
        else:
            date_lpis = {
                "gasit_in_lpis":       False,
                "cod_parcela":         cod_parcela,
                "fermier":             "NECUNOSCUT",
                "cultura_declarata":   "nedeclarat",
                "suprafata_declarata": 0.0,
                "uat":                 "N/A",
            }
        return date_lpis


# ─── AGENT 3: CONFORMITATE ────────────────────────────────────────────────────

class AgentConformitate:
    """
    Compara rezultatele drone cu datele LPIS declarate.
    Input:  rezultat AgentDrona + rezultat AgentLPIS
    Output: dict cu verdict PAC, penalizare, explicatie
    """
    NUME = "Agent Conformitate"
    ICONA = "CHK"

    # Praguri vegetatie per cultura (ExG mediu asteptat)
    PRAGURI_CULTURA = {
        "grau":             (0.05, 0.35),
        "porumb":           (0.08, 0.45),
        "rapita":           (0.03, 0.30),
        "floarea_soarelui": (0.05, 0.38),
        "lucerna":          (0.10, 0.50),
        "pasune":           (0.08, 0.50),
        "orz":              (0.05, 0.32),
        "triticale":        (0.05, 0.33),
        "nedeclarat":       (0.0,  1.0),
    }

    def proceseaza(self, rez_drona: dict, rez_lpis: dict) -> dict:
        if not rez_lpis["gasit_in_lpis"]:
            return {
                "verdict":       "NECONFORM",
                "cod":           "LPIS_MISS",
                "penalizare_pct": 100,
                "explicatie":    "Parcela nu exista in LPIS. Cerere invalidata.",
                "recomandat":    "Verificare identitate parcela pe teren.",
            }

        exg      = rez_drona["exg_mediu"]
        pct_veg  = rez_drona["pct_vegetatie"]
        cultura  = rez_lpis["cultura_declarata"]
        prag_min, prag_max = self.PRAGURI_CULTURA.get(cultura, (0.0, 1.0))

        probleme = []
        penalizare = 0

        # Verificare vegetatie insuficienta
        if pct_veg < 20:
            probleme.append(f"Vegetatie detectata {pct_veg:.1f}% < prag minim 20% PAC")
            penalizare += 50

        # Verificare indice ExG in afara intervalului asteptat
        if exg < prag_min:
            probleme.append(f"ExG={exg:.3f} sub minimul pentru {cultura} ({prag_min})")
            penalizare += 25
        elif exg > prag_max:
            probleme.append(f"ExG={exg:.3f} peste maximul pentru {cultura} ({prag_max})")
            penalizare += 10

        # Verdict
        if penalizare == 0:
            verdict     = "CONFORM"
            cod         = "OK"
            explicatie  = f"Parcela conforma PAC. Cultura {cultura} detectata corect."
            recomandat  = "Nicio actiune necesara."
        elif penalizare <= 25:
            verdict     = "ATENTIE"
            cod         = "WARN"
            explicatie  = " | ".join(probleme)
            recomandat  = "Monitorizare suplimentara recomandata. Control teren optional."
        else:
            verdict     = "NECONFORM"
            cod         = "NON_CONF"
            explicatie  = " | ".join(probleme)
            recomandat  = "Control obligatoriu pe teren. Notificare fermier conf. Art. 68 Reg. UE 2021/2116."

        return {
            "verdict":        verdict,
            "cod":            cod,
            "penalizare_pct": penalizare,
            "explicatie":     explicatie,
            "recomandat":     recomandat,
            "probleme":       probleme,
        }


# ─── AGENT 4: RAPORTARE ───────────────────────────────────────────────────────

class AgentRaportare:
    """
    Aduna rezultatele tuturor agentilor si genereaza raportul Word oficial.
    Input:  rezultatele celor 3 agenti anteriori
    Output: bytes (fisier .docx)
    """
    NUME = "Agent Raportare"
    ICONA = "DOC"

    def proceseaza(self, rez_drona: dict, rez_lpis: dict, rez_conf: dict) -> bytes:
        doc = Document()

        # Stil pagina
        for sectiune in doc.sections:
            sectiune.top_margin    = Cm(2.5)
            sectiune.bottom_margin = Cm(2.5)
            sectiune.left_margin   = Cm(2.5)
            sectiune.right_margin  = Cm(2.5)

        def stil_titlu(paragraph, text, size=14, bold=True, color=None):
            paragraph.clear()
            run = paragraph.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def adauga_rand_tabel(tabel, valori, bold=False, bg=None):
            rand = tabel.add_row()
            for i, val in enumerate(valori):
                celula = rand.cells[i]
                celula.text = str(val)
                run = celula.paragraphs[0].runs[0] if celula.paragraphs[0].runs else celula.paragraphs[0].add_run(str(val))
                run.bold = bold
                run.font.size = Pt(9)

        # Antet
        antet_p = doc.add_paragraph()
        stil_titlu(antet_p, "AGENTIA DE PLATI SI INTERVENTIE PENTRU AGRICULTURA", 11, color=(39, 174, 96))
        subant_p = doc.add_paragraph()
        stil_titlu(subant_p, "Centrul Judetean Gorj — Str. I.C. Pompilian nr. 51, Targu Jiu", 9, bold=False)
        doc.add_paragraph()

        # Titlu raport
        titlu_p = doc.add_paragraph()
        stil_titlu(titlu_p, "RAPORT DE CONTROL UAV — SISTEM AGROVISION", 14, color=(41, 128, 185))
        doc.add_paragraph()

        # Date generale
        tabel_info = doc.add_table(rows=0, cols=2)
        tabel_info.style = "Table Grid"
        info = [
            ("Cod parcela",         rez_drona["cod_parcela"]),
            ("Fermier",             rez_lpis.get("fermier", "N/A")),
            ("UAT",                 rez_lpis.get("uat", "N/A")),
            ("Cultura declarata",   rez_lpis.get("cultura_declarata", "N/A")),
            ("Suprafata declarata", f"{rez_lpis.get('suprafata_declarata', 0):.2f} ha"),
            ("Data controlului",    date.today().strftime("%d.%m.%Y")),
            ("Ora procesare",       datetime.now().strftime("%H:%M:%S")),
            ("Rezolutie imagine",   rez_drona["rezolutie"]),
        ]
        for et, val in info:
            rand = tabel_info.add_row()
            rand.cells[0].text = et
            rand.cells[1].text = val
            rand.cells[0].paragraphs[0].runs[0].bold = True
            rand.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            rand.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph()

        # Rezultate drone
        doc.add_heading("1. Rezultate analiza drone", level=1)
        tabel_indici = doc.add_table(rows=0, cols=2)
        tabel_indici.style = "Table Grid"
        indici = [
            ("Indice ExG (vegetatie)",  f"{rez_drona['exg_mediu']:.4f}"),
            ("Indice VARI",             f"{rez_drona['vari_mediu']:.4f}"),
            ("Indice GLI",              f"{rez_drona['gli_mediu']:.4f}"),
            ("Procent vegetatie",       f"{rez_drona['pct_vegetatie']:.1f}%"),
        ]
        for et, val in indici:
            rand = tabel_indici.add_row()
            rand.cells[0].text = et
            rand.cells[1].text = val
            rand.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            rand.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph()

        # Verdict
        doc.add_heading("2. Verdict conformitate PAC", level=1)
        culori_verdict = {"CONFORM": "CONFORM", "ATENTIE": "ATENTIE", "NECONFORM": "NECONFORM"}
        verdict_p = doc.add_paragraph()
        run_v = verdict_p.add_run(f"VERDICT: {rez_conf['verdict']} (penalizare: {rez_conf['penalizare_pct']}%)")
        run_v.bold = True
        run_v.font.size = Pt(12)
        if rez_conf["verdict"] == "CONFORM":
            run_v.font.color.rgb = RGBColor(39, 174, 96)
        elif rez_conf["verdict"] == "ATENTIE":
            run_v.font.color.rgb = RGBColor(230, 126, 34)
        else:
            run_v.font.color.rgb = RGBColor(231, 76, 60)

        doc.add_paragraph(f"Explicatie: {rez_conf['explicatie']}")
        doc.add_paragraph(f"Recomandat: {rez_conf['recomandat']}")
        doc.add_paragraph()

        # Baza legala
        doc.add_heading("3. Baza legala", level=1)
        doc.add_paragraph("- Regulamentul (UE) 2021/2116 al Parlamentului European si al Consiliului")
        doc.add_paragraph("- Regulamentul (UE) 2022/1173 — sistemul integrat de administrare si control (IACS)")
        doc.add_paragraph("- Ordinul MADR privind utilizarea sistemelor UAV in controalele APIA")
        doc.add_paragraph()

        # Semnatura
        doc.add_heading("4. Semnatura inspector", level=1)
        semn = doc.add_table(rows=2, cols=2)
        semn.style = "Table Grid"
        semn.rows[0].cells[0].text = "Inspector responsabil:"
        semn.rows[0].cells[1].text = "Prof. Asoc. Dr. Oliviu M. Gamulescu"
        semn.rows[1].cells[0].text = "Data:"
        semn.rows[1].cells[1].text = date.today().strftime("%d.%m.%Y")
        for rand in semn.rows:
            for celula in rand.cells:
                celula.paragraphs[0].runs[0].font.size = Pt(9)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()


# ─── ORCHESTRATOR ─────────────────────────────────────────────────────────────

class OrchestratorAGROVISION:
    """
    Coordoneaza toti agentii in ordine.
    Nu stie detaliile fiecarui agent — doar ii apeleaza si paseaza rezultatele.
    """

    def __init__(self):
        self.agent_drona        = AgentDrona()
        self.agent_lpis         = AgentLPIS()
        self.agent_conformitate = AgentConformitate()
        self.agent_raportare    = AgentRaportare()

    def ruleaza(self, imagine: Image.Image, cod_parcela: str,
                status_ui=None) -> dict:
        """
        Pipeline complet:
        imagine + cod_parcela → rez_drona → rez_lpis → rez_conf → raport_word
        """
        rezultate = {}

        def log(msg):
            if status_ui:
                status_ui.write(msg)

        # Agent 1
        log(f"[{AgentDrona.ICONA}] {AgentDrona.NUME} — analizeaza imaginea...")
        time.sleep(0.4)
        rezultate["drona"] = self.agent_drona.proceseaza(imagine, cod_parcela)
        log(f"   ExG={rezultate['drona']['exg_mediu']:.3f} | Vegetatie={rezultate['drona']['pct_vegetatie']:.1f}%")

        # Agent 2
        log(f"[{AgentLPIS.ICONA}] {AgentLPIS.NUME} — cauta parcela in LPIS...")
        time.sleep(0.3)
        rezultate["lpis"] = self.agent_lpis.proceseaza(cod_parcela)
        if rezultate["lpis"]["gasit_in_lpis"]:
            log(f"   Gasit: {rezultate['lpis']['fermier']} | {rezultate['lpis']['cultura_declarata']} | {rezultate['lpis']['suprafata_declarata']} ha")
        else:
            log("   ATENTIE: parcela negasita in LPIS!")

        # Agent 3
        log(f"[{AgentConformitate.ICONA}] {AgentConformitate.NUME} — evalueaza conformitatea PAC...")
        time.sleep(0.3)
        rezultate["conformitate"] = self.agent_conformitate.proceseaza(
            rezultate["drona"], rezultate["lpis"]
        )
        log(f"   Verdict: {rezultate['conformitate']['verdict']} (penalizare {rezultate['conformitate']['penalizare_pct']}%)")

        # Agent 4
        log(f"[{AgentRaportare.ICONA}] {AgentRaportare.NUME} — genereaza raport Word...")
        time.sleep(0.3)
        rezultate["raport_docx"] = self.agent_raportare.proceseaza(
            rezultate["drona"], rezultate["lpis"], rezultate["conformitate"]
        )
        log("   Raport Word generat.")

        return rezultate


# ─── INTERFATA STREAMLIT ──────────────────────────────────────────────────────

st.title("Multi-Agenti Drona — AGROVISION")
st.markdown("**Ziua 37** | Arhitectura multi-agent: 4 agenti specializati coordonati de un orchestrator")

st.info(
    "**Arhitectura:** Orchestrator → AgentDrona → AgentLPIS → AgentConformitate → AgentRaportare  \n"
    "Fiecare agent primeste rezultatul agentului anterior si adauga informatia lui."
)

# ── Sidebar ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.header("Parametri control")
    cod_parcela = st.selectbox("Parcela LPIS", list(LPIS_DB.keys()))
    st.markdown("---")
    st.markdown(f"**Fermier:** {LPIS_DB[cod_parcela]['fermier']}")
    st.markdown(f"**Cultura:** {LPIS_DB[cod_parcela]['cultura_declarata']}")
    st.markdown(f"**Suprafata:** {LPIS_DB[cod_parcela]['suprafata_declarata']} ha")
    st.markdown(f"**UAT:** {LPIS_DB[cod_parcela]['uat']}")

# ── Upload imagine ────────────────────────────────────────────────────────────
col1, col2 = st.columns([2, 1])

with col1:
    imagine_file = st.file_uploader(
        "Incarca imaginea drone (JPG/PNG)",
        type=["jpg", "jpeg", "png"]
    )

with col2:
    st.markdown("**Nu ai o imagine drone?**")
    genereaza_test = st.button("Genereaza imagine test", use_container_width=True)
    st.caption("Creeaza o imagine RGB sintetica pentru demo")

# Imagine test generata
if genereaza_test or "imagine_test" not in st.session_state:
    rand = np.random.RandomState(42)
    img_arr = np.zeros((256, 256, 3), dtype=np.uint8)
    # Zone de vegetatie (verde)
    img_arr[30:120, 20:200, 1] = rand.randint(100, 200, (90, 180))
    img_arr[30:120, 20:200, 0] = rand.randint(20, 80, (90, 180))
    img_arr[30:120, 20:200, 2] = rand.randint(20, 60, (90, 180))
    # Zone de sol
    img_arr[130:220, 40:180, 0] = rand.randint(120, 180, (90, 140))
    img_arr[130:220, 40:180, 1] = rand.randint(100, 150, (90, 140))
    img_arr[130:220, 40:180, 2] = rand.randint(60, 100, (90, 140))
    st.session_state["imagine_test"] = Image.fromarray(img_arr)

if genereaza_test:
    st.success("Imagine test generata (256x256 px, vegetatie + sol)")

# ── Ruleaza orchestratorul ────────────────────────────────────────────────────
st.markdown("---")
porneste = st.button("Porneste toti agentii", type="primary", use_container_width=True)

if porneste:
    # Stabileste imaginea de folosit
    if imagine_file is not None:
        imagine = Image.open(imagine_file).convert("RGB")
    elif "imagine_test" in st.session_state:
        imagine = st.session_state["imagine_test"]
    else:
        st.error("Incarca o imagine sau genereaza una de test.")
        st.stop()

    orchestrator = OrchestratorAGROVISION()

    with st.status("Orchestrator AGROVISION — pornesc agentii...", expanded=True) as status_ui:
        rezultate = orchestrator.ruleaza(imagine, cod_parcela, status_ui)
        verdict   = rezultate["conformitate"]["verdict"]
        culoare_status = "complete" if verdict == "CONFORM" else "error"
        status_ui.update(
            label=f"Toti agentii finalizati — Verdict: {verdict}",
            state=culoare_status,
            expanded=False
        )

    st.session_state["rezultate_agenti"] = rezultate

# ── Afisare rezultate ─────────────────────────────────────────────────────────
if "rezultate_agenti" in st.session_state:
    rez = st.session_state["rezultate_agenti"]
    rez_drona = rez["drona"]
    rez_lpis  = rez["lpis"]
    rez_conf  = rez["conformitate"]

    st.markdown("---")
    st.subheader("Rezultate agenti")

    tab1, tab2, tab3, tab4 = st.tabs([
        "UAV Agent Drona",
        "DB Agent LPIS",
        "CHK Agent Conformitate",
        "DOC Agent Raportare"
    ])

    with tab1:
        st.markdown("**Indici de vegetatie calculati:**")
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("ExG",  f"{rez_drona['exg_mediu']:.4f}")
        col2.metric("VARI", f"{rez_drona['vari_mediu']:.4f}")
        col3.metric("GLI",  f"{rez_drona['gli_mediu']:.4f}")
        col4.metric("Vegetatie", f"{rez_drona['pct_vegetatie']:.1f}%")
        st.image(rez_drona["heatmap_png"], caption="Harti indici vegetatie (ExG / VARI / GLI)")

    with tab2:
        st.markdown("**Date din baza LPIS:**")
        if rez_lpis["gasit_in_lpis"]:
            st.success("Parcela gasita in LPIS")
            df_lpis = pd.DataFrame([{
                "Fermier":             rez_lpis["fermier"],
                "Cultura declarata":   rez_lpis["cultura_declarata"],
                "Suprafata (ha)":      rez_lpis["suprafata_declarata"],
                "UAT":                 rez_lpis["uat"],
            }])
            st.dataframe(df_lpis, use_container_width=True)
        else:
            st.error("Parcela NEGASITA in LPIS — cerere invalidata!")

    with tab3:
        verdict = rez_conf["verdict"]
        if verdict == "CONFORM":
            st.success(f"VERDICT: {verdict}")
        elif verdict == "ATENTIE":
            st.warning(f"VERDICT: {verdict}")
        else:
            st.error(f"VERDICT: {verdict}")

        col1, col2 = st.columns(2)
        col1.metric("Penalizare", f"{rez_conf['penalizare_pct']}%")
        col2.metric("Cod", rez_conf["cod"])
        st.markdown(f"**Explicatie:** {rez_conf['explicatie']}")
        st.markdown(f"**Recomandat:** {rez_conf['recomandat']}")

    with tab4:
        st.markdown("**Raport Word generat de AgentRaportare:**")
        st.download_button(
            label="Descarca Raport Word",
            data=rez["raport_docx"],
            file_name=f"Raport_Control_{cod_parcela}_{date.today().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary"
        )
        st.info("Raportul contine: date parcela, indici drone, verdict PAC, baza legala, semnatura.")

# ── Teorie ────────────────────────────────────────────────────────────────────
with st.expander("Concepte cheie Ziua 37"):
    st.markdown("""
    ### Arhitectura multi-agent

    **Fiecare agent = clasa Python cu metoda proceseaza()**
    ```python
    class AgentDrona:
        def proceseaza(self, imagine, cod_parcela) -> dict:
            # analiza imagine
            return {"exg_mediu": ..., "pct_vegetatie": ..., ...}

    class AgentLPIS:
        def proceseaza(self, cod_parcela) -> dict:
            # cauta in baza de date
            return {"fermier": ..., "cultura_declarata": ..., ...}
    ```

    **Orchestratorul nu stie detaliile — doar apeleaza agentii:**
    ```python
    rez_drona = self.agent_drona.proceseaza(imagine, cod)
    rez_lpis  = self.agent_lpis.proceseaza(cod)
    rez_conf  = self.agent_conformitate.proceseaza(rez_drona, rez_lpis)
    raport    = self.agent_raportare.proceseaza(rez_drona, rez_lpis, rez_conf)
    ```

    **st.status() — progres in timp real:**
    ```python
    with st.status("Procesez...", expanded=True) as s:
        s.write("Pasul 1...")
        # lucru
        s.update(label="Gata!", state="complete")
    ```

    **De ce multi-agent?**
    - Fiecare agent poate fi inlocuit independent (ex: AgentLPIS real cu API APIA)
    - Usor de testat fiecare agent separat
    - Scalabil: poti adauga AgentMeteo, AgentSentinel2, AgentNotificare
    """)
