"""
AGROVISION — Raport Control Teren GIS
Ziua 33 | Autor: Prof. Asoc. Dr. Oliviu Mihnea Gamulescu

Scop:
    Generare raport oficial Word pentru control pe teren APIA.
    Selectie parcele, ruta optima, coordonate Stereo 70,
    exportul in format docx gata de semnat si arhivat.

Combina:
    - Analiza spatiala (Ziua 32) — ruta greedy, haversine
    - Export GIS (Ziua 31) — coordonate Stereo 70
    - Rapoarte Word (Ziua 20) — document oficial APIA
"""

import streamlit as st
import folium
from folium import plugins
from streamlit_folium import st_folium
import pandas as pd
import io
import math
import zipfile
from datetime import date, datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ─── CONFIGURARE PAGINA ───────────────────────────────────────────────────────
st.set_page_config(
    page_title="Raport Control Teren | AGROVISION",
    page_icon="GIS",
    layout="wide"
)

# ─── DATE LPIS GORJ ───────────────────────────────────────────────────────────
PARCELE = [
    {"cod": "GJ_78258-1675", "fermier": "Popescu Ion",      "suprafata": 4.32,
     "cultura": "grau",    "uat": "Targu Jiu",    "lat": 45.0421, "lon": 23.2718,
     "status": "CONFORM",    "risc": 1, "ndvi": 0.72},
    {"cod": "GJ_78301-0892", "fermier": "Ionescu Maria",    "suprafata": 6.78,
     "cultura": "porumb",  "uat": "Rovinari",     "lat": 44.9183, "lon": 23.1645,
     "status": "CONFORM",    "risc": 1, "ndvi": 0.68},
    {"cod": "GJ_78445-2341", "fermier": "Dumitrescu Gh.",   "suprafata": 3.15,
     "cultura": "rapita",  "uat": "Motru",        "lat": 44.8067, "lon": 22.9876,
     "status": "NECONFORM",  "risc": 3, "ndvi": 0.31},
    {"cod": "GJ_78512-0077", "fermier": "Stanescu Petre",   "suprafata": 8.90,
     "cultura": "grau",    "uat": "Bumbesti-Jiu", "lat": 45.1823, "lon": 23.3912,
     "status": "CONFORM",    "risc": 2, "ndvi": 0.61},
    {"cod": "GJ_78634-1129", "fermier": "Olteanu Vasile",   "suprafata": 2.44,
     "cultura": "lucerna", "uat": "Novaci",       "lat": 45.3012, "lon": 23.6734,
     "status": "CONFORM",    "risc": 1, "ndvi": 0.79},
    {"cod": "GJ_78720-3388", "fermier": "Marinescu Ana",    "suprafata": 5.67,
     "cultura": "floarea", "uat": "Targu Jiu",    "lat": 45.0198, "lon": 23.2456,
     "status": "NECONFORM",  "risc": 3, "ndvi": 0.28},
    {"cod": "GJ_79001-0445", "fermier": "Petrescu Dan",     "suprafata": 7.23,
     "cultura": "grau",    "uat": "Turceni",      "lat": 44.8734, "lon": 23.4012,
     "status": "CONFORM",    "risc": 2, "ndvi": 0.58},
    {"cod": "GJ_79234-1876", "fermier": "Constantin Gh.",   "suprafata": 1.98,
     "cultura": "lucerna", "uat": "Aninoasa",     "lat": 45.0867, "lon": 23.5219,
     "status": "CONFORM",    "risc": 1, "ndvi": 0.81},
    {"cod": "GJ_79567-2290", "fermier": "Florescu Elena",   "suprafata": 9.45,
     "cultura": "porumb",  "uat": "Rovinari",     "lat": 44.9045, "lon": 23.1823,
     "status": "NECONFORM",  "risc": 3, "ndvi": 0.24},
    {"cod": "GJ_80980-2611", "fermier": "Draghici Marin",   "suprafata": 6.64,
     "cultura": "lucerna", "uat": "Targu Jiu",    "lat": 45.0534, "lon": 23.2901,
     "status": "CONFORM",    "risc": 1, "ndvi": 0.77},
]

RISC_LABEL = {1: "SCAZUT", 2: "MEDIU", 3: "RIDICAT"}
RISC_CULOARE = {1: "#28a745", 2: "#ffc107", 3: "#dc3545"}


# ─── FUNCTII SPATIALE ─────────────────────────────────────────────────────────

def distanta_km(lat1: float, lon1: float, lat2: float, lon2: float) -> float:
    """Distanta Haversine intre doua puncte GPS, in kilometri."""
    R = 6371.0
    phi1, phi2 = math.radians(lat1), math.radians(lat2)
    dphi = math.radians(lat2 - lat1)
    dlam = math.radians(lon2 - lon1)
    a = math.sin(dphi / 2)**2 + math.cos(phi1) * math.cos(phi2) * math.sin(dlam / 2)**2
    return R * 2 * math.asin(math.sqrt(a))


def ruta_optima_greedy(parcele: list) -> list:
    """
    Ruta greedy — incepe de la parcela cu risc maxim,
    continua la cea mai apropiata nevizitata.
    Returneaza lista de indici in ordinea vizitarii.
    """
    if not parcele:
        return []
    # Punct de start = parcela cu risc maxim
    start = max(range(len(parcele)), key=lambda i: parcele[i]["risc"])
    vizitate = [start]
    ramase = set(range(len(parcele))) - {start}
    while ramase:
        curent = vizitate[-1]
        cel_mai_aproape = min(
            ramase,
            key=lambda j: distanta_km(
                parcele[curent]["lat"], parcele[curent]["lon"],
                parcele[j]["lat"],      parcele[j]["lon"]
            )
        )
        vizitate.append(cel_mai_aproape)
        ramase.remove(cel_mai_aproape)
    return vizitate


def wgs84_to_stereo70(lat: float, lon: float):
    """
    Conversie aproximativa WGS84 -> Stereo 70.
    Daca pyproj e disponibil, foloseste transformatorul exact.
    Altfel, returneaza None.
    """
    try:
        from pyproj import Transformer
        t = Transformer.from_crs("EPSG:4326", "EPSG:31700", always_xy=True)
        x, y = t.transform(lon, lat)
        return round(x, 2), round(y, 2)
    except Exception:
        return None, None


def construieste_harta_teren(parcele_selectate: list, ruta: list) -> folium.Map:
    """Harta cu parcele selectate, marcate dupa risc, cu ruta trasa."""
    if not parcele_selectate:
        centru = [45.05, 23.28]
        return folium.Map(location=centru, zoom_start=10, tiles="CartoDB positron")

    centru_lat = sum(p["lat"] for p in parcele_selectate) / len(parcele_selectate)
    centru_lon = sum(p["lon"] for p in parcele_selectate) / len(parcele_selectate)

    m = folium.Map(location=[centru_lat, centru_lon],
                   zoom_start=10, tiles="CartoDB positron")

    # Traseaza ruta
    if len(ruta) > 1:
        coords_ruta = [
            [parcele_selectate[i]["lat"], parcele_selectate[i]["lon"]]
            for i in ruta
        ]
        folium.PolyLine(
            coords_ruta,
            color="#0052A5",
            weight=3,
            opacity=0.8,
            tooltip="Ruta inspectie"
        ).add_to(m)

    # Marcheaza parcelele
    for ord_viz, idx in enumerate(ruta):
        p = parcele_selectate[idx]
        culoare = RISC_CULOARE[p["risc"]]
        x_s, y_s = wgs84_to_stereo70(p["lat"], p["lon"])
        stereo_txt = f"X={x_s:,.0f} / Y={y_s:,.0f}" if x_s else "N/A"

        popup_html = f"""
        <div style="font-family:Arial;font-size:12px;min-width:220px;">
            <b style="color:#0052A5">#{ord_viz + 1} — {p['cod']}</b><br>
            <b>Fermier:</b> {p['fermier']}<br>
            <b>UAT:</b> {p['uat']}<br>
            <b>Cultura:</b> {p['cultura']} | {p['suprafata']} ha<br>
            <b>NDVI:</b> {p['ndvi']}<br>
            <b>Risc:</b> <span style="color:{culoare};font-weight:bold">
                {RISC_LABEL[p['risc']]}
            </span><br>
            <b>Stereo 70:</b> {stereo_txt}
        </div>
        """
        folium.CircleMarker(
            location=[p["lat"], p["lon"]],
            radius=12,
            color=culoare,
            fill=True,
            fill_color=culoare,
            fill_opacity=0.8,
            popup=folium.Popup(popup_html, max_width=260),
            tooltip=f"#{ord_viz+1} {p['fermier']} | {RISC_LABEL[p['risc']]}"
        ).add_to(m)

        # Numar ordine vizitare
        folium.Marker(
            location=[p["lat"], p["lon"]],
            icon=folium.DivIcon(
                html=f'<div style="font-weight:bold;font-size:11px;'
                     f'color:white;background:{culoare};border-radius:50%;'
                     f'width:20px;height:20px;display:flex;'
                     f'align-items:center;justify-content:center;'
                     f'margin-top:-10px;margin-left:-10px;">'
                     f'{ord_viz + 1}</div>',
                icon_size=(20, 20)
            )
        ).add_to(m)

    plugins.Fullscreen(position="topright").add_to(m)
    plugins.MeasureControl(position="bottomleft",
                            primary_length_unit="meters").add_to(m)
    return m


# ─── GENERARE RAPORT WORD ────────────────────────────────────────────────────

def genereaza_raport_word(
    parcele_selectate: list,
    ruta: list,
    inspector: str,
    nr_raport: str,
    data_control: date
) -> bytes:
    """
    Genereaza document Word oficial pentru control pe teren APIA.
    Format: antet institutie, date inspector, tabel parcele cu Stereo 70.
    """
    doc = Document()

    # Margini pagina
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── ANTET ──────────────────────────────────────────────────────────────────
    antet = doc.add_paragraph()
    antet.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = antet.add_run("AGENTIA DE PLATI SI INTERVENTIE PENTRU AGRICULTURA")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 82, 165)

    doc.add_paragraph(
        "Centrul Judetean Gorj — Serviciul Control pe Teren"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── TITLU RAPORT ───────────────────────────────────────────────────────────
    titlu = doc.add_paragraph()
    titlu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = titlu.add_run(f"RAPORT DE CONTROL PE TEREN Nr. {nr_raport}")
    run_t.bold = True
    run_t.font.size = Pt(14)

    subtitlu = doc.add_paragraph()
    subtitlu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitlu.add_run(
        f"Data: {data_control.strftime('%d.%m.%Y')} | "
        f"Inspector: {inspector}"
    ).italic = True

    doc.add_paragraph()

    # ── SECTIUNEA 1: DATE GENERALE ─────────────────────────────────────────────
    h1 = doc.add_paragraph()
    r1 = h1.add_run("1. DATE GENERALE")
    r1.bold = True
    r1.font.size = Pt(11)

    conforme    = sum(1 for p in parcele_selectate if p["status"] == "CONFORM")
    neconforme  = len(parcele_selectate) - conforme
    sup_total   = sum(p["suprafata"] for p in parcele_selectate)
    dist_total  = sum(
        distanta_km(
            parcele_selectate[ruta[i]]["lat"], parcele_selectate[ruta[i]]["lon"],
            parcele_selectate[ruta[i+1]]["lat"], parcele_selectate[ruta[i+1]]["lon"]
        )
        for i in range(len(ruta) - 1)
    ) if len(ruta) > 1 else 0

    date_gen = [
        ("Nr. parcele selectate",     str(len(parcele_selectate))),
        ("Suprafata totala",           f"{sup_total:.2f} ha"),
        ("Parcele CONFORM",            str(conforme)),
        ("Parcele NECONFORM",          str(neconforme)),
        ("Distanta totala ruta",       f"{dist_total:.1f} km"),
        ("UAT-uri acoperite",          str(len({p['uat'] for p in parcele_selectate}))),
        ("Data intocmirii raportului", data_control.strftime("%d.%m.%Y")),
        ("Inspector",                  inspector),
    ]

    tbl_gen = doc.add_table(rows=1, cols=2)
    tbl_gen.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_gen.style = "Table Grid"
    hdr = tbl_gen.rows[0].cells
    hdr[0].text = "Parametru"
    hdr[1].text = "Valoare"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True

    for param, val in date_gen:
        row = tbl_gen.add_row().cells
        row[0].text = param
        row[1].text = val

    doc.add_paragraph()

    # ── SECTIUNEA 2: PARCELE CONTROLATE ────────────────────────────────────────
    h2 = doc.add_paragraph()
    r2 = h2.add_run("2. LISTA PARCELE CONTROLATE (ordine inspectie)")
    r2.bold = True
    r2.font.size = Pt(11)

    coloane = ["Nr.", "Cod LPIS", "Fermier", "UAT", "Cultura",
               "Sup. (ha)", "NDVI", "Risc", "Status",
               "X Stereo 70", "Y Stereo 70"]

    tbl = doc.add_table(rows=1, cols=len(coloane))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    # Header
    hdr_cells = tbl.rows[0].cells
    for i, col in enumerate(coloane):
        hdr_cells[i].text = col
        run_h = hdr_cells[i].paragraphs[0].runs[0]
        run_h.bold = True
        run_h.font.size = Pt(8)

    # Date parcele in ordinea rutei
    for ord_viz, idx in enumerate(ruta):
        p = parcele_selectate[idx]
        x_s, y_s = wgs84_to_stereo70(p["lat"], p["lon"])
        x_txt = f"{x_s:,.0f}" if x_s else "N/A"
        y_txt = f"{y_s:,.0f}" if y_s else "N/A"

        row = tbl.add_row().cells
        valori = [
            str(ord_viz + 1),
            p["cod"],
            p["fermier"],
            p["uat"],
            p["cultura"],
            f"{p['suprafata']:.2f}",
            str(p["ndvi"]),
            RISC_LABEL[p["risc"]],
            p["status"],
            x_txt,
            y_txt,
        ]
        for i, val in enumerate(valori):
            row[i].paragraphs[0].add_run(val).font.size = Pt(8)

    doc.add_paragraph()

    # ── SECTIUNEA 3: RUTA OPTIMA ───────────────────────────────────────────────
    h3 = doc.add_paragraph()
    r3 = h3.add_run("3. RUTA OPTIMA DE INSPECTIE")
    r3.bold = True
    r3.font.size = Pt(11)

    ruta_txt = " → ".join(
        f"#{i+1} {parcele_selectate[idx]['fermier']}"
        for i, idx in enumerate(ruta)
    )
    doc.add_paragraph(ruta_txt)

    if len(ruta) > 1:
        doc.add_paragraph()
        doc.add_paragraph("Distante intre puncte succesive:")
        for i in range(len(ruta) - 1):
            p1 = parcele_selectate[ruta[i]]
            p2 = parcele_selectate[ruta[i+1]]
            d = distanta_km(p1["lat"], p1["lon"], p2["lat"], p2["lon"])
            doc.add_paragraph(
                f"  {p1['fermier']} → {p2['fermier']}: {d:.2f} km",
                style="List Bullet"
            )
        doc.add_paragraph(f"  TOTAL: {dist_total:.1f} km", style="List Bullet")

    doc.add_paragraph()

    # ── SECTIUNEA 4: CONCLUZII ─────────────────────────────────────────────────
    h4 = doc.add_paragraph()
    r4 = h4.add_run("4. CONCLUZII SI RECOMANDARI")
    r4.bold = True
    r4.font.size = Pt(11)

    risc_max = [p for p in parcele_selectate if p["risc"] == 3]
    if risc_max:
        doc.add_paragraph(
            f"Au fost identificate {len(risc_max)} parcele cu risc PAC RIDICAT, "
            f"care necesita verificare prioritara la teren: "
            + ", ".join(p["cod"] for p in risc_max) + "."
        )
    if neconforme > 0:
        doc.add_paragraph(
            f"{neconforme} parcele au status NECONFORM — "
            "se recomanda verificare documente justificative si masuratori GPS pe teren."
        )

    doc.add_paragraph()

    # ── SEMNATURA ──────────────────────────────────────────────────────────────
    sem = doc.add_paragraph()
    sem.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sem.add_run(
        f"Inspector: {inspector}\n"
        f"Data: {data_control.strftime('%d.%m.%Y')}\n\n"
        "Semnatura: ___________________"
    )

    doc.add_paragraph()
    doc.add_paragraph()

    nota = doc.add_paragraph()
    nota.add_run(
        "Generat automat de AGROVISION — Ziua 33 | "
        "Coordonate Stereo 70 (EPSG:31700) | "
        "Prof. Asoc. Dr. Oliviu Mihnea Gamulescu, UCB Targu Jiu 2026"
    ).font.size = Pt(8)
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Salveaza in buffer
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# INTERFATA STREAMLIT
# ═══════════════════════════════════════════════════════════════════════════════

st.title("Ziua 33 - Raport Control Teren GIS")
st.markdown(
    "**Document oficial Word** pentru control pe teren APIA — "
    "selectie parcele, ruta optima, coordonate Stereo 70, export docx."
)

# ─── SIDEBAR ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.header("Configurare Raport")

    inspector = st.text_input(
        "Inspector",
        value="Prof. Asoc. Dr. Oliviu Mihnea Gamulescu",
        help="Numele inspectorului care efectueaza controlul"
    )
    nr_raport = st.text_input(
        "Nr. raport",
        value=f"CT-{date.today().strftime('%Y%m%d')}-001"
    )
    data_ctrl = st.date_input("Data control", value=date.today())

    st.divider()
    st.markdown("**Filtru risc:**")
    filtre_risc = st.multiselect(
        "Arata parcele cu risc",
        options=["SCAZUT", "MEDIU", "RIDICAT"],
        default=["MEDIU", "RIDICAT"],
        help="Selecteaza nivelurile de risc de inclus"
    )

    st.divider()
    st.markdown("**Legenda risc:**")
    for cod, lab in RISC_LABEL.items():
        culoare = RISC_CULOARE[cod]
        st.markdown(
            f'<span style="color:{culoare};font-size:16px;">●</span> {lab}',
            unsafe_allow_html=True
        )

# ─── FILTREAZA PARCELE ────────────────────────────────────────────────────────
risc_map = {"SCAZUT": 1, "MEDIU": 2, "RIDICAT": 3}
coduri_risc = [risc_map[r] for r in filtre_risc]
parcele_filtrate = [p for p in PARCELE if p["risc"] in coduri_risc]

# ─── TABS ─────────────────────────────────────────────────────────────────────
tab1, tab2, tab3, tab4 = st.tabs([
    "Selectie Parcele",
    "Harta Ruta",
    "Raport Word",
    "Teorie GIS"
])

# ── TAB 1: SELECTIE ───────────────────────────────────────────────────────────
with tab1:
    st.subheader("Selectie Parcele pentru Control")

    if not parcele_filtrate:
        st.warning("Nicio parcela corespunde filtrelor selectate. Alege alt nivel de risc din sidebar.")
        st.stop()

    # KPI
    k1, k2, k3, k4 = st.columns(4)
    with k1:
        st.metric("Parcele disponibile", len(parcele_filtrate))
    with k2:
        neconf_f = sum(1 for p in parcele_filtrate if p["status"] == "NECONFORM")
        st.metric("Neconforme", neconf_f)
    with k3:
        sup_f = sum(p["suprafata"] for p in parcele_filtrate)
        st.metric("Suprafata totala", f"{sup_f:.1f} ha")
    with k4:
        risc_f = sum(1 for p in parcele_filtrate if p["risc"] == 3)
        st.metric("Risc RIDICAT", risc_f)

    st.divider()

    # Tabel cu checkbox selectie
    df_f = pd.DataFrame(parcele_filtrate)
    df_f["risc_label"] = df_f["risc"].map(RISC_LABEL)

    def color_status(val):
        if val == "CONFORM":
            return "background-color:#d4edda;color:#155724"
        return "background-color:#f8d7da;color:#721c24"

    def color_risc(val):
        if val == "RIDICAT":
            return "background-color:#f8d7da;color:#721c24;font-weight:bold"
        if val == "MEDIU":
            return "background-color:#fff3cd;color:#856404"
        return "background-color:#d4edda;color:#155724"

    st.dataframe(
        df_f[["cod", "fermier", "uat", "cultura", "suprafata",
              "ndvi", "risc_label", "status"]]
        .rename(columns={
            "cod": "Cod LPIS", "fermier": "Fermier", "uat": "UAT",
            "cultura": "Cultura", "suprafata": "Sup.(ha)",
            "ndvi": "NDVI", "risc_label": "Risc", "status": "Status"
        })
        .style
        .map(color_status, subset=["Status"])
        .map(color_risc, subset=["Risc"])
        .format({"Sup.(ha)": "{:.2f}"}),
        use_container_width=True,
        height=380
    )

    st.info(
        "Parcelele afisate corespund filtrului de risc din sidebar. "
        "Toate parcelele vizibile vor fi incluse in raport si pe harta."
    )

    # Coordonate Stereo 70
    st.subheader("Coordonate Stereo 70 (EPSG:31700)")
    stereo_data = []
    for p in parcele_filtrate:
        x_s, y_s = wgs84_to_stereo70(p["lat"], p["lon"])
        stereo_data.append({
            "Cod LPIS": p["cod"],
            "Fermier":  p["fermier"],
            "Lat WGS84": p["lat"],
            "Lon WGS84": p["lon"],
            "X Stereo 70": f"{x_s:,.0f}" if x_s else "N/A",
            "Y Stereo 70": f"{y_s:,.0f}" if y_s else "N/A",
        })
    st.dataframe(pd.DataFrame(stereo_data), use_container_width=True)


# ── TAB 2: HARTA RUTA ─────────────────────────────────────────────────────────
with tab2:
    st.subheader("Harta Ruta Optima de Inspectie")

    ruta = ruta_optima_greedy(parcele_filtrate)

    # Calcul distante
    dist_segmente = []
    dist_total_km = 0.0
    for i in range(len(ruta) - 1):
        p1 = parcele_filtrate[ruta[i]]
        p2 = parcele_filtrate[ruta[i+1]]
        d = distanta_km(p1["lat"], p1["lon"], p2["lat"], p2["lon"])
        dist_segmente.append({
            "Pas": i + 1,
            "De la":   f"{p1['fermier']} ({p1['uat']})",
            "La":      f"{p2['fermier']} ({p2['uat']})",
            "Dist (km)": round(d, 2)
        })
        dist_total_km += d

    # KPI ruta
    r1, r2, r3 = st.columns(3)
    with r1:
        st.metric("Parcele in ruta", len(ruta))
    with r2:
        st.metric("Distanta totala", f"{dist_total_km:.1f} km")
    with r3:
        st.metric("Nr. segmente", len(dist_segmente))

    st.divider()

    m = construieste_harta_teren(parcele_filtrate, ruta)
    st_folium(m, width=None, height=500, returned_objects=[])

    if dist_segmente:
        st.subheader("Detalii ruta")
        df_ruta = pd.DataFrame(dist_segmente)
        st.dataframe(
            df_ruta.style.format({"Dist (km)": "{:.2f}"}),
            use_container_width=True,
            hide_index=True
        )
        st.markdown(f"**Distanta totala:** {dist_total_km:.1f} km")

        # Export CSV ruta
        csv_ruta = df_ruta.to_csv(index=False, encoding="utf-8")
        st.download_button(
            label="Descarca CSV ruta",
            data=csv_ruta.encode("utf-8"),
            file_name=f"Ruta_Control_{date.today().strftime('%Y%m%d')}.csv",
            mime="text/csv"
        )


# ── TAB 3: RAPORT WORD ────────────────────────────────────────────────────────
with tab3:
    st.subheader("Generare Raport Oficial Word")

    st.markdown("""
    Raportul Word contine:
    - **Antet APIA** — Centrul Judetean Gorj
    - **Date generale** — nr. parcele, suprafata, distanta ruta
    - **Tabel parcele** — cu coordonate Stereo 70 si status PAC
    - **Ruta de inspectie** — ordinea si distantele intre puncte
    - **Concluzii** — parcele cu risc ridicat, recomandari
    - **Semnatura inspector**
    """)

    st.divider()

    ruta_raport = ruta_optima_greedy(parcele_filtrate)

    if not inspector.strip():
        st.warning("Completeaza numele inspectorului in sidebar.")
    else:
        col_r1, col_r2 = st.columns(2)
        with col_r1:
            st.markdown(f"**Inspector:** {inspector}")
            st.markdown(f"**Nr. raport:** {nr_raport}")
            st.markdown(f"**Data:** {data_ctrl.strftime('%d.%m.%Y')}")
            st.markdown(f"**Parcele incluse:** {len(parcele_filtrate)}")
        with col_r2:
            st.markdown(f"**Suprafata totala:** {sum(p['suprafata'] for p in parcele_filtrate):.2f} ha")
            neconf_r = sum(1 for p in parcele_filtrate if p["status"] == "NECONFORM")
            st.markdown(f"**Neconforme PAC:** {neconf_r}")
            risc_rid = sum(1 for p in parcele_filtrate if p["risc"] == 3)
            st.markdown(f"**Risc RIDICAT:** {risc_rid} parcele")

        st.divider()

        docx_bytes = genereaza_raport_word(
            parcele_filtrate, ruta_raport,
            inspector, nr_raport, data_ctrl
        )

        st.download_button(
            label="Descarca Raport Control Teren (.docx)",
            data=docx_bytes,
            file_name=f"Raport_Control_Teren_{nr_raport.replace('/', '-')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary"
        )
        st.success(
            "Raportul este gata. Descarca fisierul .docx, "
            "verifica datele si semneaza manual dupa tiparire."
        )


# ── TAB 4: TEORIE GIS ─────────────────────────────────────────────────────────
with tab4:
    st.subheader("Concepte GIS aplicate in APIA")

    st.markdown("""
### Fluxul GIS pentru control pe teren APIA

```
Date GPS (WGS84)
    ↓
Conversie Stereo 70 (EPSG:31700)
    ↓
Suprapunere pe harti oficiale ANCPI/LPIS
    ↓
Calcul distante Haversine + ruta optima
    ↓
Raport oficial Word + Shapefile
    ↓
Arhivare IACS / dosar fermier
```

### Sistemele de coordonate folosite in Romania

| Sistem | EPSG | Unitate | Utilizare |
|--------|------|---------|-----------|
| WGS84 | 4326 | grade | GPS, Google Maps, drone |
| Stereo 70 | 31700 | metri | harti oficiale Romania, cadastru |
| ETRS89 | 4258 | grade | INSPIRE, reteaua europeana |
| UTM zona 35N | 32635 | metri | harti militare, topografie |

### Formula Haversine — distanta pe suprafata Pamantului

```
d = 2R · arcsin√(sin²(Δφ/2) + cosφ₁ · cosφ₂ · sin²(Δλ/2))

Unde:
  R = 6371 km (raza medie Pamant)
  φ = latitudine in radiani
  λ = longitudine in radiani
  Δφ, Δλ = diferente de latitudine/longitudine
```

### Ruta optima greedy pentru inspectori

Algoritmul folosit:
1. **Start** — parcela cu risc PAC cel mai ridicat
2. **Urmator** — cel mai aproape punct nevizitat (distanta Haversine)
3. **Repeta** pana toate parcelele sunt vizitate

Nu este solutia absolut optima (NP-hard — TSP), dar in practica
este cu **15-20% mai scurta** decat ordinea aleatorie si
se calculeaza instant pentru 10-50 parcele.

### Stereo 70 — de ce este obligatoriu la APIA

- **LPIS** (Land Parcel Identification System) Romania — date in Stereo 70
- **Hartile cadastrale** ANCPI — Stereo 70
- **Masuratori topografice** pentru contestatii PAC — Stereo 70
- Eroarea de conversie WGS84 → Stereo 70: **sub 1 metru** cu pyproj
- Eroarea de conversie manuala (formula simpla): **5-20 metri** (inacceptabila)
    """)

    st.info(
        "Cele 3 module GIS (Zilele 31-33) acopera tot fluxul de lucru "
        "al unui inspector APIA: vizualizare harti oficiale (Z31), "
        "analiza spatiala si rute (Z32), raport oficial generat automat (Z33)."
    )


# ─── FOOTER ───────────────────────────────────────────────────────────────────
st.divider()
st.caption(
    "Ziua 33 - AGROVISION | Raport Control Teren GIS | "
    "Word + Stereo 70 + Ruta Optima | "
    "Prof. Asoc. Dr. Oliviu Mihnea Gamulescu, UCB Targu Jiu 2026"
)
