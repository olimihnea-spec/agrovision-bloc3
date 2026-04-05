"""
BLOC 3 — Deep Learning YOLOv8, Ziua 12
Inferenta cu modelul antrenat — detectie pe imagini noi, parametri, vizualizare, export
Autor: Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj

CONCEPT CHEIE:
  Inferenta = folosesti modelul antrenat (best.pt) pe imagini pe care nu le-a vazut niciodata.
  Nu mai antrenezi nimic — doar "intrebi" modelul: "ce vezi in aceasta imagine?"

  Parametri inferenta:
    conf  = confidence threshold (0.5 = afiseaza doar detectii cu >50% siguranta)
    iou   = IoU threshold pentru NMS (0.45 standard — elimina BBox-uri duplicate)
    imgsz = dimensiunea la care se redimensioneaza imaginea (640 standard)

  Rezultate:
    result.boxes.xyxy   = coordonate absolute BBox (x1, y1, x2, y2)
    result.boxes.conf   = scoruri de incredere per detectie
    result.boxes.cls    = clasa detectata (index intreg)
    result.plot()       = imaginea cu BBox-uri desenate automat
"""

import streamlit as st
import numpy as np
from PIL import Image, ImageDraw, ImageFont
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from io import BytesIO
import random
from collections import Counter
from datetime import date

st.set_page_config(page_title="Inferenta YOLO — Ziua 12", layout="wide")

st.markdown("""
<style>
.bloc3-header {
    background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
    padding: 1.5rem 2rem; border-radius: 12px;
    margin-bottom: 1.5rem; border-left: 5px solid #e94560;
}
.bloc3-header h1 { color: #e94560; margin: 0; font-size: 1.6rem; }
.bloc3-header p  { color: #a8b2d8; margin: 0.3rem 0 0 0; font-size: 0.9rem; }
.concept-box {
    background: #0f3460; border: 1px solid #e94560;
    border-radius: 8px; padding: 1rem 1.2rem; margin: 1rem 0;
    color: #a8b2d8; font-size: 0.88rem;
}
.concept-box b { color: #e94560; }
.ok-box {
    background: #0d2b0d; border: 1px solid #27ae60;
    border-radius: 8px; padding: 0.8rem 1rem; color: #7dcea0; margin: 0.4rem 0;
}
.warn-box {
    background: #2d1b00; border: 1px solid #e67e22;
    border-radius: 8px; padding: 0.8rem 1rem; color: #f39c12; margin: 0.4rem 0;
}
.det-card {
    background: #16213e; border-left: 4px solid #e94560;
    border-radius: 6px; padding: 0.6rem 1rem; margin: 0.3rem 0;
    font-family: monospace; font-size: 0.85rem; color: #a8b2d8;
}
.det-card b { color: #e94560; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="bloc3-header">
  <h1>Ziua 12 — Inferenta cu Modelul Antrenat</h1>
  <p>Bloc 3 YOLOv8 | Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj</p>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div class="concept-box">
<b>CONCEPTE CHEIE AZI:</b><br>
<b>Inferenta</b> = aplici modelul antrenat pe imagini noi (nu mai antrenezi nimic)<br>
<b>conf</b> = confidence threshold — afiseaza doar detectii cu scor > conf (0.5 standard)<br>
<b>IoU</b> = Intersection over Union — masoara cat se suprapun doua BBox-uri<br>
<b>NMS</b> = Non-Maximum Suppression — elimina BBox-uri duplicate pe acelasi obiect<br>
<b>result.boxes</b> = obiect cu toate detectiile: coordonate, clase, scoruri<br>
<b>result.plot()</b> = imaginea cu BBox-uri desenate automat de Ultralytics
</div>
""", unsafe_allow_html=True)

# ── Sectiunea 1: Codul real ──────────────────────────────────────────────────

st.header("1. Codul real pentru inferenta")

st.code("""
from ultralytics import YOLO
from PIL import Image

# Incarca modelul antrenat (best.pt de la Ziua 11)
model = YOLO("runs/train/agrovision_v1/weights/best.pt")

# Inferenta pe o singura imagine
results = model.predict(
    source="imagine_noua.jpg",
    conf=0.5,       # afiseaza doar detectii cu >50% siguranta
    iou=0.45,       # prag NMS — elimina BBox-uri duplicate
    imgsz=640,      # dimensiune imagine
    save=False,     # nu salva automat (gestionam noi rezultatele)
    verbose=False
)

# Acceseaza rezultatele primei imagini
result = results[0]

# Coordonate BBox (x1, y1, x2, y2) — pixeli absoluti
print(result.boxes.xyxy)

# Scoruri de incredere per detectie
print(result.boxes.conf)

# Clasa detectata (index intreg)
print(result.boxes.cls)

# Imaginea cu BBox-uri desenate automat
img_cu_bbox = result.plot()  # numpy array BGR

# Inferenta pe un folder intreg
results_batch = model.predict(source="folder_imagini/", conf=0.5, save=True)
""", language="python")

# ── Sectiunea 2: Demo inferenta ──────────────────────────────────────────────

st.header("2. Demo inferenta (simulare vizuala)")

st.markdown("""
<div class="concept-box">
<b>Cum functioneaza demo-ul:</b> Simulam detectiile unui model antrenat pe imagini drone agricole.
Imaginea de intrare este o imagine sintetica sau incarcata de tine.
Detectiile sunt simulate cu parametrii pe care ii alegi mai jos.
</div>
""", unsafe_allow_html=True)

# Configurare clase
col_cfg1, col_cfg2 = st.columns(2)
with col_cfg1:
    class_names_in = st.text_input("Clase model", "vegetatie,sol_gol,apa,cultura")
    class_names = [c.strip() for c in class_names_in.split(",") if c.strip()]
    conf_thresh = st.slider("Confidence threshold (conf)", 0.1, 0.9, 0.5, 0.05,
        help="Detectiile cu scor sub acest prag sunt ignorate")
with col_cfg2:
    iou_thresh = st.slider("IoU threshold (NMS)", 0.1, 0.9, 0.45, 0.05,
        help="BBox-uri cu IoU > acest prag sunt eliminate ca duplicate")
    n_detectii_max = st.slider("Max detectii simulate", 3, 20, 8,
        help="Cate obiecte sa simuleze modelul ca a gasit")

# Upload imagine sau sintetica
uploaded_img = st.file_uploader("Incarca imagine drone (optional)", type=["jpg", "jpeg", "png"])

btn_inferenta = st.button("Ruleaza inferenta", type="primary")

if btn_inferenta:
    random.seed(None)  # random de fiecare data
    np.random.seed(None)

    # Pregatire imagine
    if uploaded_img:
        img_pil = Image.open(uploaded_img).convert("RGB")
        # redimensionare daca prea mare
        W, H = img_pil.size
        if max(W, H) > 1280:
            factor = 1280 / max(W, H)
            img_pil = img_pil.resize((int(W*factor), int(H*factor)), Image.LANCZOS)
    else:
        # Imagine sintetica drone-like
        W, H = 640, 480
        arr = np.zeros((H, W, 3), dtype=np.uint8)
        # fond — sol agricol
        arr[:, :] = [120, 100, 60]
        # zone de vegetatie
        for _ in range(8):
            x0 = random.randint(0, W-100)
            y0 = random.randint(0, H-80)
            x1 = min(x0 + random.randint(60, 150), W)
            y1 = min(y0 + random.randint(50, 120), H)
            green_val = random.randint(60, 120)
            arr[y0:y1, x0:x1] = [30, green_val, 20]
        # o zona de apa
        arr[20:80, 400:560] = [20, 60, 120]
        # zgomot
        noise = np.random.randint(-15, 15, arr.shape, dtype=np.int16)
        arr = np.clip(arr.astype(np.int16) + noise, 0, 255).astype(np.uint8)
        img_pil = Image.fromarray(arr)

    W, H = img_pil.size

    # Simulare detectii
    CULORI = ["#e94560", "#27ae60", "#3498db", "#f39c12", "#9b59b6", "#1abc9c", "#e74c3c", "#2ecc71"]

    detectii_brute = []
    for _ in range(n_detectii_max + random.randint(0, 5)):
        cls_id = random.randint(0, len(class_names)-1)
        conf   = round(random.uniform(0.25, 0.98), 3)
        xc = random.uniform(0.1, 0.9)
        yc = random.uniform(0.1, 0.9)
        w  = random.uniform(0.08, 0.35)
        h  = random.uniform(0.06, 0.30)
        x1 = max(0, xc - w/2)
        y1 = max(0, yc - h/2)
        x2 = min(1.0, xc + w/2)
        y2 = min(1.0, yc + h/2)
        detectii_brute.append({"cls": cls_id, "conf": conf,
                                "x1": x1, "y1": y1, "x2": x2, "y2": y2})

    # Aplica confidence threshold
    dupa_conf = [d for d in detectii_brute if d["conf"] >= conf_thresh]

    # Simulare NMS simplu — elimina BBox-uri cu IoU mare
    def iou(a, b):
        ix1 = max(a["x1"], b["x1"]); iy1 = max(a["y1"], b["y1"])
        ix2 = min(a["x2"], b["x2"]); iy2 = min(a["y2"], b["y2"])
        inter = max(0, ix2-ix1) * max(0, iy2-iy1)
        area_a = (a["x2"]-a["x1"]) * (a["y2"]-a["y1"])
        area_b = (b["x2"]-b["x1"]) * (b["y2"]-b["y1"])
        union = area_a + area_b - inter
        return inter / union if union > 0 else 0

    dupa_conf.sort(key=lambda d: d["conf"], reverse=True)
    detectii_finale = []
    for d in dupa_conf:
        suprapus = False
        for d2 in detectii_finale:
            if d["cls"] == d2["cls"] and iou(d, d2) > iou_thresh:
                suprapus = True
                break
        if not suprapus:
            detectii_finale.append(d)

    # Deseneaza pe imagine
    img_draw = img_pil.copy()
    draw = ImageDraw.Draw(img_draw)

    for det in detectii_finale:
        cls_id = det["cls"]
        conf   = det["conf"]
        color  = CULORI[cls_id % len(CULORI)]
        lbl    = class_names[cls_id] if cls_id < len(class_names) else f"cls_{cls_id}"

        px1 = int(det["x1"] * W); py1 = int(det["y1"] * H)
        px2 = int(det["x2"] * W); py2 = int(det["y2"] * H)

        # BBox
        draw.rectangle([px1, py1, px2, py2], outline=color, width=3)
        # eticheta
        text = f"{lbl} {conf:.2f}"
        tw = len(text) * 7 + 4
        draw.rectangle([px1, py1-18, px1+tw, py1], fill=color)
        draw.text((px1+2, py1-16), text, fill="white")

    # Afisare
    col_img, col_tbl = st.columns([2, 1])
    with col_img:
        st.image(img_draw, caption=f"Inferenta: {len(detectii_finale)} detectii (din {len(detectii_brute)} brute)", use_container_width=True)

    with col_tbl:
        st.markdown("**Detectii finale:**")
        for i, det in enumerate(detectii_finale, 1):
            lbl = class_names[det["cls"]] if det["cls"] < len(class_names) else f"cls_{det['cls']}"
            color = CULORI[det["cls"] % len(CULORI)]
            st.markdown(f'<div class="det-card"><b>#{i} {lbl}</b><br>conf={det["conf"]:.3f}<br>BBox: [{det["x1"]:.2f},{det["y1"]:.2f},{det["x2"]:.2f},{det["y2"]:.2f}]</div>', unsafe_allow_html=True)

    # ── Statistici detectii ──────────────────────────────────────────────────

    st.markdown("---")
    st.subheader("Statistici detectii")

    col_s1, col_s2, col_s3, col_s4 = st.columns(4)
    col_s1.metric("Detectii brute", len(detectii_brute))
    col_s2.metric("Dupa conf filter", len(dupa_conf))
    col_s3.metric("Dupa NMS", len(detectii_finale), help="Detectii unice, fara duplicate")
    eliminate_nms = len(dupa_conf) - len(detectii_finale)
    col_s4.metric("Eliminate NMS", eliminate_nms)

    # Distributie clase detectate
    cls_counter = Counter(d["cls"] for d in detectii_finale)
    if cls_counter:
        fig_det, ax_det = plt.subplots(figsize=(7, 3))
        fig_det.patch.set_facecolor("#16213e")
        ax_det.set_facecolor("#0f3460")

        cls_lbls = [class_names[cid] if cid < len(class_names) else f"cls_{cid}" for cid in sorted(cls_counter)]
        cls_cnts = [cls_counter[cid] for cid in sorted(cls_counter)]
        cols_bar = [CULORI[cid % len(CULORI)] for cid in sorted(cls_counter)]

        bars = ax_det.bar(cls_lbls, cls_cnts, color=cols_bar, edgecolor="white", linewidth=0.5)
        for bar, cnt in zip(bars, cls_cnts):
            ax_det.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05,
                        str(cnt), ha="center", va="bottom", color="white", fontsize=9)

        ax_det.set_title("Distributia claselor detectate", color="white", fontsize=10)
        ax_det.set_ylabel("Nr. detectii", color="white")
        ax_det.tick_params(colors="white")
        for sp in ax_det.spines.values(): sp.set_edgecolor("#0f3460")
        plt.tight_layout()

        buf_det = BytesIO()
        fig_det.savefig(buf_det, dpi=150, bbox_inches="tight", facecolor="#16213e")
        buf_det.seek(0)
        st.image(buf_det, use_container_width=True)
        plt.close(fig_det)

    # Distributie scoruri de incredere
    if detectii_finale:
        fig_conf, ax_conf = plt.subplots(figsize=(7, 2.5))
        fig_conf.patch.set_facecolor("#16213e")
        ax_conf.set_facecolor("#0f3460")
        confs = [d["conf"] for d in detectii_finale]
        ax_conf.hist(confs, bins=10, range=(0, 1), color="#e94560", edgecolor="white", linewidth=0.5, alpha=0.85)
        ax_conf.axvline(conf_thresh, color="yellow", linewidth=2, linestyle="--", label=f"Prag conf={conf_thresh}")
        ax_conf.set_title("Distributia scorurilor de incredere", color="white", fontsize=10)
        ax_conf.set_xlabel("Confidence", color="white")
        ax_conf.set_ylabel("Nr. detectii", color="white")
        ax_conf.legend(fontsize=8, labelcolor="white", facecolor="#0f3460", edgecolor="white")
        ax_conf.tick_params(colors="white")
        for sp in ax_conf.spines.values(): sp.set_edgecolor("#0f3460")
        plt.tight_layout()

        buf_conf = BytesIO()
        fig_conf.savefig(buf_conf, dpi=150, bbox_inches="tight", facecolor="#16213e")
        buf_conf.seek(0)
        st.image(buf_conf, use_container_width=True)
        plt.close(fig_conf)

    # ── Export imagine + raport ──────────────────────────────────────────────

    st.markdown("---")
    st.subheader("Export rezultate")

    col_ex1, col_ex2 = st.columns(2)

    with col_ex1:
        buf_img_dl = BytesIO()
        img_draw.save(buf_img_dl, format="JPEG", quality=95)
        buf_img_dl.seek(0)
        st.download_button(
            "Descarca imagine cu detectii",
            data=buf_img_dl,
            file_name=f"Detectii_YOLO_{date.today().strftime('%Y%m%d')}.jpg",
            mime="image/jpeg"
        )

    with col_ex2:
        if st.button("Genereaza raport Word inferenta"):
            from docx import Document
            from docx.shared import Pt, RGBColor, Cm, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            for section in doc.sections:
                section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

            titlu = doc.add_heading("RAPORT INFERENTA YOLOv8", level=1)
            titlu.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in titlu.runs:
                run.font.color.rgb = RGBColor(0x23, 0x6a, 0x8e)

            doc.add_paragraph(f"Data: {date.today().strftime('%d.%m.%Y')}")
            doc.add_paragraph("Autor: Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj")
            doc.add_paragraph()

            doc.add_heading("Parametri inferenta", level=2)
            tbl_p = doc.add_table(rows=3, cols=2)
            tbl_p.style = "Table Grid"
            for i, (k, v) in enumerate([
                ("Confidence threshold", str(conf_thresh)),
                ("IoU threshold (NMS)", str(iou_thresh)),
                ("Clase model", ", ".join(class_names)),
            ]):
                tbl_p.rows[i].cells[0].text = k
                tbl_p.rows[i].cells[1].text = v

            doc.add_paragraph()
            doc.add_heading("Rezultate inferenta", level=2)
            tbl_r = doc.add_table(rows=4, cols=2)
            tbl_r.style = "Table Grid"
            for i, (k, v) in enumerate([
                ("Detectii brute", str(len(detectii_brute))),
                ("Dupa confidence filter", str(len(dupa_conf))),
                ("Dupa NMS (finale)", str(len(detectii_finale))),
                ("Eliminate NMS", str(eliminate_nms)),
            ]):
                tbl_r.rows[i].cells[0].text = k
                tbl_r.rows[i].cells[1].text = v

            doc.add_paragraph()
            doc.add_heading("Detectii individuale", level=2)
            tbl_d = doc.add_table(rows=1 + len(detectii_finale), cols=4)
            tbl_d.style = "Table Grid"
            for j, hdr in enumerate(["Nr.", "Clasa", "Confidence", "BBox (x1,y1,x2,y2)"]):
                tbl_d.rows[0].cells[j].text = hdr
            for i, det in enumerate(detectii_finale, 1):
                lbl = class_names[det["cls"]] if det["cls"] < len(class_names) else f"cls_{det['cls']}"
                tbl_d.rows[i].cells[0].text = str(i)
                tbl_d.rows[i].cells[1].text = lbl
                tbl_d.rows[i].cells[2].text = f"{det['conf']:.3f}"
                tbl_d.rows[i].cells[3].text = f"[{det['x1']:.3f},{det['y1']:.3f},{det['x2']:.3f},{det['y2']:.3f}]"

            buf_word = BytesIO()
            doc.save(buf_word)
            buf_word.seek(0)
            st.download_button(
                "Descarca raport Word",
                data=buf_word,
                file_name=f"Raport_Inferenta_YOLO_{date.today().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# ── Sectiunea 3: Efectul conf si IoU ────────────────────────────────────────

st.markdown("---")
st.header("3. Efectul parametrilor conf si IoU")

st.markdown("""
<div class="concept-box">
Acesta este cel mai important concept de inteles la inferenta: <b>conf si IoU controleaza
ce vede si ce raporteaza modelul tau.</b>
</div>
""", unsafe_allow_html=True)

col_exp1, col_exp2 = st.columns(2)
with col_exp1:
    st.markdown("""
    **Confidence (conf):**
    | Valoare | Efect |
    |---|---|
    | 0.2 | Gaseste tot, inclusiv fals pozitive |
    | 0.5 | Echilibrat (standard) |
    | 0.8 | Gaseste doar ce e sigur, rata recall mica |

    **Regula:** La explorare → conf mic. La productie → conf mare.
    """)
with col_exp2:
    st.markdown("""
    **IoU / NMS:**
    | Valoare | Efect |
    |---|---|
    | 0.3 | Elimina agresiv — putine BBox-uri |
    | 0.45 | Standard — echilibrat |
    | 0.7 | Permite suprapuneri — obiecte aglomerate |

    **Regula:** Obiecte aglomerate (culturi dense) → IoU mai mare.
    """)

# Grafic vizual conf vs detectii
fig_exp, axes_exp = plt.subplots(1, 2, figsize=(12, 3.5))
fig_exp.patch.set_facecolor("#16213e")
np.random.seed(99)
n_sim = 200
confs_sim = np.random.beta(2, 2, n_sim)

for ax, title, color_above, color_below in [
    (axes_exp[0], "Efect Confidence Threshold", "#27ae60", "#e74c3c"),
    (axes_exp[1], "Distributia Scorurilor Model", "#3498db", "#e94560"),
]:
    ax.set_facecolor("#0f3460")
    ax.set_title(title, color="white", fontsize=9)
    ax.tick_params(colors="white")
    for sp in ax.spines.values(): sp.set_edgecolor("#0f3460")

ax = axes_exp[0]
thresholds = np.linspace(0.1, 0.95, 50)
n_detectii_per_thresh = [(confs_sim >= t).sum() for t in thresholds]
ax.plot(thresholds, n_detectii_per_thresh, color="#e94560", linewidth=2)
ax.axvline(0.5, color="yellow", linewidth=1.5, linestyle="--", label="conf=0.5 (standard)")
ax.fill_between(thresholds, n_detectii_per_thresh, alpha=0.2, color="#e94560")
ax.set_xlabel("Confidence threshold", color="white")
ax.set_ylabel("Nr. detectii acceptate", color="white")
ax.legend(fontsize=8, labelcolor="white", facecolor="#0f3460", edgecolor="white")

ax = axes_exp[1]
ax.hist(confs_sim, bins=20, color="#3498db", edgecolor="white", linewidth=0.4, alpha=0.85)
ax.axvline(0.5, color="yellow", linewidth=1.5, linestyle="--", label="conf=0.5")
ax.set_xlabel("Scor incredere", color="white")
ax.set_ylabel("Nr. detectii", color="white")
ax.legend(fontsize=8, labelcolor="white", facecolor="#0f3460", edgecolor="white")

plt.tight_layout()
buf_exp = BytesIO()
fig_exp.savefig(buf_exp, dpi=150, bbox_inches="tight", facecolor="#16213e")
buf_exp.seek(0)
st.image(buf_exp, use_container_width=True)
plt.close(fig_exp)

# ── Rezumat lectie ────────────────────────────────────────────────────────────

st.markdown("---")
with st.expander("Rezumat Ziua 12 — Ce am invatat"):
    st.markdown("""
**Inferenta YOLOv8** — fluxul complet:

```
imagine_noua.jpg
      |
      v
model = YOLO("best.pt")           # incarca modelul antrenat
results = model.predict(img, conf=0.5, iou=0.45)
      |
      v
result.boxes.xyxy    # coordonate absolute
result.boxes.conf    # scoruri incredere
result.boxes.cls     # clasa detectata
result.plot()        # imagine cu BBox-uri desenate
```

| Pas | Ce face |
|---|---|
| Forward pass | Imaginea trece prin retea, genereaza ~25000 candidati |
| Confidence filter | Elimina candidatii cu scor < conf |
| NMS | Elimina BBox-uri duplicate pe acelasi obiect |
| Rezultat final | Detectii curate, gata de afisare sau export |

**Urmatoarea zi — Ziua 13:** Evaluare model — confusion matrix, PR curve, F1-confidence curve
    """)
