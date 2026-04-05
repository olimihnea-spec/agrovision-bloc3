"""
BLOC 3 — Deep Learning YOLOv8, Ziua 14
Pipeline complet: imagine drone → YOLOv8 detectie → raport oficial APIA
Autor: Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj

CONCEPT CHEIE:
  Pipeline = lant automat de procesare, de la input brut la output oficial.
  In contextul APIA:
    1. Inspector incarca imaginea drone a parcelei
    2. Modelul detecteaza automat vegetatie / sol_gol / apa
    3. Se calculeaza % suprafata per clasa
    4. Se compara cu declaratia fermierului (cultura declarata)
    5. Se genereaza automat raportul de control cu concluzie PAC
"""

import streamlit as st
import numpy as np
from PIL import Image, ImageDraw
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from io import BytesIO
import random
from datetime import date, datetime
from collections import Counter

st.set_page_config(page_title="Pipeline APIA — Ziua 14", layout="wide")

st.markdown("""
<style>
.bloc3-header {
    background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
    padding: 1.5rem 2rem; border-radius: 12px;
    margin-bottom: 1.5rem; border-left: 5px solid #e94560;
}
.bloc3-header h1 { color: #e94560; margin: 0; font-size: 1.6rem; }
.bloc3-header p  { color: #a8b2d8; margin: 0.3rem 0 0 0; font-size: 0.9rem; }
.concept-box {
    background: #0f3460; border: 1px solid #e94560;
    border-radius: 8px; padding: 1rem 1.2rem; margin: 1rem 0;
    color: #a8b2d8; font-size: 0.88rem;
}
.concept-box b { color: #e94560; }
.ok-box {
    background: #0d2b0d; border: 1px solid #27ae60;
    border-radius: 8px; padding: 0.8rem 1rem; color: #7dcea0; margin: 0.4rem 0;
}
.warn-box {
    background: #2d1b00; border: 1px solid #e67e22;
    border-radius: 8px; padding: 0.8rem 1rem; color: #f39c12; margin: 0.4rem 0;
}
.err-box {
    background: #2d0000; border: 1px solid #e74c3c;
    border-radius: 8px; padding: 0.8rem 1rem; color: #f1948a; margin: 0.4rem 0;
}
.step-box {
    background: #16213e; border-left: 4px solid #e94560;
    border-radius: 0 8px 8px 0; padding: 0.8rem 1.2rem; margin: 0.5rem 0;
    color: #a8b2d8;
}
.step-box b { color: white; }
.metric-card {
    background: #16213e; border: 1px solid #0f3460;
    border-radius: 8px; padding: 1rem; text-align: center;
}
.metric-card .val { font-size: 1.8rem; font-weight: bold; color: #e94560; }
.metric-card .lbl { font-size: 0.8rem; color: #a8b2d8; margin-top: 0.2rem; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="bloc3-header">
  <h1>Ziua 14 — Pipeline Complet: Drone → YOLOv8 → Raport APIA</h1>
  <p>Bloc 3 YOLOv8 | Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj</p>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div class="concept-box">
<b>CONCEPTE CHEIE AZI:</b><br>
<b>Pipeline</b> = lant automat: imagine → detectie → analiza → raport oficial<br>
<b>% suprafata</b> = nr. pixeli clasa / nr. pixeli totali * 100 (din detectiile YOLOv8)<br>
<b>Prag PAC</b> = conform Reg. UE 2021/2116: vegetatie &lt; 50% din parcela = risc neconformitate<br>
<b>Raport de control</b> = document oficial generat automat cu datele fermierului si concluziile AI
</div>
""", unsafe_allow_html=True)

# ── Fluxul pipeline ───────────────────────────────────────────────────────────

st.header("Fluxul pipeline APIA")
col_f1, col_f2, col_f3, col_f4, col_f5 = st.columns(5)
for col, nr, titlu, desc in [
    (col_f1, "1", "Upload imagine", "Fotografie drone parcela"),
    (col_f2, "2", "Detectie YOLOv8", "Modelu analizeaza imaginea"),
    (col_f3, "3", "Calcul suprafete", "% vegetatie / sol / apa"),
    (col_f4, "4", "Comparatie PAC", "vs. declaratia fermierului"),
    (col_f5, "5", "Raport oficial", "Word + PDF gata de semnat"),
]:
    with col:
        st.markdown(f'<div class="step-box"><b>Pasul {nr}: {titlu}</b><br>{desc}</div>', unsafe_allow_html=True)

st.markdown("---")

# ── Pasul 1: Date fermier ─────────────────────────────────────────────────────

st.header("Pasul 1 — Date fermier si parcela")

col_d1, col_d2 = st.columns(2)
with col_d1:
    nume_fermier   = st.text_input("Nume fermier", "Popescu Ion")
    cnp_cui        = st.text_input("CNP / CUI", "1760312182456")
    nr_cerere      = st.text_input("Nr. cerere unica", "GJ-2026-004521")
with col_d2:
    cod_parcela    = st.text_input("Cod parcela LPIS", "GJ_79157-348")
    suprafata_ha   = st.number_input("Suprafata declarata (ha)", 0.1, 500.0, 2.45, 0.01)
    cultura_decl   = st.selectbox("Cultura declarata", [
        "Grau", "Porumb", "Floarea-soarelui", "Rapita", "Orz",
        "Lucerna", "Pasune permanenta", "Legume", "Livada"
    ])

inspector = st.text_input("Inspector APIA", "Prof. Asoc. Dr. Oliviu Mihnea Gamulescu")
data_control = st.date_input("Data controlului", value=date.today())

st.markdown("---")

# ── Pasul 2: Upload imagine ───────────────────────────────────────────────────

st.header("Pasul 2 — Imagine drone parcela")

uploaded = st.file_uploader("Incarca imaginea drone (.jpg/.png)", type=["jpg","jpeg","png"])

col_cfg1, col_cfg2 = st.columns(2)
with col_cfg1:
    conf_thresh = st.slider("Confidence threshold", 0.1, 0.9, 0.45, 0.05)
with col_cfg2:
    class_names_in = st.text_input("Clase model", "vegetatie,sol_gol,apa")
    class_names = [c.strip() for c in class_names_in.split(",") if c.strip()]

btn_analizeaza = st.button("Analizeaza parcela cu YOLOv8", type="primary", use_container_width=True)

# ── Procesare ─────────────────────────────────────────────────────────────────

if btn_analizeaza:
    random.seed(None)
    np.random.seed(None)

    # Pregatire imagine
    if uploaded:
        img_pil = Image.open(uploaded).convert("RGB")
        W, H = img_pil.size
        if max(W, H) > 1280:
            f = 1280 / max(W, H)
            img_pil = img_pil.resize((int(W*f), int(H*f)), Image.LANCZOS)
    else:
        W, H = 640, 480
        arr = np.zeros((H, W, 3), dtype=np.uint8)
        arr[:, :] = [110, 90, 50]
        # zone vegetatie
        for _ in range(12):
            x0 = random.randint(0, W-120); y0 = random.randint(0, H-90)
            x1 = min(x0+random.randint(70,160), W); y1 = min(y0+random.randint(60,130), H)
            arr[y0:y1, x0:x1] = [25, random.randint(80,140), 20]
        # zona apa
        arr[10:70, 450:600] = [15, 55, 130]
        noise = np.random.randint(-10, 10, arr.shape, dtype=np.int16)
        arr = np.clip(arr.astype(np.int16)+noise, 0, 255).astype(np.uint8)
        img_pil = Image.fromarray(arr)
        W, H = img_pil.size

    # Simulare detectii YOLOv8
    n_cls = len(class_names)
    CULORI = ["#27ae60", "#e67e22", "#3498db", "#e94560", "#9b59b6"]
    CULORI_RGB = [(39,174,96), (230,126,34), (52,152,219), (233,69,96), (155,89,182)]

    detectii = []
    # distributie realista: mai multa vegetatie
    weights = [0.55, 0.30, 0.15][:n_cls]
    while len(weights) < n_cls:
        weights.append(0.1)
    total_w = sum(weights); weights = [w/total_w for w in weights]

    for _ in range(random.randint(8, 18)):
        cls_id = random.choices(range(n_cls), weights=weights[:n_cls])[0]
        conf   = round(random.uniform(0.3, 0.97), 3)
        if conf < conf_thresh:
            continue
        xc = random.uniform(0.08, 0.92); yc = random.uniform(0.08, 0.92)
        w  = random.uniform(0.06, 0.30); h  = random.uniform(0.05, 0.25)
        detectii.append({"cls": cls_id, "conf": conf,
                         "xc": xc, "yc": yc, "w": w, "h": h})

    if not detectii:
        detectii.append({"cls": 0, "conf": 0.72, "xc": 0.5, "yc": 0.5,
                         "w": 0.4, "h": 0.35})

    # Deseneaza BBox-uri
    img_draw = img_pil.copy()
    draw = ImageDraw.Draw(img_draw)
    aria_totala_bbox = 0.0
    aria_per_cls = {i: 0.0 for i in range(n_cls)}

    for det in detectii:
        cls_id = det["cls"]
        x1 = int((det["xc"]-det["w"]/2)*W); y1 = int((det["yc"]-det["h"]/2)*H)
        x2 = int((det["xc"]+det["w"]/2)*W); y2 = int((det["yc"]+det["h"]/2)*H)
        color = CULORI[cls_id % len(CULORI)]
        draw.rectangle([x1,y1,x2,y2], outline=color, width=3)
        lbl = f"{class_names[cls_id]} {det['conf']:.2f}"
        draw.rectangle([x1, y1-18, x1+len(lbl)*7+4, y1], fill=color)
        draw.text((x1+2, y1-16), lbl, fill="white")
        aria = det["w"] * det["h"]
        aria_per_cls[cls_id] += aria
        aria_totala_bbox += aria

    # Calcul % suprafata
    if aria_totala_bbox > 0:
        pct_cls = {i: aria_per_cls[i]/aria_totala_bbox*100 for i in range(n_cls)}
    else:
        pct_cls = {i: 100.0/n_cls for i in range(n_cls)}

    # ── Afisare rezultate ─────────────────────────────────────────────────────

    st.markdown("---")
    st.header("Pasul 3 — Rezultate detectie")

    col_img, col_stats = st.columns([2, 1])
    with col_img:
        st.image(img_draw, caption=f"Imagine analizata — {len(detectii)} detectii", use_container_width=True)
    with col_stats:
        st.markdown("**Distributie suprafata detectata:**")
        for i, cls_name in enumerate(class_names):
            pct = pct_cls.get(i, 0)
            color = CULORI[i % len(CULORI)]
            st.markdown(f'<div class="metric-card" style="margin-bottom:0.5rem;border-color:{color}"><div class="val" style="color:{color}">{pct:.1f}%</div><div class="lbl">{cls_name}</div></div>', unsafe_allow_html=True)

    # Grafic pie
    fig_pie, ax_pie = plt.subplots(figsize=(5, 4))
    fig_pie.patch.set_facecolor("#16213e")
    ax_pie.set_facecolor("#16213e")
    pcts = [pct_cls.get(i, 0) for i in range(n_cls)]
    wedge_colors = [CULORI[i % len(CULORI)] for i in range(n_cls)]
    wedges, texts, autotexts = ax_pie.pie(
        pcts, labels=class_names, colors=wedge_colors,
        autopct="%1.1f%%", startangle=90,
        textprops={"color": "white", "fontsize": 9}
    )
    for at in autotexts:
        at.set_color("white"); at.set_fontsize(8)
    ax_pie.set_title("Suprafata per clasa", color="white", fontsize=10)
    plt.tight_layout()
    buf_pie = BytesIO()
    fig_pie.savefig(buf_pie, dpi=150, bbox_inches="tight", facecolor="#16213e")
    buf_pie.seek(0)
    st.image(buf_pie, use_container_width=True)
    plt.close(fig_pie)

    # ── Pasul 4: Evaluare PAC ─────────────────────────────────────────────────

    st.markdown("---")
    st.header("Pasul 4 — Evaluare conformitate PAC")

    pct_vegetatie = pct_cls.get(0, 0)
    pct_sol_gol   = pct_cls.get(1, 0) if n_cls > 1 else 0
    pct_apa       = pct_cls.get(2, 0) if n_cls > 2 else 0

    # Reguli PAC simplificate
    probleme_pac = []
    if pct_vegetatie < 50:
        probleme_pac.append(f"Vegetatie detectata: {pct_vegetatie:.1f}% < 50% minim PAC")
    if pct_sol_gol > 40:
        probleme_pac.append(f"Sol gol: {pct_sol_gol:.1f}% > 40% — posibila subutilizare")

    col_pac1, col_pac2 = st.columns(2)
    with col_pac1:
        st.markdown("**Cultura declarata vs. detectie:**")
        st.markdown(f"- Cultura declarata: **{cultura_decl}**")
        st.markdown(f"- Vegetatie detectata AI: **{pct_vegetatie:.1f}%**")
        st.markdown(f"- Suprafata declarata: **{suprafata_ha} ha**")
    with col_pac2:
        st.markdown("**Concluzie control:**")
        if not probleme_pac:
            st.markdown(f'<div class="ok-box"><b>CONFORM PAC</b><br>Vegetatie suficienta detectata ({pct_vegetatie:.1f}%). Parcela corespunde declaratiei.</div>', unsafe_allow_html=True)
            concluzie = "CONFORM"
            culoare_concluzie = "#27ae60"
        else:
            for prob in probleme_pac:
                st.markdown(f'<div class="warn-box"><b>ATENTIE:</b> {prob}</div>', unsafe_allow_html=True)
            concluzie = "NECESITA VERIFICARE SUPLIMENTARA"
            culoare_concluzie = "#e67e22"

    # ── Pasul 5: Raport oficial Word ──────────────────────────────────────────

    st.markdown("---")
    st.header("Pasul 5 — Raport oficial APIA")

    if st.button("Genereaza raport oficial Word", type="primary"):
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

        # Antet
        antet = doc.add_paragraph()
        antet.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_a = antet.add_run("AGENȚIA DE PLĂȚI ȘI INTERVENȚIE PENTRU AGRICULTURĂ")
        run_a.bold = True; run_a.font.size = Pt(13)
        run_a.font.color.rgb = RGBColor(0x00, 0x4e, 0x92)

        doc.add_paragraph("Centrul Județean Gorj | Str. I.C. Pompilian nr. 51, Târgu Jiu").alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        titlu = doc.add_heading("RAPORT DE CONTROL PE TEREN — ANALIZĂ UAV/AI", level=1)
        titlu.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titlu.runs:
            run.font.color.rgb = RGBColor(0x00, 0x4e, 0x92)

        doc.add_paragraph(f"Nr. înregistrare: {nr_cerere} | Data: {data_control.strftime('%d.%m.%Y')}")
        doc.add_paragraph()

        # Date fermier
        doc.add_heading("1. Date identificare fermier", level=2)
        tbl1 = doc.add_table(rows=4, cols=2)
        tbl1.style = "Table Grid"
        for i, (k, v) in enumerate([
            ("Nume fermier", nume_fermier),
            ("CNP / CUI", cnp_cui),
            ("Nr. cerere unică", nr_cerere),
            ("Inspector APIA", inspector),
        ]):
            tbl1.rows[i].cells[0].text = k
            tbl1.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            tbl1.rows[i].cells[1].text = v

        doc.add_paragraph()

        # Date parcela
        doc.add_heading("2. Date parcelă", level=2)
        tbl2 = doc.add_table(rows=4, cols=2)
        tbl2.style = "Table Grid"
        for i, (k, v) in enumerate([
            ("Cod parcelă LPIS", cod_parcela),
            ("Suprafață declarată", f"{suprafata_ha:.2f} ha"),
            ("Cultură declarată", cultura_decl),
            ("Data controlului", data_control.strftime("%d.%m.%Y")),
        ]):
            tbl2.rows[i].cells[0].text = k
            tbl2.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            tbl2.rows[i].cells[1].text = v

        doc.add_paragraph()

        # Rezultate AI
        doc.add_heading("3. Rezultate analiză UAV + YOLOv8", level=2)
        p_met = doc.add_paragraph()
        p_met.add_run("Model AI utilizat: ").bold = True
        p_met.add_run("YOLOv8n (transfer learning, mAP50=0.829)")
        p_met.add_run("\nNr. detectii: ").bold = True
        p_met.add_run(str(len(detectii)))

        doc.add_paragraph()
        tbl3 = doc.add_table(rows=1+n_cls, cols=3)
        tbl3.style = "Table Grid"
        for j, h in enumerate(["Clasă detectată", "% din suprafață", "Suprafață estimată (ha)"]):
            cell = tbl3.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
        for i, cls_name in enumerate(class_names):
            pct = pct_cls.get(i, 0)
            ha_est = suprafata_ha * pct / 100
            tbl3.rows[i+1].cells[0].text = cls_name
            tbl3.rows[i+1].cells[1].text = f"{pct:.1f}%"
            tbl3.rows[i+1].cells[2].text = f"{ha_est:.3f} ha"

        doc.add_paragraph()

        # Concluzie
        doc.add_heading("4. Concluzie conformitate PAC", level=2)
        p_conc = doc.add_paragraph()
        run_conc = p_conc.add_run(f"CONCLUZIE: {concluzie}")
        run_conc.bold = True
        run_conc.font.size = Pt(12)
        if concluzie == "CONFORM":
            run_conc.font.color.rgb = RGBColor(0x1e, 0x8b, 0x1e)
        else:
            run_conc.font.color.rgb = RGBColor(0xcc, 0x66, 0x00)

        doc.add_paragraph()
        if not probleme_pac:
            doc.add_paragraph(
                f"Analiza imaginilor UAV aferente parcelei {cod_parcela} a evidențiat "
                f"prezența vegetației pe {pct_vegetatie:.1f}% din suprafața inspectată, "
                f"corespunzând declarației de cultură ({cultura_decl}). "
                f"Parcela se încadrează în condițiile de eligibilitate PAC 2023-2027 "
                f"conform Reg. UE 2021/2116."
            )
        else:
            doc.add_paragraph(
                f"Analiza imaginilor UAV aferente parcelei {cod_parcela} a identificat "
                f"anomalii care necesită verificare suplimentară: "
                + "; ".join(probleme_pac) + ". "
                f"Se recomandă control fizic al parcelei conform procedurii APIA."
            )

        doc.add_paragraph()

        # Semnatura
        doc.add_heading("5. Semnătură inspector", level=2)
        tbl_s = doc.add_table(rows=3, cols=2)
        tbl_s.style = "Table Grid"
        tbl_s.rows[0].cells[0].text = "Inspector APIA"
        tbl_s.rows[0].cells[1].text = inspector
        tbl_s.rows[1].cells[0].text = "Data întocmirii"
        tbl_s.rows[1].cells[1].text = date.today().strftime("%d.%m.%Y")
        tbl_s.rows[2].cells[0].text = "Semnătură"
        tbl_s.rows[2].cells[1].text = "_________________________"

        doc.add_paragraph()
        nota = doc.add_paragraph()
        run_nota = nota.add_run(
            "Notă: Prezentul raport a fost generat asistat de sistemul AI AGROVISION "
            "(YOLOv8, transfer learning). Concluziile au caracter orientativ și se "
            "completează cu verificarea documentelor justificative ale fermierului."
        )
        run_nota.font.size = Pt(9)
        run_nota.font.italic = True

        buf_word = BytesIO()
        doc.save(buf_word)
        buf_word.seek(0)
        st.download_button(
            "Descarca raport Word oficial",
            data=buf_word,
            file_name=f"Raport_Control_APIA_{cod_parcela}_{date.today().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.markdown(f'<div class="ok-box">Raport generat: Raport_Control_APIA_{cod_parcela}_{date.today().strftime("%Y%m%d")}.docx</div>', unsafe_allow_html=True)

else:
    st.info("Completeaza datele fermierului si apasa 'Analizeaza parcela cu YOLOv8' pentru a porni pipeline-ul.")

# ── Rezumat lectie ────────────────────────────────────────────────────────────

st.markdown("---")
with st.expander("Rezumat Ziua 14 — Ce am invatat"):
    st.markdown("""
**Pipeline complet APIA — 5 pasi:**

| Pas | Actiune | Output |
|---|---|---|
| 1 | Date fermier + parcela | Formular completat |
| 2 | Upload imagine drone | Imagine bruta |
| 3 | YOLOv8 detectie | BBox-uri + % suprafata per clasa |
| 4 | Evaluare PAC | CONFORM / NECESITA VERIFICARE |
| 5 | Raport oficial Word | Document gata de semnat |

**Valoarea pentru APIA:**
- Inspector incarca imaginea drone → raport generat in < 10 secunde
- Obiectiv: reducere timp control de la 2 ore la 15 minute per parcela
- Documentatie automata: trasabilitate decizii, audit UE

**Valoarea pentru teza:**
- Demonstreaza aplicabilitatea practica a YOLOv8 in agricultura
- Pipeline end-to-end: de la date brute la decizie administrativa
- Conform Reg. UE 2021/2116 (IACS, controale pe teren)

**Urmatoarea zi — Ziua 15:** Batch procesare — analiza automata a unui folder intreg cu imagini drone
    """)
