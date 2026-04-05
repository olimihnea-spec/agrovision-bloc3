"""
AGROVISION — Sinteza Finala Bloc 3 YOLOv8
Ziua 30 | Autor: Prof. Asoc. Dr. Oliviu Mihnea Gamulescu

Scop:
    Pagina de sinteza a intregului Bloc 3 (30 zile YOLOv8 + ML + GIS).
    Demo live complet, certificat Word, statistici aplicatie,
    roadmap urmatorii pasi.
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import datetime
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─── CONFIGURARE PAGINA ───────────────────────────────────────────────────────
st.set_page_config(
    page_title="Sinteza Bloc 3 | AGROVISION",
    page_icon="30",
    layout="wide"
)

# ─── DATE BLOC 3 ──────────────────────────────────────────────────────────────
ZILE_BLOC3 = [
    {"ziua": 1,  "titlu": "Arhitectura YOLOv8",         "categorie": "Deep Learning",  "completat": True},
    {"ziua": 2,  "titlu": "Prima Detectie",              "categorie": "Deep Learning",  "completat": True},
    {"ziua": 3,  "titlu": "Detectie Drone",              "categorie": "Deep Learning",  "completat": True},
    {"ziua": 4,  "titlu": "Batch Analiza",               "categorie": "Procesare",      "completat": True},
    {"ziua": 5,  "titlu": "Format Dataset YOLO",         "categorie": "Dataset",        "completat": True},
    {"ziua": 6,  "titlu": "Adnotare LabelImg",           "categorie": "Dataset",        "completat": True},
    {"ziua": 7,  "titlu": "Split Dataset",               "categorie": "Dataset",        "completat": True},
    {"ziua": 8,  "titlu": "Tiling Imagini",              "categorie": "Dataset",        "completat": True},
    {"ziua": 9,  "titlu": "Augmentare Dataset",          "categorie": "Dataset",        "completat": True},
    {"ziua": 10, "titlu": "Validare Dataset",            "categorie": "Dataset",        "completat": True},
    {"ziua": 11, "titlu": "Antrenament YOLOv8",          "categorie": "Deep Learning",  "completat": True},
    {"ziua": 12, "titlu": "Inferenta YOLO",              "categorie": "Deep Learning",  "completat": True},
    {"ziua": 13, "titlu": "Evaluare Model",              "categorie": "Deep Learning",  "completat": True},
    {"ziua": 14, "titlu": "Pipeline APIA",               "categorie": "APIA",           "completat": True},
    {"ziua": 15, "titlu": "Batch Procesare",             "categorie": "Procesare",      "completat": True},
    {"ziua": 16, "titlu": "Comparatie Temporala",        "categorie": "Procesare",      "completat": True},
    {"ziua": 17, "titlu": "Export GIS",                  "categorie": "GIS",            "completat": True},
    {"ziua": 18, "titlu": "Dashboard AGROVISION",        "categorie": "Dashboard",      "completat": True},
    {"ziua": 19, "titlu": "Autentificare Roluri",        "categorie": "Securitate",     "completat": True},
    {"ziua": 20, "titlu": "Rapoarte PDF",                "categorie": "Export",         "completat": True},
    {"ziua": 21, "titlu": "Deployment Cloud",            "categorie": "Cloud",          "completat": True},
    {"ziua": 22, "titlu": "Baza Date SQLite",            "categorie": "Date",           "completat": True},
    {"ziua": 23, "titlu": "Notificari Email",            "categorie": "Comunicare",     "completat": True},
    {"ziua": 24, "titlu": "API REST FastAPI",            "categorie": "API",            "completat": True},
    {"ziua": 25, "titlu": "Hugging Face Hub",            "categorie": "Cloud",          "completat": True},
    {"ziua": 26, "titlu": "Analiza NDVI Spectral",       "categorie": "Spectral",       "completat": True},
    {"ziua": 27, "titlu": "Clustering K-Means",          "categorie": "ML",             "completat": True},
    {"ziua": 28, "titlu": "Random Forest Clasificare",   "categorie": "ML",             "completat": True},
    {"ziua": 29, "titlu": "Random Forest Regresie",      "categorie": "ML",             "completat": True},
    {"ziua": 30, "titlu": "Sinteza Bloc 3",              "categorie": "Sinteza",        "completat": True},
]

COMPETENTE = [
    "Arhitectura si antrenament YOLOv8 (mAP50=0.829)",
    "Pipeline complet drone → detectie → raport APIA",
    "Dataset: adnotare, augmentare, validare, split",
    "Export GIS: GeoJSON, Shapefile Stereo70, GPX",
    "Autentificare roluri SHA-256, securitate Streamlit",
    "Baza de date SQLite, rapoarte PDF, notificari email",
    "API REST FastAPI cu Swagger UI",
    "Model pe Hugging Face Hub, deployment Streamlit Cloud",
    "Indici spectrali: ExG, VARI, GLI, NGRDI, NDVI_sim",
    "Machine Learning: K-Means, Random Forest clasificare+regresie",
]

CULORI_CATEGORIE = {
    "Deep Learning": "#E74C3C",
    "Dataset":       "#E67E22",
    "Procesare":     "#F1C40F",
    "APIA":          "#0052A5",
    "GIS":           "#27AE60",
    "Dashboard":     "#1ABC9C",
    "Securitate":    "#8E44AD",
    "Export":        "#2980B9",
    "Cloud":         "#16A085",
    "Date":          "#D35400",
    "Comunicare":    "#C0392B",
    "API":           "#7F8C8D",
    "Spectral":      "#2ECC71",
    "ML":            "#3498DB",
    "Sinteza":       "#F39C12",
}


def genereaza_certificat() -> bytes:
    """Genereaza certificat Word de absolvire Bloc 3."""
    doc = Document()

    # Margini
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Antet
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSITATEA \"CONSTANTIN BRANCUSI\" TARGU JIU")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 82, 165)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Facultatea de Inginerie | Departamentul Energie, Mediu si Agroturism")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Titlu certificat
    titlu = doc.add_paragraph()
    titlu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = titlu.add_run("CERTIFICAT DE COMPETENTE DIGITALE")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor(0, 82, 165)

    subtitlu = doc.add_paragraph()
    subtitlu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitlu.add_run("Bloc 3 — Deep Learning YOLOv8 si Machine Learning Agricol")
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(52, 73, 94)

    doc.add_paragraph()

    # Beneficiar
    ben = doc.add_paragraph()
    ben.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ben.add_run("Se certifica faptul ca\n").font.size = Pt(11)
    br = ben.add_run("Prof. Asoc. Dr. Oliviu Mihnea Gamulescu")
    br.bold = True
    br.font.size = Pt(15)
    br.font.color.rgb = RGBColor(0, 82, 165)

    ben2 = doc.add_paragraph()
    ben2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    b2r = ben2.add_run(
        "Inspector Principal APIA CJ Gorj | Profesor Asociat UCB Targu Jiu"
    )
    b2r.font.size = Pt(10)
    b2r.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Descriere
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = desc.add_run(
        "a parcurs cu succes programul de 30 de zile\n"
        "\"AGROVISION — YOLOv8 si Machine Learning in Agricultura de Precizie\"\n"
        "dobandind urmatoarele competente:"
    )
    dr.font.size = Pt(11)

    doc.add_paragraph()

    # Competente
    for i, comp in enumerate(COMPETENTE, 1):
        p_comp = doc.add_paragraph(style="List Number")
        p_comp.paragraph_format.left_indent = Inches(0.5)
        rc = p_comp.add_run(comp)
        rc.font.size = Pt(10)

    doc.add_paragraph()

    # Metrici cheie
    metrici = doc.add_paragraph()
    metrici.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = metrici.add_run(
        "Model antrenat: YOLOv8n | mAP50 = 0.829 | Dataset: 7 imagini drone reale Gorj\n"
        "Aplicatie live: https://agrovision-bloc3-8qydbmd2z3zgmpqk4ygtsg.streamlit.app"
    )
    mr.bold = True
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0, 82, 165)

    doc.add_paragraph()

    # Data si semnatura
    data_azi = datetime.date.today().strftime("%d %B %Y")
    tabel_semn = doc.add_table(rows=1, cols=2)
    tabel_semn.columns[0].width = Inches(3)
    tabel_semn.columns[1].width = Inches(3)

    c0 = tabel_semn.cell(0, 0).paragraphs[0]
    c0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c0r = c0.add_run(f"Targu Jiu, {data_azi}")
    c0r.font.size = Pt(10)

    c1 = tabel_semn.cell(0, 1).paragraphs[0]
    c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c1r = c1.add_run("Prof. Asoc. Dr. Oliviu Mihnea Gamulescu\n_______________________")
    c1r.font.size = Pt(10)

    # Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        "Sistem AGROVISION | YOLOv8 | Streamlit | Python | scikit-learn | "
        "UCB Targu Jiu | APIA CJ Gorj | 2026"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(150, 150, 150)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# INTERFATA STREAMLIT
# ═══════════════════════════════════════════════════════════════════════════════

st.title("Ziua 30 - Sinteza Finala Bloc 3 YOLOv8")
st.markdown(
    "**30 de zile de Deep Learning, Machine Learning si GIS aplicat in agricultura.** "
    "Sistem complet AGROVISION — de la imaginea drone la raportul oficial APIA."
)

# ─── KPI BLOC 3 ───────────────────────────────────────────────────────────────
k1, k2, k3, k4, k5, k6 = st.columns(6)
with k1:
    st.metric("Zile parcurse", "30 / 30")
with k2:
    st.metric("Module create", "30")
with k3:
    st.metric("mAP50 model", "0.829")
with k4:
    st.metric("Linii de cod", "~16.000")
with k5:
    st.metric("Competente", "10")
with k6:
    st.metric("Cost total", "0 RON")

st.divider()

# ─── TABS ─────────────────────────────────────────────────────────────────────
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "Progres 30 Zile",
    "Harta Competente",
    "Demo Live",
    "Certificat",
    "Roadmap"
])

# ── TAB 1: PROGRES ────────────────────────────────────────────────────────────
with tab1:
    st.subheader("Toate cele 30 de zile — Bloc 3 YOLOv8")

    df_zile = pd.DataFrame(ZILE_BLOC3)

    # Grafic timeline pe categorii
    fig_timeline = px.bar(
        df_zile,
        x="ziua", y=[1]*30,
        color="categorie",
        color_discrete_map=CULORI_CATEGORIE,
        title="Distributia zilelor pe categorii tematice",
        labels={"ziua": "Ziua", "y": "", "categorie": "Categorie"},
        hover_data=["titlu"]
    )
    fig_timeline.update_layout(
        height=250, showlegend=True,
        yaxis_visible=False, plot_bgcolor="#F8F9FA",
        bargap=0.1
    )
    st.plotly_chart(fig_timeline, use_container_width=True)

    # Distributie pe categorii
    df_cat = df_zile.groupby("categorie").size().reset_index(name="zile")
    df_cat = df_cat.sort_values("zile", ascending=False)

    col_pie, col_list = st.columns([1, 1])
    with col_pie:
        fig_pie = px.pie(
            df_cat, values="zile", names="categorie",
            color="categorie",
            color_discrete_map=CULORI_CATEGORIE,
            title="Distributia pe categorii"
        )
        fig_pie.update_layout(height=360)
        st.plotly_chart(fig_pie, use_container_width=True)

    with col_list:
        st.markdown("#### Toate zilele")
        for z in ZILE_BLOC3:
            culoare = CULORI_CATEGORIE.get(z["categorie"], "#95A5A6")
            st.markdown(
                f'<div style="padding:3px 8px;margin:2px 0;'
                f'border-left:3px solid {culoare};font-size:13px;">'
                f'<b>Ziua {z["ziua"]}</b> — {z["titlu"]} '
                f'<span style="color:{culoare};font-size:11px;">'
                f'[{z["categorie"]}]</span>'
                f'</div>',
                unsafe_allow_html=True
            )

# ── TAB 2: HARTA COMPETENTE ───────────────────────────────────────────────────
with tab2:
    st.subheader("Harta Competentelor Dobandite")

    # Radar chart competente
    categorii_radar = list(df_cat.set_index("categorie")["zile"].to_dict().keys())
    valori_radar    = list(df_cat.set_index("categorie")["zile"].to_dict().values())

    fig_radar = go.Figure(go.Scatterpolar(
        r=valori_radar + [valori_radar[0]],
        theta=categorii_radar + [categorii_radar[0]],
        fill="toself",
        fillcolor="rgba(0, 82, 165, 0.2)",
        line=dict(color="#0052A5", width=2),
        name="Zile per categorie"
    ))
    fig_radar.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 6])),
        title="Profil de competente Bloc 3",
        height=450
    )
    st.plotly_chart(fig_radar, use_container_width=True)

    st.markdown("#### Cele 10 competente dobandite")
    for i, comp in enumerate(COMPETENTE, 1):
        st.markdown(
            f'<div style="padding:8px 12px;margin:4px 0;'
            f'background:#EBF5FB;border-left:4px solid #0052A5;'
            f'border-radius:4px;font-size:13px;">'
            f'<b>{i}.</b> {comp}</div>',
            unsafe_allow_html=True
        )

# ── TAB 3: DEMO LIVE ──────────────────────────────────────────────────────────
with tab3:
    st.subheader("Demo Live — Sistemul AGROVISION in Actiune")

    st.markdown(
        "Aplicatia AGROVISION este **live pe internet**, accesibila oricand, gratuit."
    )

    col_url, col_qr = st.columns([2, 1])
    with col_url:
        st.markdown(
            """
            <div style="background:#EBF5FB;border:2px solid #0052A5;
            border-radius:10px;padding:20px;text-align:center;">
            <div style="font-size:14px;color:#666;margin-bottom:8px;">
            URL Public AGROVISION</div>
            <div style="font-size:15px;font-weight:bold;color:#0052A5;
            word-break:break-all;">
            https://agrovision-bloc3-8qydbmd2z3zgmpqk4ygtsg.streamlit.app
            </div>
            <div style="font-size:12px;color:#27AE60;margin-top:8px;">
            Live 24/7 | Gratuit | GitHub + Streamlit Cloud + HuggingFace
            </div>
            </div>
            """,
            unsafe_allow_html=True
        )

        st.markdown("#### Ce functioneaza live:")
        functionalitati = [
            ("Dashboard AGROVISION", "Harta Folium, KPI, detectie live"),
            ("Autentificare roluri", "admin / inspector / viewer"),
            ("Rapoarte PDF", "Generat din browser, descarcat instant"),
            ("Baza de date SQLite", "Sesiuni si detectii salvate"),
            ("API FastAPI", "Pornit separat pe port 8000/8001"),
            ("Hugging Face Hub", "Model descarcat automat din cloud"),
            ("Indici NDVI spectral", "ExG, VARI, GLI, NGRDI din RGB"),
            ("K-Means clustering", "3 clustere pe 10 parcele LPIS Gorj"),
            ("Random Forest", "Clasificare risc + regresie productie"),
        ]
        for titlu, desc in functionalitati:
            st.markdown(f"**{titlu}** — {desc}")

    with col_qr:
        st.markdown(
            """
            <div style="background:#F8F9FA;border:1px solid #ddd;
            border-radius:10px;padding:20px;text-align:center;">
            <div style="font-size:40px;">🌐</div>
            <div style="font-size:12px;color:#666;margin-top:8px;">
            Deschide in browser sau<br>trimite adresa colegilor
            </div>
            </div>
            """,
            unsafe_allow_html=True
        )

        st.markdown("#### Stack tehnologic:")
        st.markdown("- Python 3.x")
        st.markdown("- Streamlit")
        st.markdown("- YOLOv8 (Ultralytics)")
        st.markdown("- scikit-learn")
        st.markdown("- Folium + Plotly")
        st.markdown("- FastAPI + SQLite")
        st.markdown("- HuggingFace Hub")
        st.markdown("- GitHub (public)")

# ── TAB 4: CERTIFICAT ─────────────────────────────────────────────────────────
with tab4:
    st.subheader("Certificat de Competente Digitale — Bloc 3")

    st.markdown(
        "Descarca certificatul Word care atesta parcurgerea celor 30 de zile "
        "de Deep Learning YOLOv8 si Machine Learning agricol."
    )

    col_cert, col_prev = st.columns([1, 1])

    with col_cert:
        cert_data = genereaza_certificat()
        st.download_button(
            label="Descarca Certificat Word",
            data=cert_data,
            file_name=f"Certificat_Bloc3_YOLOv8_Gamulescu_{datetime.date.today()}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary"
        )

        st.markdown("#### Certificatul contine:")
        st.markdown("- Antet UCB Targu Jiu")
        st.markdown("- Numele complet si functia")
        st.markdown("- Cele 10 competente dobandite")
        st.markdown("- Metrici model: mAP50=0.829")
        st.markdown("- URL aplicatie live")
        st.markdown("- Data si loc semnatura")

    with col_prev:
        st.markdown("#### Previzualizare continut:")
        for i, comp in enumerate(COMPETENTE, 1):
            st.markdown(f"{i}. {comp}")

# ── TAB 5: ROADMAP ────────────────────────────────────────────────────────────
with tab5:
    st.subheader("Roadmap — Ce urmeaza dupa Bloc 3")

    etape = [
        {
            "etapa": "Bloc 3 completat",
            "perioada": "Aprilie 2026",
            "status": "FINALIZAT",
            "culoare": "#27AE60",
            "actiuni": [
                "30 module YOLOv8 + ML + GIS",
                "Aplicatie live pe Streamlit Cloud",
                "Model pe HuggingFace Hub",
                "mAP50=0.829 pe date reale Gorj"
            ]
        },
        {
            "etapa": "Articol IEEE FINE 2026",
            "perioada": "Aug 2026",
            "status": "TRIMIS",
            "culoare": "#F39C12",
            "actiuni": [
                "Paper ID: IEEE_FINE_2026_paper_28",
                "Conferinta: Osaka, Japonia",
                "Asteptam decizia reviewerilor",
                "Dupa acceptare → PCE UEFISCDI"
            ]
        },
        {
            "etapa": "Proiect PCE UEFISCDI 2026",
            "perioada": "Oct 2026",
            "status": "IN PREGATIRE",
            "culoare": "#3498DB",
            "actiuni": [
                "Conditie: publicatie ISI acceptata",
                "Parteneri: UCB + APIA + Prefectura + DAJ",
                "Buget: ~620.000 RON (3 ani)",
                "Aplicatie AGROVISION = demonstrator"
            ]
        },
        {
            "etapa": "Horizon Europe",
            "perioada": "2028-2029",
            "status": "PLANIFICAT",
            "culoare": "#9B59B6",
            "actiuni": [
                "Cluster 6: Food, Bioeconomy, Agriculture",
                "Consortiu minim 3 state membre UE",
                "Buget: 3-5 milioane EUR",
                "Strategie: PCE 2026 → ISI 2027 → Horizon 2028"
            ]
        },
    ]

    for etapa in etape:
        with st.expander(
            f"{etapa['etapa']} — {etapa['perioada']} "
            f"[{etapa['status']}]",
            expanded=(etapa["status"] in ["FINALIZAT", "TRIMIS"])
        ):
            st.markdown(
                f'<div style="border-left:4px solid {etapa["culoare"]};'
                f'padding:4px 12px;">' +
                "".join([f"<p style='margin:4px 0'>• {a}</p>"
                         for a in etapa["actiuni"]]) +
                "</div>",
                unsafe_allow_html=True
            )

    st.divider()
    st.markdown(
        """
        ### Mesaj final

        In 30 de zile ai construit un sistem AI complet pentru agricultura de precizie:
        - Ai antrenat un model YOLOv8 pe imagini drone reale din judetul Gorj
        - Ai creat o aplicatie web profesionala cu 30 de module functionale
        - Ai publicat modelul pe cloud si aplicatia pe internet
        - Ai trimis un articol la o conferinta IEEE internationala

        **Combinatia unica in Romania:**
        APIA 20+ ani + Doctorat drone/AI + Profesor UCB =
        avantaj competitiv real pentru proiecte europene finantate.
        """
    )

# ─── FOOTER ───────────────────────────────────────────────────────────────────
st.divider()
st.markdown(
    """
    <div style='text-align:center;padding:16px;
    background:linear-gradient(135deg,#0052A5,#1ABC9C);
    border-radius:10px;color:white;'>
    <b>AGROVISION — Bloc 3 Complet</b><br>
    30 zile | YOLOv8 | scikit-learn | Streamlit | FastAPI | GIS |
    Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj | 2026
    </div>
    """,
    unsafe_allow_html=True
)
