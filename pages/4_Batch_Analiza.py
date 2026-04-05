"""
BLOC 3 — Deep Learning YOLOv8, Ziua 4
Analiza batch imagini drone — procesare multipla + raport Excel + Word
Autor: Prof. Asoc. Dr. Oliviu Mihnea Gamulescu | UCB Targu Jiu | APIA CJ Gorj

CONCEPT CHEIE:
  st.file_uploader(accept_multiple_files=True) — upload mai multe fisiere odata
  for fisier in fisiere: ...                   — procesare in bucla
  st.progress(i/n)                             — bara de progres
  pd.ExcelWriter cu mai multe sheet-uri        — raport Excel complet
  python-docx cu tabel colorat per imagine     — raport Word oficial
"""

import streamlit as st
import numpy as np
import pandas as pd
from PIL import Image
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from io import BytesIO
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ─── Configurare pagina ───────────────────────────────────────────────────────
st.set_page_config(page_title="Batch Analiza Drone — Ziua 4", layout="wide")

st.markdown("""
<style>
.titlu { color:#6a1b9a; font-size:1.05rem; font-weight:700;
         border-bottom:2px solid #6a1b9a; padding-bottom:4px; margin-bottom:12px; }
.card         { background:#f3e5f5; border-left:4px solid #6a1b9a;
                border-radius:6px; padding:12px 16px; margin-bottom:8px; }
.card-conform { background:#e8f5e9; border-left:4px solid #2e7d32;
                border-radius:6px; padding:10px 14px; margin-bottom:6px; }
.card-mediu   { background:#fff8e1; border-left:4px solid #f9a825;
                border-radius:6px; padding:10px 14px; margin-bottom:6px; }
.card-ridicat { background:#ffebee; border-left:4px solid #c62828;
                border-radius:6px; padding:10px 14px; margin-bottom:6px; }
</style>
""", unsafe_allow_html=True)

st.title("Ziua 4 — Analiza Batch Imagini Drone")
st.markdown("**Procesare multipla: mai multe imagini odata → raport Excel + Word automat**")
st.markdown("---")

# ─── SIDEBAR ──────────────────────────────────────────────────────────────────
st.sidebar.markdown("### Parametri analiza")
prag_veg   = st.sidebar.slider("Prag vegetatie ExG", 0.0, 0.4, 0.10, 0.01)
prag_risc  = st.sidebar.slider("Prag risc PAC (%)",  5,   40,  20,   1,
                                help="Parcele sub acest % de vegetatie = RISC RIDICAT")
inspector  = st.sidebar.text_input("Inspector", "Consilier Superior Oliviu Gamulescu")
nr_raport  = st.sidebar.text_input("Nr. raport", f"APIA-GJ-{date.today().strftime('%Y%m%d')}-B")

# ─── UPLOAD MULTIPLE ──────────────────────────────────────────────────────────
st.markdown('<p class="titlu">Pas 1 — Incarca imaginile drone</p>',
            unsafe_allow_html=True)

fisiere = st.file_uploader(
    "Selecteaza una sau mai multe imagini (JPG, PNG, TIF)",
    type=["jpg","jpeg","png","tif","tiff"],
    accept_multiple_files=True
)

if not fisiere:
    st.markdown("""
    <div class="card">
        <strong>Cum functioneaza batch processing:</strong><br>
        1. Selectezi mai multe imagini odata (Ctrl+Click sau Shift+Click)<br>
        2. Aplicatia le proceseaza pe rand, automat<br>
        3. La final primesti un raport Excel + Word cu toate rezultatele<br><br>
        <strong>Util pentru:</strong> inspectorii APIA care verifica zeci de parcele zilnic.
    </div>
    """, unsafe_allow_html=True)
    st.stop()

st.markdown(f"**{len(fisiere)} imagini incarcate.** Se proceseaza...")

# ─── PROCESARE BATCH ──────────────────────────────────────────────────────────

def analizeaza_imagine(img_pil, prag):
    img_np = np.array(img_pil.convert("RGB"), dtype=np.float32) / 255.0
    R, G, B = img_np[:,:,0], img_np[:,:,1], img_np[:,:,2]
    eps = 1e-6

    ExG  = 2*G - R - B
    VARI = (G - R) / (G + R - B + eps)
    GLI  = (2*G - R - B) / (2*G + R + B + eps)

    masca = (ExG > prag) & (VARI > prag)
    pct   = float(masca.mean() * 100)

    return {
        "exg_medie":  round(float(ExG.mean()), 4),
        "vari_medie": round(float(VARI.mean()), 4),
        "gli_medie":  round(float(GLI.mean()), 4),
        "pct_vegetatie": round(pct, 2),
        "rezolutie": f"{img_pil.width}x{img_pil.height}",
        "masca": masca,
        "exg": ExG,
    }

rezultate = []
progres   = st.progress(0)
status    = st.empty()

for i, fisier in enumerate(fisiere):
    status.markdown(f"Se proceseaza: **{fisier.name}** ({i+1}/{len(fisiere)})")
    img_pil = Image.open(fisier)
    rez     = analizeaza_imagine(img_pil, prag_veg)
    pct     = rez["pct_vegetatie"]

    if pct >= 60:
        risc = "CONFORM"
    elif pct >= prag_risc:
        risc = "MEDIU"
    else:
        risc = "RIDICAT"

    rezultate.append({
        "Nr.":           i + 1,
        "Fisier":        fisier.name,
        "Rezolutie":     rez["rezolutie"],
        "ExG medie":     rez["exg_medie"],
        "VARI medie":    rez["vari_medie"],
        "GLI medie":     rez["gli_medie"],
        "Vegetatie (%)": pct,
        "Risc PAC":      risc,
        "_masca":        rez["masca"],
        "_exg":          rez["exg"],
        "_pil":          img_pil,
    })
    progres.progress((i + 1) / len(fisiere))

status.empty()
progres.empty()

df = pd.DataFrame([{k: v for k, v in r.items() if not k.startswith("_")}
                   for r in rezultate])

# ─── METRICI GLOBALE ──────────────────────────────────────────────────────────
st.markdown("---")
st.markdown('<p class="titlu">Rezultate batch</p>', unsafe_allow_html=True)

n_total   = len(df)
n_conform = (df["Risc PAC"] == "CONFORM").sum()
n_mediu   = (df["Risc PAC"] == "MEDIU").sum()
n_ridicat = (df["Risc PAC"] == "RIDICAT").sum()

c1, c2, c3, c4 = st.columns(4)
c1.metric("Total imagini",   n_total)
c2.metric("CONFORM",         n_conform, delta=f"{n_conform/n_total:.0%}")
c3.metric("Risc MEDIU",      n_mediu,   delta=f"{n_mediu/n_total:.0%}")
c4.metric("Risc RIDICAT",    n_ridicat, delta=f"{n_ridicat/n_total:.0%}")

# Tabel rezultate cu culori
def coloreaza_risc(val):
    if val == "RIDICAT": return "background-color:#ffcdd2"
    if val == "MEDIU":   return "background-color:#fff9c4"
    return "background-color:#c8e6c9"

st.dataframe(
    df.style.applymap(coloreaza_risc, subset=["Risc PAC"]),
    use_container_width=True, hide_index=True
)

st.markdown("---")

# ─── GRAFICE BATCH ────────────────────────────────────────────────────────────
st.markdown('<p class="titlu">Grafice comparative</p>', unsafe_allow_html=True)

col_g1, col_g2 = st.columns(2)

with col_g1:
    # Bar vegetatie per imagine
    fig1, ax1 = plt.subplots(figsize=(5, 3.5))
    culori = ["#c62828" if r=="RIDICAT" else "#f9a825" if r=="MEDIU" else "#2e7d32"
              for r in df["Risc PAC"]]
    bars = ax1.bar(range(n_total), df["Vegetatie (%)"], color=culori, edgecolor="white")
    ax1.axhline(prag_risc, color="#c62828", linestyle="--",
                linewidth=1.5, label=f"Prag risc {prag_risc}%")
    ax1.axhline(60, color="#2e7d32", linestyle="--",
                linewidth=1.5, label="Prag conform 60%")
    ax1.set_xticks(range(n_total))
    ax1.set_xticklabels([r["Fisier"][:12] for r in rezultate],
                        rotation=45, ha="right", fontsize=7)
    ax1.set_ylabel("Vegetatie (%)")
    ax1.set_title("Vegetatie per parcela")
    ax1.legend(fontsize=8)
    ax1.spines[["top","right"]].set_visible(False)
    plt.tight_layout()
    buf1 = BytesIO(); fig1.savefig(buf1, dpi=150, bbox_inches="tight")
    buf1.seek(0); plt.close()
    st.image(buf1, use_container_width=True)

with col_g2:
    # Pie risc
    dist = df["Risc PAC"].value_counts()
    fig2, ax2 = plt.subplots(figsize=(5, 3.5))
    culori_pie = {"CONFORM":"#2e7d32","MEDIU":"#f9a825","RIDICAT":"#c62828"}
    ax2.pie(dist.values, labels=dist.index, autopct="%1.0f%%",
            colors=[culori_pie.get(k,"#888") for k in dist.index],
            startangle=90)
    ax2.set_title(f"Distributie risc — {n_total} parcele")
    plt.tight_layout()
    buf2 = BytesIO(); fig2.savefig(buf2, dpi=150, bbox_inches="tight")
    buf2.seek(0); plt.close()
    st.image(buf2, use_container_width=True)

st.markdown("---")

# ─── PREVIZUALIZARE HARTI ─────────────────────────────────────────────────────
if n_total <= 6:
    st.markdown('<p class="titlu">Harti vegetatie per parcela</p>',
                unsafe_allow_html=True)

    cols = st.columns(min(n_total, 3))
    for i, (rez, col) in enumerate(zip(rezultate, cols * 2)):
        with col:
            masca  = rez["_masca"]
            img_np = np.array(rez["_pil"].convert("RGB"))
            overlay = img_np.copy()
            overlay[masca]  = [0, 200, 0]
            overlay[~masca] = (overlay[~masca] * 0.6 +
                               np.array([160,80,0])*0.4).clip(0,255).astype(np.uint8)
            risc = rez["Risc PAC"]
            culoare_titlu = {"CONFORM":"green","MEDIU":"orange","RIDICAT":"red"}
            st.image(overlay, use_container_width=True,
                     caption=f"{rez['Fisier'][:20]} — {rez['Vegetatie (%)']:.1f}%")
            st.markdown(f"<center><strong style='color:{culoare_titlu[risc]}'>"
                        f"{risc}</strong></center>", unsafe_allow_html=True)

st.markdown("---")

# ─── EXPORT ───────────────────────────────────────────────────────────────────
st.markdown('<p class="titlu">Export rapoarte</p>', unsafe_allow_html=True)

col_e1, col_e2 = st.columns(2)

# Export Excel
with col_e1:
    buf_xl = BytesIO()
    with pd.ExcelWriter(buf_xl, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Toate_parcelele")
        df[df["Risc PAC"]=="RIDICAT"].to_excel(writer, index=False,
                                                sheet_name="Risc_Ridicat")
        df[df["Risc PAC"]=="MEDIU"].to_excel(writer, index=False,
                                              sheet_name="Risc_Mediu")
        df[df["Risc PAC"]=="CONFORM"].to_excel(writer, index=False,
                                                sheet_name="Conform")
        # Sheet statistici
        stats = pd.DataFrame({
            "Indicator": ["Total imagini","CONFORM","Risc MEDIU","Risc RIDICAT",
                          "Vegetatie medie (%)","ExG mediu","VARI mediu"],
            "Valoare":   [n_total, n_conform, n_mediu, n_ridicat,
                          round(df["Vegetatie (%)"].mean(),2),
                          round(df["ExG medie"].mean(),4),
                          round(df["VARI medie"].mean(),4)]
        })
        stats.to_excel(writer, index=False, sheet_name="Statistici")
    buf_xl.seek(0)

    st.download_button(
        "Descarca Excel (4 sheet-uri)",
        data=buf_xl,
        file_name=f"Batch_APIA_{date.today().strftime('%Y%m%d')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# Export Word
with col_e2:
    if DOCX_OK:
        if st.button("Genereaza Raport Word", type="primary"):
            doc = Document()
            for sec in doc.sections:
                sec.top_margin = sec.bottom_margin = Cm(2.5)
                sec.left_margin = sec.right_margin = Cm(2.5)

            # Titlu
            t = doc.add_heading("Raport Analiza Batch Imagini Drone", level=1)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            t.runs[0].font.color.rgb = RGBColor(0x6a, 0x1b, 0x9a)

            p = doc.add_paragraph(f"Nr. raport: {nr_raport} | "
                                  f"Data: {date.today().strftime('%d.%m.%Y')} | "
                                  f"Inspector: {inspector}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = RGBColor(0x55,0x55,0x55)

            doc.add_paragraph()

            # Rezumat
            doc.add_heading("Rezumat executiv", level=2)
            doc.add_paragraph(
                f"Au fost analizate {n_total} imagini drone utilizand indicii spectrale "
                f"ExG si VARI pentru detectia vegetatiei. Pragul de risc PAC aplicat: "
                f"{prag_risc}% vegetatie minima. Rezultate: {n_conform} parcele CONFORME, "
                f"{n_mediu} cu RISC MEDIU, {n_ridicat} cu RISC RIDICAT."
            )

            doc.add_paragraph()

            # Tabel rezultate
            doc.add_heading("Rezultate per parcela", level=2)
            cols_tabel = ["Nr.", "Fisier", "Vegetatie (%)", "Risc PAC"]
            tabel = doc.add_table(rows=1, cols=len(cols_tabel))
            tabel.style = "Table Grid"

            # Header
            hdr = tabel.rows[0].cells
            for cell, col_name in zip(hdr, cols_tabel):
                cell.text = col_name
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(10)

            # Date
            culori_risc_rgb = {
                "CONFORM":  RGBColor(0x2e, 0x7d, 0x32),
                "MEDIU":    RGBColor(0xf9, 0xa8, 0x25),
                "RIDICAT":  RGBColor(0xc6, 0x28, 0x28),
            }
            for _, row_data in df.iterrows():
                row = tabel.add_row().cells
                vals = [str(row_data["Nr."]), row_data["Fisier"],
                        f"{row_data['Vegetatie (%)']:.1f}%", row_data["Risc PAC"]]
                for cell, val in zip(row, vals):
                    cell.text = val
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    if val in culori_risc_rgb:
                        cell.paragraphs[0].runs[0].font.color.rgb = culori_risc_rgb[val]
                        cell.paragraphs[0].runs[0].font.bold = True

            doc.add_paragraph()
            doc.add_paragraph(f"Inspector: {inspector}").runs[0].bold = True
            doc.add_paragraph(f"Data: {date.today().strftime('%d.%m.%Y')}")

            buf_w = BytesIO()
            doc.save(buf_w)
            buf_w.seek(0)

            st.download_button(
                "Descarca Raport Word",
                data=buf_w,
                file_name=f"Raport_Batch_{date.today().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.warning("Instaleaza python-docx: pip install python-docx")

st.markdown("---")

# ─── Concept Ziua 4 ───────────────────────────────────────────────────────────
with st.expander("Conceptul Zilei 4 — Batch processing cu st.file_uploader multiplu"):
    st.markdown("""
**Concept cheie: `accept_multiple_files=True`** — permite upload mai multor fisiere odata.
""")
    st.code("""
# Upload multiple fisiere
fisiere = st.file_uploader(
    "Selecteaza imagini",
    type=["jpg","png","tif"],
    accept_multiple_files=True   # <-- cheia
)

# Bara de progres
progres = st.progress(0)

for i, fisier in enumerate(fisiere):
    img = Image.open(fisier)
    # ... procesare ...
    progres.progress((i + 1) / len(fisiere))

progres.empty()  # ascunde bara dupa finalizare
    """, language="python")
    st.info("**st.progress(valoare)** — valoarea trebuie sa fie intre 0.0 si 1.0. "
            "Folosim (i+1)/total pentru a actualiza progresul la fiecare imagine.")
