"""
AGROVISION -- Multi-Agenti Avansat
Ziua 38 | Autor: Prof. Asoc. Dr. Oliviu Mihnea Gamulescu

Diferente fata de Ziua 37:
  - PARALELISM: AgentDrona si AgentLPIS ruleaza simultan (threading)
  - ORCHESTRATOR DINAMIC: traseul depinde de rezultate (risc ridicat -> AgentAlerta)
  - MEMORIE AGENT: AgentLPIS salveaza rezultate in SQLite intre sesiuni
  - RETRY: daca un agent esueaza, incearca de max 3 ori
  - AgentAlerta (NOU): activat doar cand riscul e ridicat

CONCEPT CHEIE -- orchestrator dinamic:
    Nu apeleaza agentii intotdeauna in aceeasi ordine.
    Citeste rezultatele si decide: daca risc_pac == "RISC RIDICAT"
    --> activeaza AgentAlerta inainte de raportare.
    Altfel --> merge direct la AgentRaportare.

    Traseu normal:    Drona + LPIS (paralel) --> Conformitate --> Raportare
    Traseu alertat:   Drona + LPIS (paralel) --> Conformitate --> Alerta --> Raportare

CONCEPT CHEIE -- paralelism cu threading:
    import threading
    t1 = threading.Thread(target=agent1.proceseaza, args=(input,), kwargs={"rezultat": r})
    t2 = threading.Thread(target=agent2.proceseaza, args=(input,), kwargs={"rezultat": r})
    t1.start(); t2.start()
    t1.join(); t2.join()   # asteapta sa termine ambii

CONCEPT CHEIE -- memorie SQLite in agent:
    Fiecare agent poate deschide o conexiune SQLite si salva/citi date.
    La urmatoarea rulare, agentul stie ce a facut inainte.
"""

import streamlit as st
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from PIL import Image
import io
import time
import random
import threading
import sqlite3
from datetime import date, datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── CONFIGURARE ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Multi-Agenti Avansat | AGROVISION",
    page_icon="38",
    layout="wide"
)

# ── DATE LPIS GORJ ─────────────────────────────────────────────────────────────
PARCELE_LPIS = {
    "GJ_78258": {"fermier": "Ionescu Gheorghe", "uat": "Targu Jiu",    "ha_declarate": 4.2,  "cultura": "grau"},
    "GJ_79157": {"fermier": "Popescu Maria",    "uat": "Rovinari",     "ha_declarate": 2.8,  "cultura": "porumb"},
    "GJ_80341": {"fermier": "Constantin Vasile","uat": "Motru",        "ha_declarate": 6.1,  "cultura": "floarea-soarelui"},
    "GJ_81092": {"fermier": "Olteanu Florin",   "uat": "Novaci",       "ha_declarate": 3.5,  "cultura": "lucerna"},
    "GJ_82174": {"fermier": "Draghici Ion",     "uat": "Targu Jiu",    "ha_declarate": 5.0,  "cultura": "rapita"},
    "GJ_83015": {"fermier": "Stanescu Elena",   "uat": "Bumbesti-Jiu", "ha_declarate": 1.9,  "cultura": "pasune"},
}

DB_PATH = "agrovision_agenti.db"

# ── BAZA DE DATE MEMORIE ───────────────────────────────────────────────────────
def init_db():
    """Creeaza tabela pentru memoria agentilor daca nu exista."""
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpis_cache (
            cod_parcela TEXT PRIMARY KEY,
            fermier     TEXT,
            uat         TEXT,
            ha_declarate REAL,
            cultura     TEXT,
            ultima_verificare TEXT
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS log_agenti (
            id        INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT,
            agent     TEXT,
            cod_parcela TEXT,
            actiune   TEXT,
            rezultat  TEXT
        )
    """)
    conn.commit()
    conn.close()

init_db()

# ─────────────────────────────────────────────────────────────────────────────
# AGENTI
# ─────────────────────────────────────────────────────────────────────────────

class AgentDrona:
    """
    Analizeaza imaginea drone si calculeaza indici de vegetatie.
    Ruleaza in paralel cu AgentLPIS.
    """
    def __init__(self):
        self.nume = "AgentDrona"

    def proceseaza(self, imagine_arr, rezultat: dict):
        """
        imagine_arr: numpy array (H, W, 3) cu valorile RGB
        rezultat:    dict comun in care scrie rezultatele (thread-safe pt. chei diferite)
        """
        time.sleep(0.3)  # simuleaza procesare imagine

        r = imagine_arr[:, :, 0].astype(float)
        g = imagine_arr[:, :, 1].astype(float)
        b = imagine_arr[:, :, 2].astype(float)

        eps = 1e-6
        exg  = (2 * g - r - b)
        vari = (g - r) / (g - r + b + eps)

        masca_verde = (exg > 20) & (g > r) & (g > b)
        pct_verde   = float(masca_verde.mean() * 100)

        exg_mean  = float(np.clip(exg[masca_verde].mean() / 255, 0, 1)) if masca_verde.any() else 0.0
        vari_mean = float(np.clip(vari[masca_verde].mean(), 0, 1))       if masca_verde.any() else 0.0

        rezultat["drona"] = {
            "pct_verde": round(pct_verde, 2),
            "exg_mean":  round(exg_mean, 3),
            "vari_mean": round(vari_mean, 3),
        }

    def log(self, conn, cod, pct_verde):
        conn.execute(
            "INSERT INTO log_agenti (timestamp, agent, cod_parcela, actiune, rezultat) VALUES (?,?,?,?,?)",
            (datetime.now().isoformat(), self.nume, cod, "analiza_imagine", f"pct_verde={pct_verde}")
        )


class AgentLPIS:
    """
    Verifica parcela in baza LPIS.
    Are MEMORIE SQLite -- daca a mai verificat parcela, returneaza din cache.
    Ruleaza in paralel cu AgentDrona.
    """
    def __init__(self):
        self.nume = "AgentLPIS"

    def proceseaza(self, cod_parcela: str, rezultat: dict):
        """Cauta parcela in cache SQLite. Daca nu e, citeste din LPIS si salveaza."""
        conn = sqlite3.connect(DB_PATH)
        row = conn.execute(
            "SELECT * FROM lpis_cache WHERE cod_parcela = ?", (cod_parcela,)
        ).fetchone()

        if row:
            din_cache = True
            date_lpis = {
                "fermier":     row[1],
                "uat":         row[2],
                "ha_declarate": row[3],
                "cultura":     row[4],
            }
        else:
            din_cache = False
            time.sleep(0.2)  # simuleaza acces baza LPIS
            date_lpis = PARCELE_LPIS.get(cod_parcela, {
                "fermier": "Necunoscut", "uat": "N/A",
                "ha_declarate": 0.0, "cultura": "N/A"
            })
            conn.execute(
                """INSERT OR REPLACE INTO lpis_cache
                   (cod_parcela, fermier, uat, ha_declarate, cultura, ultima_verificare)
                   VALUES (?,?,?,?,?,?)""",
                (cod_parcela, date_lpis["fermier"], date_lpis["uat"],
                 date_lpis["ha_declarate"], date_lpis["cultura"],
                 datetime.now().isoformat())
            )
            conn.commit()

        conn.execute(
            "INSERT INTO log_agenti (timestamp, agent, cod_parcela, actiune, rezultat) VALUES (?,?,?,?,?)",
            (datetime.now().isoformat(), self.nume, cod_parcela,
             "cache_hit" if din_cache else "lpis_query",
             date_lpis["cultura"])
        )
        conn.commit()
        conn.close()

        rezultat["lpis"] = {**date_lpis, "din_cache": din_cache}


class AgentConformitate:
    """
    Compara rezultatele AgentDrona cu datele AgentLPIS.
    Emite verdict PAC conform Reg. UE 2021/2116.
    Include RETRY: daca datele lipsesc, incearca de max 3 ori.
    """
    PRAG_VERDE = 20.0   # % vegetatie minima PAC

    def __init__(self):
        self.nume = "AgentConformitate"

    def proceseaza(self, rezultat: dict, max_retry: int = 3) -> dict:
        """Returneaza verdictul si nivelul de risc."""
        for tentativa in range(1, max_retry + 1):
            if "drona" in rezultat and "lpis" in rezultat:
                break
            time.sleep(0.1 * tentativa)
        else:
            return {"verdict": "DATE INSUFICIENTE", "risc_pac": "NECUNOSCUT",
                    "detalii": "Agentii Drona/LPIS nu au furnizat date.", "tentative": max_retry}

        pct_verde    = rezultat["drona"]["pct_verde"]
        ha_declarate = rezultat["lpis"]["ha_declarate"]
        cultura      = rezultat["lpis"]["cultura"]

        # Calcul ha vegetate (estimare)
        ha_vegetate = round(ha_declarate * pct_verde / 100, 2)
        diferenta   = round(ha_vegetate - ha_declarate, 2)

        if pct_verde >= self.PRAG_VERDE:
            verdict  = "CONFORM"
            risc_pac = "RISC SCAZUT"
            detalii  = f"Vegetatie {pct_verde:.1f}% peste pragul PAC de {self.PRAG_VERDE}%."
        elif pct_verde >= 10:
            verdict  = "ATENTIONARE"
            risc_pac = "RISC MEDIU"
            detalii  = f"Vegetatie {pct_verde:.1f}% sub prag. Necesita control suplimentar."
        else:
            verdict  = "NECONFORM"
            risc_pac = "RISC RIDICAT"
            detalii  = f"Vegetatie {pct_verde:.1f}% -- posibila frauda/abandon. Suspendare plata recomandata."

        conn = sqlite3.connect(DB_PATH)
        conn.execute(
            "INSERT INTO log_agenti (timestamp, agent, cod_parcela, actiune, rezultat) VALUES (?,?,?,?,?)",
            (datetime.now().isoformat(), self.nume, rezultat["lpis"].get("uat", ""),
             "evaluare_conformitate", f"{verdict}|{risc_pac}")
        )
        conn.commit()
        conn.close()

        return {
            "verdict":     verdict,
            "risc_pac":    risc_pac,
            "detalii":     detalii,
            "ha_vegetate": ha_vegetate,
            "ha_declarate": ha_declarate,
            "diferenta":   diferenta,
            "cultura":     cultura,
            "tentative":   tentativa,
        }


class AgentAlerta:
    """
    NOU in Ziua 38.
    Activat de orchestrator DOAR cand risc_pac == 'RISC RIDICAT'.
    Genereaza o notificare de urgenta si o inregistreaza in log.
    """
    def __init__(self):
        self.nume = "AgentAlerta"

    def proceseaza(self, cod_parcela: str, conformitate: dict, fermier: str) -> dict:
        time.sleep(0.1)
        mesaj = (
            f"ALERTA PAC -- Parcela {cod_parcela} | Fermier: {fermier} | "
            f"Vegetatie insuficienta | Verdict: {conformitate['verdict']} | "
            f"Recomandat: suspendare plata + control teren."
        )
        conn = sqlite3.connect(DB_PATH)
        conn.execute(
            "INSERT INTO log_agenti (timestamp, agent, cod_parcela, actiune, rezultat) VALUES (?,?,?,?,?)",
            (datetime.now().isoformat(), self.nume, cod_parcela, "ALERTA_RIDICATA", mesaj)
        )
        conn.commit()
        conn.close()
        return {"alerta": True, "mesaj_alerta": mesaj}


class AgentRaportare:
    """Genereaza raportul Word final cu toate rezultatele."""
    def __init__(self):
        self.nume = "AgentRaportare"

    def proceseaza(self, cod_parcela: str, rezultat: dict,
                   conformitate: dict, alerta: dict | None) -> bytes:
        doc = Document()

        # Titlu
        h = doc.add_heading("RAPORT CONTROL AGROVISION", 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            f"Parcela: {cod_parcela}  |  Data: {date.today().strftime('%d.%m.%Y')}  |  "
            f"Fermier: {rezultat['lpis']['fermier']}  |  UAT: {rezultat['lpis']['uat']}"
        )
        doc.add_paragraph("")

        # Verdict
        doc.add_heading("1. Verdict PAC", level=1)
        p = doc.add_paragraph()
        r = p.add_run(f"Verdict: {conformitate['verdict']}  |  Risc: {conformitate['risc_pac']}")
        r.bold = True
        if conformitate["risc_pac"] == "RISC RIDICAT":
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif conformitate["risc_pac"] == "RISC MEDIU":
            r.font.color.rgb = RGBColor(0xD3, 0x89, 0x00)
        else:
            r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

        doc.add_paragraph(conformitate["detalii"])

        # Date LPIS
        doc.add_heading("2. Date LPIS", level=1)
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        t.rows[0].cells[0].text = "Camp"
        t.rows[0].cells[1].text = "Valoare"
        for camp, val in [
            ("Fermier",       rezultat["lpis"]["fermier"]),
            ("UAT",           rezultat["lpis"]["uat"]),
            ("Cultura",       conformitate["cultura"]),
            ("Ha declarate",  str(conformitate["ha_declarate"])),
            ("Ha vegetate",   str(conformitate["ha_vegetate"])),
            ("Diferenta",     str(conformitate["diferenta"])),
            ("Cache LPIS",    "DA" if rezultat["lpis"]["din_cache"] else "NU (interogare noua)"),
        ]:
            row = t.add_row()
            row.cells[0].text = camp
            row.cells[1].text = val

        # Date drona
        doc.add_heading("3. Analiza Drone", level=1)
        doc.add_paragraph(
            f"Vegetatie detectata: {rezultat['drona']['pct_verde']:.1f}%  |  "
            f"ExG: {rezultat['drona']['exg_mean']:.3f}  |  "
            f"VARI: {rezultat['drona']['vari_mean']:.3f}"
        )

        # Alerta
        if alerta and alerta.get("alerta"):
            doc.add_heading("4. ALERTA EMISA", level=1)
            p = doc.add_paragraph(alerta["mesaj_alerta"])
            p.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            p.runs[0].bold = True

        # Arhitectura agenti
        doc.add_heading("5. Traseu Agenti", level=1)
        traseu = "AgentDrona + AgentLPIS (paralel) -> AgentConformitate"
        if alerta and alerta.get("alerta"):
            traseu += " -> AgentAlerta -> AgentRaportare"
        else:
            traseu += " -> AgentRaportare"
        doc.add_paragraph(traseu)
        doc.add_paragraph(f"Tentative conformitate: {conformitate.get('tentative', 1)}")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# ORCHESTRATOR DINAMIC
# ─────────────────────────────────────────────────────────────────────────────

class OrchestratorAvansat:
    """
    Orchestrator cu traseu dinamic:
      - Lanseaza AgentDrona + AgentLPIS in PARALEL (threading)
      - Asteapta ambii agenti (join)
      - Apeleaza AgentConformitate cu retry
      - DACA risc ridicat --> activeaza AgentAlerta (traseu alertat)
      - Apeleaza AgentRaportare cu toate rezultatele
    """
    def __init__(self):
        self.drona       = AgentDrona()
        self.lpis        = AgentLPIS()
        self.conformitate = AgentConformitate()
        self.alerta      = AgentAlerta()
        self.raportare   = AgentRaportare()

    def ruleaza(self, imagine_arr, cod_parcela: str,
                status_container) -> tuple[dict, bytes]:
        rezultat = {}
        log_traseu = []

        # PASUL 1 -- Paralel: AgentDrona + AgentLPIS
        status_container.update(
            label="Pas 1/4 -- AgentDrona + AgentLPIS ruleaza in paralel...",
            state="running"
        )

        t_drona = threading.Thread(
            target=self.drona.proceseaza,
            args=(imagine_arr, rezultat)
        )
        t_lpis = threading.Thread(
            target=self.lpis.proceseaza,
            args=(cod_parcela, rezultat)
        )
        t_drona.start()
        t_lpis.start()
        t_drona.join()
        t_lpis.join()

        log_traseu.append("AgentDrona + AgentLPIS (paralel) -- COMPLET")

        # PASUL 2 -- AgentConformitate (cu retry)
        status_container.update(
            label="Pas 2/4 -- AgentConformitate evalueaza...",
            state="running"
        )
        conformitate = self.conformitate.proceseaza(rezultat)
        log_traseu.append(f"AgentConformitate -- {conformitate['verdict']}")

        # PASUL 3 -- DINAMIC: alerta doar daca risc ridicat
        alerta_rezultat = None
        if conformitate["risc_pac"] == "RISC RIDICAT":
            status_container.update(
                label="Pas 3/4 -- RISC RIDICAT! AgentAlerta activat...",
                state="running"
            )
            alerta_rezultat = self.alerta.proceseaza(
                cod_parcela, conformitate, rezultat["lpis"]["fermier"]
            )
            log_traseu.append("AgentAlerta -- ALERTA EMISA")
        else:
            log_traseu.append("AgentAlerta -- omis (risc acceptabil)")

        # PASUL 4 -- Raportare
        status_container.update(
            label="Pas 4/4 -- AgentRaportare genereaza documentul...",
            state="running"
        )
        doc_bytes = self.raportare.proceseaza(
            cod_parcela, rezultat, conformitate, alerta_rezultat
        )
        log_traseu.append("AgentRaportare -- RAPORT GENERAT")

        status_container.update(label="Procesare completa!", state="complete")

        return {
            "rezultat":     rezultat,
            "conformitate": conformitate,
            "alerta":       alerta_rezultat,
            "log_traseu":   log_traseu,
        }, doc_bytes


# ─────────────────────────────────────────────────────────────────────────────
# INTERFATA STREAMLIT
# ─────────────────────────────────────────────────────────────────────────────

st.title("Multi-Agenti Avansat -- Ziua 38")
st.caption("Paralelism | Orchestrator dinamic | Memorie SQLite | Retry | AgentAlerta")

tab1, tab2, tab3, tab4 = st.tabs([
    "Procesare Parcela",
    "Log Agenti (SQLite)",
    "Cache LPIS",
    "Teorie -- Agenti Avansati",
])

# ── TAB 1: PROCESARE ──────────────────────────────────────────────────────────
with tab1:
    col1, col2 = st.columns([1, 2])

    with col1:
        st.subheader("Configurare misiune")
        cod_selectat = st.selectbox(
            "Parcela LPIS",
            list(PARCELE_LPIS.keys()),
            help="Alege parcela pe care o analizezi"
        )

        imagine_upload = st.file_uploader(
            "Imagine drone (optional)",
            type=["jpg", "jpeg", "png"],
            help="Daca nu uploadezi, se genereaza o imagine sintetica"
        )

        # Slider risc manual (pentru demonstratie)
        pct_verde_manual = st.slider(
            "Simuleaza % vegetatie (daca nu ai imagine)",
            min_value=0, max_value=100, value=45,
            help="Sub 20% = RISC RIDICAT | 20-50% = RISC MEDIU | >50% = CONFORM"
        )

        btn_ruleaza = st.button("Lanseaza orchestratorul", type="primary", use_container_width=True)

    with col2:
        if btn_ruleaza:
            # Pregateste imaginea
            if imagine_upload:
                img = Image.open(imagine_upload).convert("RGB").resize((200, 200))
                imagine_arr = np.array(img)
            else:
                # Genereaza imagine sintetica cu pct_verde_manual
                imagine_arr = np.zeros((200, 200, 3), dtype=np.uint8)
                n_verzi = int(200 * 200 * pct_verde_manual / 100)
                idx = random.sample(range(200 * 200), n_verzi)
                flat = imagine_arr.reshape(-1, 3)
                for i in idx:
                    flat[i] = [20, 120 + random.randint(0, 80), 20]  # verde
                imagine_arr = flat.reshape(200, 200, 3)

            orchestrator = OrchestratorAvansat()

            with st.status("Orchestratorul lucreaza...", expanded=True) as status:
                rezultate, doc_bytes = orchestrator.ruleaza(
                    imagine_arr, cod_selectat, status
                )

            # Afisare rezultate
            conf   = rezultate["conformitate"]
            drona  = rezultate["rezultat"]["drona"]
            lpis   = rezultate["rezultat"]["lpis"]
            alerta = rezultate["alerta"]

            # Verdict principal
            culoare_verdict = {
                "CONFORM":     "green",
                "ATENTIONARE": "orange",
                "NECONFORM":   "red",
            }.get(conf["verdict"], "gray")

            st.markdown(
                f"<h2 style='color:{culoare_verdict};text-align:center'>"
                f"{conf['verdict']} -- {conf['risc_pac']}</h2>",
                unsafe_allow_html=True
            )
            st.info(conf["detalii"])

            if alerta:
                st.error(f"ALERTA EMISA: {alerta['mesaj_alerta']}")

            # KPI-uri
            k1, k2, k3, k4 = st.columns(4)
            k1.metric("Vegetatie", f"{drona['pct_verde']:.1f}%")
            k2.metric("ExG", f"{drona['exg_mean']:.3f}")
            k3.metric("Ha vegetate", f"{conf['ha_vegetate']:.2f}")
            k4.metric("Diferenta ha", f"{conf['diferenta']:+.2f}")

            st.caption(f"Cache LPIS: {'DA (din memorie)' if lpis['din_cache'] else 'NU (interogat acum)'}")

            # Traseu agenti
            st.subheader("Traseul orchestratorului")
            for pas in rezultate["log_traseu"]:
                st.markdown(f"- {pas}")

            # Imagine sintetica
            fig, ax = plt.subplots(figsize=(3, 3))
            ax.imshow(imagine_arr)
            ax.set_title(f"Imagine {cod_selectat}", fontsize=9)
            ax.axis("off")
            st.pyplot(fig)
            plt.close(fig)

            # Download raport
            st.download_button(
                "Descarca Raport Word",
                data=doc_bytes,
                file_name=f"Raport_Agenti_{cod_selectat}_{date.today()}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

# ── TAB 2: LOG AGENTI ──────────────────────────────────────────────────────────
with tab2:
    st.subheader("Log actiuni agenti (SQLite)")
    st.caption("Fiecare agent inregistreaza ce a facut. Persistenta intre sesiuni.")

    conn = sqlite3.connect(DB_PATH)
    df_log = pd.read_sql(
        "SELECT timestamp, agent, cod_parcela, actiune, rezultat FROM log_agenti ORDER BY id DESC LIMIT 100",
        conn
    )
    conn.close()

    if df_log.empty:
        st.info("Nu exista inregistrari inca. Ruleaza orchestratorul in Tab 1.")
    else:
        # Filtre
        agenti_unici = ["Toti"] + sorted(df_log["agent"].unique().tolist())
        filtru_agent = st.selectbox("Filtreaza dupa agent", agenti_unici)
        if filtru_agent != "Toti":
            df_log = df_log[df_log["agent"] == filtru_agent]

        st.dataframe(df_log, use_container_width=True)

        # Statistici
        st.subheader("Actiuni per agent")
        fig2, ax2 = plt.subplots(figsize=(6, 3))
        conn2 = sqlite3.connect(DB_PATH)
        df_stat = pd.read_sql(
            "SELECT agent, COUNT(*) as n FROM log_agenti GROUP BY agent", conn2
        )
        conn2.close()
        if not df_stat.empty:
            ax2.barh(df_stat["agent"], df_stat["n"], color="#27AE60")
            ax2.set_xlabel("Nr. actiuni")
            ax2.set_title("Activitate per agent")
            st.pyplot(fig2)
        plt.close(fig2)

        if st.button("Sterge tot log-ul"):
            conn3 = sqlite3.connect(DB_PATH)
            conn3.execute("DELETE FROM log_agenti")
            conn3.commit()
            conn3.close()
            st.success("Log sters.")
            st.rerun()

# ── TAB 3: CACHE LPIS ─────────────────────────────────────────────────────────
with tab3:
    st.subheader("Cache LPIS (memoria AgentLPIS)")
    st.caption(
        "Prima data cand AgentLPIS verifica o parcela, salveaza in SQLite. "
        "La urmatoarea rulare, returneaza din cache -- fara sa mai interoge LPIS."
    )

    conn = sqlite3.connect(DB_PATH)
    df_cache = pd.read_sql(
        "SELECT cod_parcela, fermier, uat, ha_declarate, cultura, ultima_verificare FROM lpis_cache",
        conn
    )
    conn.close()

    if df_cache.empty:
        st.info("Cache gol. Ruleaza orchestratorul de cateva ori pentru a-l popula.")
    else:
        st.dataframe(df_cache, use_container_width=True)
        st.success(f"{len(df_cache)} parcele in cache. Urmatoarele rulari vor fi mai rapide.")

        if st.button("Goleste cache LPIS"):
            conn4 = sqlite3.connect(DB_PATH)
            conn4.execute("DELETE FROM lpis_cache")
            conn4.commit()
            conn4.close()
            st.success("Cache golit.")
            st.rerun()

# ── TAB 4: TEORIE ─────────────────────────────────────────────────────────────
with tab4:
    st.subheader("Arhitectura Multi-Agent Avansata")

    st.markdown("""
### De ce arhitectura multi-agent?

Un singur script face totul secvential -- daca un pas dureaza, totul asteapta.
Arhitectura multi-agent imparte munca:
- **Paralelism** -- doi agenti lucreaza simultan, economisind timp
- **Specializare** -- fiecare agent stie sa faca un singur lucru, bine
- **Flexibilitate** -- orchestratorul decide traseul in functie de rezultate

---

### Comparatie Ziua 37 vs Ziua 38

| Caracteristica | Ziua 37 | Ziua 38 |
|---|---|---|
| Agenti | 4, in secventa | 5, cu paralelism |
| Orchestrator | Fix (apeleaza in ordine) | Dinamic (decide traseul) |
| Memorie | Fara | SQLite persistenta |
| Retry | Fara | Max 3 tentative |
| Alerta | Fara | AgentAlerta (risc ridicat) |

---

### Paralelism cu threading

```python
import threading

rezultat = {}  # dict comun, scris de ambii agenti

t1 = threading.Thread(target=agent_drona.proceseaza, args=(imagine, rezultat))
t2 = threading.Thread(target=agent_lpis.proceseaza,  args=(cod, rezultat))

t1.start(); t2.start()   # pornesc simultan
t1.join();  t2.join()    # asteapta sa termine ambii

# dupa join(), rezultat contine datele de la amandoi agentii
```

**Atentie:** Fiecare agent scrie la chei diferite in dict (`rezultat["drona"]` si
`rezultat["lpis"]`) -- nu exista conflict intre fire de executie.

---

### Orchestrator dinamic

```python
conformitate = agent_conf.proceseaza(rezultat)

if conformitate["risc_pac"] == "RISC RIDICAT":
    # traseu alertat
    alerta = agent_alerta.proceseaza(...)
else:
    # traseu normal
    alerta = None

raport = agent_raportare.proceseaza(..., alerta)
```

Orchestratorul nu stie dinainte ce va face -- decide pe baza datelor.
Acesta este principiul fundamental al sistemelor multi-agent reactive.

---

### Memorie agent cu SQLite

```python
conn = sqlite3.connect("agrovision_agenti.db")
row = conn.execute("SELECT * FROM lpis_cache WHERE cod_parcela = ?", (cod,)).fetchone()

if row:
    return din_baza  # raspuns instant
else:
    date = interogheaza_lpis(cod)  # interogare reala
    conn.execute("INSERT INTO lpis_cache ...", date)
    conn.commit()
```

La prima rulare: 0.2s (interogare). La urmatoarele: <0.01s (cache).

---

### Retry automat

```python
for tentativa in range(1, max_retry + 1):
    if "drona" in rezultat and "lpis" in rezultat:
        break
    time.sleep(0.1 * tentativa)  # asteapta progresiv mai mult
```

Daca un agent mai lent nu a terminat inca, conformitatea asteapta pana la 3 incercari.

---

### Relevanta academica

Aceasta arhitectura este folosita in:
- **UAV swarms** -- mai multe drone coordonate de un orchestrator central
- **Sisteme IACS** -- agenti specializati per tip de control (teledetectie, teren, documente)
- **Smart farming** -- agenti de senzori (sol, apa, aer) + agent de decizie + agent de notificare

**Referinte:**
- Wooldridge, M. (2009). *An Introduction to MultiAgent Systems*. Wiley.
- Reg. UE 2021/2116 -- Art. 70: sisteme de monitorizare a suprafetelor agricole
""")
